#!/usr/bin/env python3
"""
Simple test to verify DOCX template generation works
"""

from docxtpl import DocxTemplate
from docx import Document
import io
import tempfile
import os

def create_simple_template():
    """Create a simple DOCX template for testing"""
    # Create a simple document
    doc = Document()
    doc.add_heading('AM AI SAFE Assessment Report', 0)
    
    # Add some template placeholders
    doc.add_paragraph('Organization: {{org.name}}')
    doc.add_paragraph('Assessment Date: {{assessment.date}}')
    doc.add_paragraph('Overall Score: {{overall.score}}%')
    doc.add_paragraph('Maturity Tier: {{overall.tier}}')
    
    # Add actions section
    doc.add_heading('High Priority Actions', level=1)
    doc.add_paragraph('{% for action in actions.high %}• {{action.text}}{% endfor %}')
    
    # Save to temporary file
    temp_path = '/tmp/simple_template.docx'
    doc.save(temp_path)
    return temp_path

def test_simple_template():
    """Test the simple template"""
    template_path = create_simple_template()
    
    try:
        # Load template
        doc = DocxTemplate(template_path)
        
        # Test data
        context = {
            'org': {'name': 'Test Organization'},
            'assessment': {'date': '2024-01-01'},
            'overall': {'score': 75, 'tier': 'Good'},
            'actions': {
                'high': [
                    {'text': 'Implement security measures'},
                    {'text': 'Update privacy policies'}
                ]
            }
        }
        
        # Render template
        doc.render(context)
        
        # Save result
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)
        
        result_bytes = output_buffer.getvalue()
        print(f"✅ Simple template test successful! Generated {len(result_bytes)} bytes")
        
        # Clean up
        os.remove(template_path)
        return True
        
    except Exception as e:
        print(f"❌ Simple template test failed: {str(e)}")
        if os.path.exists(template_path):
            os.remove(template_path)
        return False

if __name__ == "__main__":
    test_simple_template()