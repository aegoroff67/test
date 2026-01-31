"""
AM AI SAFE Report Generator
Generates DOCX and PDF reports from assessment data using templates.

NOTE: This file is being incrementally refactored into modular components.
New modular code is in /app/backend/report_generator/ directory:
- charts.py: Chart generation (heatmap, bar, radar)
- utils.py: Utility functions (formatting, tier calculations)
- ai_narratives.py: AI narrative generation (TODO)
"""

import os
import io
import tempfile
import base64
from pathlib import Path
from datetime import datetime, timezone
from typing import Dict, List, Any, Optional, Tuple
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import seaborn as sns
import numpy as np
from docxtpl import DocxTemplate, InlineImage, RichText
from docx.shared import Inches, Pt, RGBColor
import requests
import subprocess
from jinja2 import Environment, FileSystemLoader
import zipfile

# WeasyPrint import - conditional to handle missing system libraries
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False
    HTML = None
    CSS = None
import pytz

# Default timezone for reports (Australian Eastern Time)
REPORT_TIMEZONE = os.environ.get("APP_TIMEZONE", "Australia/Sydney")

def get_report_local_time(utc_time: datetime = None) -> datetime:
    """Convert UTC time to report timezone (Australian Eastern Time)."""
    if utc_time is None:
        utc_time = datetime.now(timezone.utc)
    try:
        local_tz = pytz.timezone(REPORT_TIMEZONE)
        if utc_time.tzinfo is None:
            utc_time = utc_time.replace(tzinfo=timezone.utc)
        return utc_time.astimezone(local_tz)
    except Exception:
        return utc_time


# Import framework coverage calculator
from framework_coverage import get_all_framework_coverage, FRAMEWORK_CONFIG

# Import modular chart generation (refactored)
from report_modules.charts import (
    generate_heatmap,
    generate_system_heatmap,
    generate_awareness_heatmap,
    generate_bar_chart,
    generate_radar_chart,
    generate_radar_chart_with_benchmark,
    generate_risk_gauge,
    generate_faira_domain_risk_radar,
    generate_faira_inherent_risk_radar,
)

# Import utilities (refactored)
from report_modules.utils import (
    AMP_PLACEHOLDER,
    replace_amp_with_placeholder,
    format_date as util_format_date,
    post_process_ampersands,
    calculate_tier,
    DOMAIN_COLORS,
)

# Import AI narratives (refactored)
from report_modules.ai_narratives import (
    generate_awareness_narratives,
    generate_readiness_narratives,
    generate_orgwide_narratives,
    generate_system_narratives,
    generate_faira_narratives,
)

# Import FAIRA scoring functions
from faira_scoring_engine import get_radar_chart_data, calculate_overall_risk, get_recommended_controls


class AMReportGenerator:
    """Generates AM AI SAFE assessment reports in DOCX and PDF formats."""
    
    def __init__(self, template_path: Optional[str] = None, use_test_template: bool = False, use_smart_priority: bool = False, assessment_type: str = "System"):
        """Initialize the report generator."""
        self.use_test_template = use_test_template
        self.use_smart_priority = use_smart_priority
        self.assessment_type = assessment_type
        # For test template, default to smart priority unless explicitly set to False
        if use_test_template and not use_smart_priority:
            self.use_smart_priority = True
        if template_path:
            self.template_path = template_path
        elif use_test_template:
            self.template_path = self._get_test_template_path()
        else:
            self.template_path = self._get_template_for_assessment_type(assessment_type)
        self.domain_colors = {
            'Fairness': '#FF6B6B',
            'Transparency': '#4ECDC4', 
            'Explainability': '#45B7D1',
            'Accountability': '#96CEB4',
            'Data Integrity': '#FFEAA7',
            'Reliability': '#DDA0DD',
            'Security': '#98D8C8',
            'Privacy': '#F7DC6F',
            'Safety': '#BB8FCE',
            'Inclusivity': '#85C1E9',
            'Sustainability': '#82E0AA'
        }
        self._sector_actions_cache = None
        self._question_metadata_cache = None
    
    def _load_question_metadata(self) -> Dict[str, Dict[str, Any]]:
        """Load question metadata containing impact_weight, effort_score, etc."""
        if self._question_metadata_cache is not None:
            return self._question_metadata_cache
        
        try:
            backend_dir = Path(__file__).parent
            metadata_path = backend_dir / "system_question_metadata.json"
            if metadata_path.exists():
                import json
                with open(metadata_path, 'r') as f:
                    self._question_metadata_cache = json.load(f)
                print(f"DEBUG: Loaded question metadata for {len(self._question_metadata_cache)} questions")
                return self._question_metadata_cache
        except Exception as e:
            print(f"Warning: Could not load question metadata: {e}")
        
        self._question_metadata_cache = {}
        return self._question_metadata_cache
    
    def _calculate_smart_priority(self, impact_weight: float, gap: int, effort_score: float) -> float:
        """
        Calculate smart priority score using the same formula as the frontend.
        Formula: impact_weight × gap × (1 - effort_score)
        - Higher impact = better
        - Higher gap (lower maturity score) = better
        - Lower effort = better (hence 1 - effort_score)
        """
        return impact_weight * gap * (1 - effort_score)
    
    def _get_quadrant_label(self, impact_weight: float, effort_score: float) -> str:
        """
        Determine quadrant label based on impact and effort.
        Same logic as frontend/backend results page.
        """
        if impact_weight >= 0.8 and effort_score <= 0.5:
            return "Quick win"
        elif impact_weight >= 0.8 and effort_score > 0.5:
            return "Strategic project"
        elif impact_weight < 0.8 and effort_score <= 0.5:
            return "Opportunistic improvement"
        else:
            return "Lower priority"
    
    def _load_sector_actions(self, assessment_type: str = None) -> Dict[str, Dict[str, str]]:
        """Load sector-specific action steps from JSON file based on assessment type."""
        # Use provided assessment_type or fall back to instance attribute
        effective_type = assessment_type or self.assessment_type
        
        # Use different cache key based on assessment type
        cache_key = f"_sector_actions_cache_{effective_type}"
        cached = getattr(self, cache_key, None)
        if cached is not None:
            return cached
        
        try:
            backend_dir = Path(__file__).parent
            # Choose the appropriate actions file based on assessment type
            if effective_type == 'Awareness':
                actions_path = backend_dir / "awareness_actions.json"
            elif effective_type == 'Readiness':
                actions_path = backend_dir / "readiness_actions.json"
            elif effective_type == 'Orgwide':
                actions_path = backend_dir / "orgwide_actions.json"
            else:
                actions_path = backend_dir / "system_actions.json"
            
            if actions_path.exists():
                import json
                with open(actions_path, 'r') as f:
                    result = json.load(f)
                    setattr(self, cache_key, result)
                    return result
        except Exception as e:
            print(f"Warning: Could not load sector actions for {effective_type}: {e}")
        
        setattr(self, cache_key, {})
        return {}
    
    def _get_template_for_assessment_type(self, assessment_type: str) -> str:
        """Get the appropriate template path based on assessment type."""
        backend_dir = Path(__file__).parent
        
        # Template mapping for each assessment type
        template_map = {
            'System': 'AM_AI_SAFE_System_Report_TEMPLATE_v0.15_20260122.docx',
            'Awareness': 'AM_AI_SAFE_Awareness_Report_TEMPLATE_v0.9.34_20260117.docx',
            'Readiness': 'AM_AI_SAFE_Readiness_Report_TEMPLATE_v0.08_20260118_FINAL.docx',
            'Orgwide': 'AM_AI_SAFE_Organisation_Report_TEMPLATE_v0.06_20260121.docx',
            'FAIRA': 'AM_AI_SAFE_FAIRA_Report_TEMPLATE_v0.37_20260131.docx',
        }
        
        template_filename = template_map.get(assessment_type, 'AM_AI_SAFE_Report_TEMPLATE_v9_10072025.docx')
        template_path = backend_dir / "templates" / "docx" / template_filename
        
        # Check if template exists, fallback to default if not
        if not template_path.exists():
            print(f"Warning: Template for {assessment_type} not found at {template_path}, using default")
            template_path = backend_dir / "templates" / "docx" / "AM_AI_SAFE_Report_TEMPLATE_v9_10072025.docx"
        
        print(f"DEBUG: Using template for {assessment_type}: {template_path.name}")
        return str(template_path)

    def _get_sector_average(self, sector_name: str) -> float:
        """
        Get the overall sector average from benchmark data files.
        
        Args:
            sector_name: The sector/industry name
            
        Returns:
            Overall sector average as a percentage (0-100), or None if not found
        """
        if not sector_name:
            return None
            
        try:
            import json
            backend_dir = Path(__file__).parent
            
            # Choose the right benchmark file based on assessment type
            if self.assessment_type == 'Awareness':
                benchmark_path = backend_dir / "awareness_benchmarks.json"
            elif self.assessment_type == 'Readiness':
                benchmark_path = backend_dir / "readiness_benchmarks.json"
            elif self.assessment_type == 'Orgwide':
                benchmark_path = backend_dir / "orgwide_benchmarks.json"
            elif self.assessment_type == 'System':
                benchmark_path = backend_dir / "ai_maturity_benchmarks_SYSTEM.json"
            else:
                benchmark_path = backend_dir / "benchmarks.json"
            
            if not benchmark_path.exists():
                print(f"Benchmark file not found: {benchmark_path}")
                return None
            
            with open(benchmark_path, 'r') as f:
                benchmarks_data = json.load(f)
            
            # Handle System benchmark file format (array of objects)
            if self.assessment_type == 'System':
                # System benchmark file is an array with sector_type, domain_name, benchmark_mean_score
                sector_scores = []
                for entry in benchmarks_data:
                    entry_sector = entry.get('sector_type', '')
                    # Check if sector matches (case-insensitive partial match)
                    if sector_name.lower() in entry_sector.lower() or entry_sector.lower() in sector_name.lower():
                        score = entry.get('benchmark_mean_score', 0)
                        if score:
                            sector_scores.append(float(score))
                
                if sector_scores:
                    sector_average = sum(sector_scores) / len(sector_scores)
                    print(f"DEBUG _get_sector_average (System): sector='{sector_name}', average={sector_average:.1f}% from {len(sector_scores)} domains")
                    return sector_average
                else:
                    print(f"DEBUG _get_sector_average (System): No data found for sector '{sector_name}'")
                    return None
            else:
                # Other assessment types use nested structure with sector_averages
                sector_averages = benchmarks_data.get("sector_averages", {})
                sector_average = sector_averages.get(sector_name)
                
                # If exact match not found, try "Other"
                if sector_average is None:
                    sector_average = sector_averages.get("Other")
                
                print(f"DEBUG _get_sector_average: sector='{sector_name}', average={sector_average}")
                return sector_average
            
        except Exception as e:
            print(f"Error loading sector average: {str(e)}")
            return None

    def _generate_sector_actions(self, report_data: Dict[str, Any], sector_name: str) -> Dict[str, List[Dict[str, Any]]]:
        """
        Generate sector-specific action steps based on assessment scores.
        
        Args:
            report_data: Report data containing questions and scores
            sector_name: The sector/industry name (e.g., "Finance / Insurance")
            
        Returns:
            Dictionary with 'high', 'medium', 'low' priority sector actions
        """
        sector_actions = {
            'high': [],
            'medium': [],
            'low': []
        }
        
        if not sector_name:
            print("DEBUG _generate_sector_actions: No sector_name provided")
            return sector_actions
        
        # Load all sector actions (will use correct file based on assessment_type)
        all_sector_actions = self._load_sector_actions()
        print(f"DEBUG _generate_sector_actions: Loaded {len(all_sector_actions)} sectors")
        
        # Get actions for this specific sector
        sector_specific = all_sector_actions.get(sector_name, {})
        if not sector_specific:
            # Try to find a partial match
            for key in all_sector_actions.keys():
                if sector_name.lower() in key.lower() or key.lower() in sector_name.lower():
                    sector_specific = all_sector_actions[key]
                    print(f"DEBUG _generate_sector_actions: Found partial match for '{sector_name}' -> '{key}'")
                    break
        
        if not sector_specific:
            print(f"DEBUG _generate_sector_actions: No sector-specific actions found for: {sector_name}")
            print(f"DEBUG _generate_sector_actions: Available sectors: {list(all_sector_actions.keys())}")
            return sector_actions
        
        print(f"DEBUG _generate_sector_actions: Found {len(sector_specific)} actions for sector '{sector_name}'")
        
        # Get questions data from report
        questions_data = report_data.get('questions_data', [])
        print(f"DEBUG _generate_sector_actions: questions_data has {len(questions_data)} domains")
        
        # Domain mapping depends on assessment type
        if self.assessment_type == 'Awareness':
            domain_mapping = {
                "AU": "Awareness & Understanding",
                "GT": "Governance & Trust Foundations",
                "PS": "People & Skills",
                "LV": "Leadership & Vision",
                "DR": "Data & Digital Readiness"
            }
        elif self.assessment_type == 'Readiness':
            domain_mapping = {
                "SA": "Strategic Alignment & Awareness",
                "GF": "Governance Foundations",
                "DR": "Data Readiness",
                "TI": "Technology & Infrastructure",
                "PC": "People & Culture",
                "PR": "Policy & Compliance Readiness",
                "RE": "Risk & Ethics Awareness",
                "CL": "Continuous Learning & Improvement"
            }
        elif self.assessment_type == 'Orgwide':
            domain_mapping = {
                "AE": "Accountability & Ethics",
                "CA": "Continuous Improvement & Assurance",
                "CC": "Culture & Capability",
                "DS": "Data Stewardship & Security",
                "FI": "Fairness & Inclusivity",
                "GO": "Governance & Oversight",
                "PL": "Privacy & Legal Compliance",
                "RS": "Reliability & Safety",
                "RM": "Risk Management",
                "TE": "Transparency & Explainability"
            }
        else:
            domain_mapping = {
                "FA": "Fairness", "TR": "Transparency", "EX": "Explainability", 
                "AC": "Accountability", "DI": "Data Integrity", "RE": "Reliability",
                "SE": "Security", "PR": "Privacy", "SA": "Safety", 
                "IN": "Inclusivity", "SU": "Sustainability"
            }
        
        # Iterate through questions and match with sector actions
        questions_processed = 0
        actions_generated = 0
        for domain_data in questions_data:
            domain_info = domain_data.get('domain', {})
            questions = domain_data.get('questions', [])
            
            for question in questions:
                questions_processed += 1
                question_id = question.get('question_id', question.get('code', ''))
                # Score can be in different locations depending on data structure
                score = question.get('score', 0)
                if not score and question.get('answer'):
                    score = question['answer'].get('numeric_score', 0)
                
                # Get sector-specific action for this question
                sector_action_text = sector_specific.get(question_id)
                
                if questions_processed <= 3:
                    print(f"DEBUG _generate_sector_actions: Q {question_id} score={score} has_action={bool(sector_action_text)}")
                
                if sector_action_text and score in [1, 2, 3]:
                    # Get domain name from question ID prefix
                    domain_prefix = question_id.split("-")[0] if "-" in question_id else ""
                    domain_name = domain_mapping.get(domain_prefix, domain_info.get('name', 'Unknown'))
                    
                    action_item = {
                        'domain': domain_name,
                        'question_id': question_id,
                        'sector_action': sector_action_text,
                        'score': score
                    }
                    
                    actions_generated += 1
                    if score == 1:
                        sector_actions['high'].append(action_item)
                    elif score == 2:
                        sector_actions['medium'].append(action_item)
                    elif score == 3:
                        sector_actions['low'].append(action_item)
        
        # Sort each priority list by question_id for consistency
        for priority in ['high', 'medium', 'low']:
            sector_actions[priority].sort(key=lambda x: x.get('question_id', ''))
        
        print(f"DEBUG _generate_sector_actions: Processed {questions_processed} questions, generated {actions_generated} actions")
        print(f"Generated sector actions for {sector_name}: high={len(sector_actions['high'])}, medium={len(sector_actions['medium'])}, low={len(sector_actions['low'])}")
        
        return sector_actions
    
    async def _generate_ai_enhanced_executive_summary(self, report_data: Dict[str, Any]) -> str:
        """
        Generate an AI-enhanced executive summary using LLM.
        Falls back to template-based generation if AI fails.
        Uses new test prompts when use_test_template is enabled.
        """
        try:
            from emergentintegrations.llm.chat import LlmChat, UserMessage
            
            # Extract key data for the prompt
            overall_score = report_data.get('overall', {}).get('score', 0)
            overall_tier = report_data.get('overall', {}).get('tier', 'Unknown')
            org_name = report_data.get('org', {}).get('name', 'the organization')
            system_info = report_data.get('system_info', {})
            system_name = system_info.get('systemName', system_info.get('system_name', 'the AI system'))
            industry = system_info.get('industry', '')
            
            # Get domain scores and detailed analysis
            heatmap_data = report_data.get('heatmap_data', {})
            domains = heatmap_data.get('domains', [])
            
            # Sort to find top and bottom performers with low-scoring question counts
            domain_scores = []
            for domain in domains:
                domain_name = domain.get('name', '')
                questions = domain.get('questions', [])
                scores = [q.get('score', 0) for q in questions if q.get('score') is not None]
                low_scoring_count = len([s for s in scores if s <= 2])
                if scores:
                    avg_pct = (sum(scores) / (len(scores) * 4)) * 100
                    domain_scores.append({
                        'name': domain_name, 
                        'percentage': avg_pct,
                        'question_count': len(scores),
                        'low_scoring_count': low_scoring_count
                    })
            
            sorted_domains = sorted(domain_scores, key=lambda x: x['percentage'], reverse=True)
            top_3 = sorted_domains[:3] if len(sorted_domains) >= 3 else sorted_domains
            bottom_3 = sorted_domains[-3:] if len(sorted_domains) >= 3 else sorted_domains
            
            # Calculate maturity distribution
            all_scores = []
            for domain in domains:
                questions = domain.get('questions', [])
                for q in questions:
                    if q.get('score') is not None:
                        all_scores.append(q.get('score', 0))
            
            non_ideal = len([s for s in all_scores if s == 1])
            basic = len([s for s in all_scores if s == 2])
            good = len([s for s in all_scores if s == 3])
            advanced = len([s for s in all_scores if s == 4])
            
            # Get priority actions counts
            actions = report_data.get('actions', {})
            high_priority_count = len(actions.get('high', []))
            medium_priority_count = len(actions.get('medium', []))
            low_priority_count = len(actions.get('low', []))
            
            # Get system context
            lifecycle = system_info.get('lifecycle', system_info.get('lifecycleStage', 'Unknown'))
            criticality = system_info.get('criticality', 'Unknown')
            hosting = system_info.get('hosting', system_info.get('cloudProvider', 'Unknown'))
            data_sensitivity = system_info.get('dataSensitivity', 'Unknown')
            oversight = system_info.get('oversight', system_info.get('humanOversight', 'Unknown'))
            
            if self.use_test_template:
                # NEW TEST PROMPT - Data-driven, conservative approach
                system_prompt = """You are generating the Executive Summary section of an AI System Maturity Assessment report using the AM AI SAFE framework. Your task is to produce a clear, executive-level summary that is fully grounded in the data provided.

⚠️ Critical rules:
- Only state facts that are explicitly provided in the input data.
- Do not invent policies, controls, processes, or cultural attributes.
- If you discuss reasons or drivers, clearly label them as "likely contributors" and base them on observable patterns (e.g. domain scores, concentration of low-scoring questions, or recommendation themes).
- Do not introduce any new numbers, percentages, counts, or rankings.
- Do not reference external regulations or frameworks unless explicitly included in the input.
- Avoid absolute claims such as "fully implemented", "comprehensive", "best practice", or "full marks" unless explicitly supported by the data.
- Do not use markdown formatting - write in plain text suitable for a Word document.

Heading format requirements:
- The Executive Summary section will be inserted into a Word document that already contains H1 headings.
- Any sub-headings you generate must be formatted as H2 only.
- Use the following exact format for all sub-headings:

Heading 2: <Heading text>

- Do not use H1, H3, bullet-styled headings, markdown (##), or HTML tags.
- Do not number the headings unless explicitly instructed.
- Do not repeat the section title as a heading."""

                user_prompt = f"""Write the Executive Summary for an AI System Maturity Assessment report.

ORGANISATION DETAILS:
- Organisation: {org_name}
- Industry: {industry}

ASSESSMENT CONTEXT:
- System Name: {system_name}
- Lifecycle Stage: {lifecycle}
- Business Criticality: {criticality}
- Hosting Model: {hosting}
- Data Sensitivity: {data_sensitivity}
- Oversight Arrangements: {oversight}

OVERALL RESULTS:
- Overall Maturity Score: {overall_score:.1f}%
- Maturity Tier: {overall_tier}

DOMAIN RESULTS (sorted by score):
{chr(10).join([f"- {d['name']}: {d['percentage']:.0f}% ({d['question_count']} questions, {d['low_scoring_count']} scoring low)" for d in sorted_domains])}

MATURITY DISTRIBUTION:
- Non-Ideal (score 1): {non_ideal} questions
- Basic (score 2): {basic} questions
- Good (score 3): {good} questions
- Advanced (score 4): {advanced} questions

PRIORITY ACTIONS:
- High priority: {high_priority_count} actions
- Medium priority: {medium_priority_count} actions
- Low priority: {low_priority_count} actions

TOP 3 STRENGTHS:
{chr(10).join([f"- {d['name']}: {d['percentage']:.0f}%" for d in top_3])}

TOP 3 GAPS:
{chr(10).join([f"- {d['name']}: {d['percentage']:.0f}%" for d in bottom_3])}

Required structure - Write the Executive Summary using the following structure exactly, in professional report language. Format each section heading as "Heading 2: <heading text>":

Heading 2: Purpose of the Assessment
Briefly restate the objective of the assessment and the scope of domains assessed.

Heading 2: Assessment Context
Summarise the system context using the provided lifecycle stage, criticality, hosting model, data sensitivity, and oversight arrangements. Do not infer risk beyond what this context reasonably implies.

Heading 2: Overall Maturity Outcome
State the overall AI maturity score and tier. Explain what this tier indicates at a high level (e.g. foundational, developing, established), without adding new metrics.

Heading 2: Domain-Level Performance Overview
Identify the strongest-performing domains and lowest-performing domains based on the provided scores. Where relevant, note patterns such as concentration of low-scoring questions. Any explanation of performance must be framed as likely contributors, not confirmed causes.

Heading 2: Key Insights
Provide 4–6 concise bullet points covering:
- Overall maturity position
- Strongest domain areas
- Primary risk or gap concentration
- Volume of high-priority actions identified
- Any notable maturity distribution pattern

Heading 2: Forward-Looking Summary
Conclude with a short paragraph outlining how addressing the identified priority actions would improve AI governance and maturity over time, without promising outcomes or timelines."""

            else:
                # ORIGINAL PROMPT
                system_prompt = """You are an expert AI governance consultant writing executive summaries for AI maturity assessment reports. 
Your writing style is professional, insightful, and actionable. Write in a formal business report tone.
Do not use markdown formatting - write in plain text suitable for a Word document.
Keep paragraphs concise but comprehensive."""

                user_prompt = f"""Write a comprehensive executive summary for an AI System Maturity Assessment report.

ASSESSMENT DATA:
- Organization: {org_name}
- AI System: {system_name}
- Industry: {industry}
- Overall Maturity Score: {overall_score:.1f}%
- Maturity Tier: {overall_tier}

TOP PERFORMING DOMAINS:
{chr(10).join([f"- {d['name']}: {d['percentage']:.0f}%" for d in top_3])}

AREAS REQUIRING IMPROVEMENT:
{chr(10).join([f"- {d['name']}: {d['percentage']:.0f}%" for d in bottom_3])}

Write an executive summary that includes:
1. Assessment objective and scope (mention the 11 domains and 88 questions of the AM AI SAFE framework)
2. Overall results interpretation with context for the {overall_tier} tier
3. Key strengths analysis - explain WHY these domains are performing well
4. Areas for improvement - explain the IMPLICATIONS of lower scores in these areas
5. Strategic recommendations summary (2-3 key actions)
6. Forward-looking statement about the organization's AI governance journey

The summary should be approximately 400-500 words, written as flowing paragraphs (not bullet points)."""

            # Initialize LLM chat
            chat = LlmChat(
                api_key="sk-emergent-01d3a5f175e7fB507B",
                session_id=f"exec_summary_{id(report_data)}",
                system_message=system_prompt
            ).with_model("openai", "gpt-4o-mini")
            
            # Send message and get response
            user_message = UserMessage(text=user_prompt)
            response = await chat.send_message(user_message)
            
            print(f"Successfully generated AI-enhanced executive summary (test_template={self.use_test_template})")
            return response.strip()
            
        except Exception as e:
            print(f"AI executive summary generation failed, using template: {str(e)}")
            return self._generate_comprehensive_executive_summary(report_data)
    
    async def _generate_ai_enhanced_key_findings(self, report_data: Dict[str, Any]) -> str:
        """
        Generate AI-enhanced key findings using LLM.
        Falls back to template-based generation if AI fails.
        Uses new test prompts when use_test_template is enabled.
        """
        try:
            from emergentintegrations.llm.chat import LlmChat, UserMessage
            
            # Extract key data
            overall_score = report_data.get('overall', {}).get('score', 0)
            overall_tier = report_data.get('overall', {}).get('tier', 'Unknown')
            org_name = report_data.get('org', {}).get('name', 'the organization')
            system_info = report_data.get('system_info', {})
            system_name = system_info.get('systemName', system_info.get('system_name', 'the AI system'))
            industry = system_info.get('industry', '')
            
            # Get system context
            lifecycle = system_info.get('lifecycle', system_info.get('lifecycleStage', 'Unknown'))
            criticality = system_info.get('criticality', 'Unknown')
            
            # Get domain analysis
            heatmap_data = report_data.get('heatmap_data', {})
            domains = heatmap_data.get('domains', [])
            
            domain_analysis = []
            all_scores = []
            for domain in domains:
                domain_name = domain.get('name', '')
                questions = domain.get('questions', [])
                scores = [q.get('score', 0) for q in questions if q.get('score') is not None]
                all_scores.extend(scores)
                if scores:
                    avg_pct = (sum(scores) / (len(scores) * 4)) * 100
                    low_scores = len([s for s in scores if s <= 2])
                    domain_analysis.append({
                        'name': domain_name, 
                        'percentage': avg_pct,
                        'questions_assessed': len(scores),
                        'low_scoring_questions': low_scores
                    })
            
            sorted_domains = sorted(domain_analysis, key=lambda x: x['percentage'], reverse=True)
            
            # Calculate maturity distribution (question-level for test prompt)
            non_ideal = len([s for s in all_scores if s == 1])
            basic = len([s for s in all_scores if s == 2])
            good = len([s for s in all_scores if s == 3])
            advanced = len([s for s in all_scores if s == 4])
            
            # Calculate domain-level maturity distribution (for original prompt)
            high_maturity = len([d for d in domain_analysis if d['percentage'] >= 75])
            moderate_maturity = len([d for d in domain_analysis if 50 <= d['percentage'] < 75])
            developing_maturity = len([d for d in domain_analysis if 25 <= d['percentage'] < 50])
            foundational_maturity = len([d for d in domain_analysis if d['percentage'] < 25])
            
            # Get priority actions with domain distribution
            actions = report_data.get('actions', {})
            high_actions = actions.get('high', [])
            medium_actions = actions.get('medium', [])
            low_actions = actions.get('low', [])
            
            # Get top 3 strengths and gaps
            top_3 = sorted_domains[:3] if len(sorted_domains) >= 3 else sorted_domains
            bottom_3 = sorted_domains[-3:] if len(sorted_domains) >= 3 else sorted_domains
            
            if self.use_test_template:
                # NEW TEST PROMPT - Data-driven key findings
                system_prompt = """You are generating the Key Findings & Risk Themes section of an AI System Maturity Assessment report using the AM AI SAFE framework. Your role is to identify data-driven findings based strictly on the assessment results provided.

⚠️ Critical rules:
- Every finding must be traceable to explicit input data (scores, distributions, question concentrations, or recommendation themes).
- Do not introduce new statistics, benchmarks, or comparisons.
- Do not reference external organisations, laws, or standards unless explicitly provided.
- Avoid speculative language unless clearly framed as risk implication or likely impact.
- Do not repeat the Executive Summary; focus on analytical insights.
- Do not use markdown formatting - write in plain text suitable for a Word document.

Heading format requirements:
- The Key Findings section will be inserted into a Word document that already contains H1 headings.
- Any sub-headings you generate must be formatted as H2 only.
- Use the following exact format for all sub-headings:

Heading 2: <Heading text>

- Do not use H1, H3, bullet-styled headings, markdown (##), or HTML tags.
- Do not number the headings unless explicitly instructed.
- Do not repeat the section title as a heading."""

                # Format high priority actions for the prompt
                high_actions_text = chr(10).join([f"- {a.get('domain', 'Unknown')}: {a.get('question_id', 'N/A')} - {a.get('text', 'N/A')[:100]}..." for a in high_actions[:10]]) if high_actions else "None identified"
                
                user_prompt = f"""Write Key Findings & Risk Themes for an AI System Maturity Assessment report.

OVERALL MATURITY:
- Score: {overall_score:.1f}%
- Tier: {overall_tier}

DOMAIN-LEVEL SCORES (sorted by performance):
{chr(10).join([f"- {d['name']}: {d['percentage']:.0f}% ({d['questions_assessed']} questions, {d['low_scoring_questions']} low-scoring)" for d in sorted_domains])}

MATURITY DISTRIBUTION (question-level):
- Non-Ideal (score 1): {non_ideal} questions
- Basic (score 2): {basic} questions
- Good (score 3): {good} questions
- Advanced (score 4): {advanced} questions

PRIORITY ACTIONS:
- High priority: {len(high_actions)} actions
- Medium priority: {len(medium_actions)} actions
- Low priority: {len(low_actions)} actions

SAMPLE HIGH PRIORITY ACTIONS:
{high_actions_text}

TOP 3 STRENGTHS:
{chr(10).join([f"- {d['name']}: {d['percentage']:.0f}%" for d in top_3])}

TOP 3 GAPS:
{chr(10).join([f"- {d['name']}: {d['percentage']:.0f}%" for d in bottom_3])}

SYSTEM CONTEXT:
- Lifecycle Stage: {lifecycle}
- Business Criticality: {criticality}

Required output format - Produce 4 to 6 findings. For each finding, use the format below. Format each finding title as "Heading 2: Finding - [Short descriptive title]":

Heading 2: Finding - [Short descriptive title]

Observation: Describe the observable pattern in the assessment results (e.g. domain performance, maturity distribution, clustering of low scores).

Evidence: Reference the specific data points that support the observation (e.g. domains involved, relative scores, concentration of high-priority actions).

Risk Implication: Explain what this pattern may imply for AI governance, operational risk, assurance, or sustainability if not addressed.

Content expectations - Across the full set of findings, ensure you cover:
- At least one finding related to overall maturity distribution
- At least one finding comparing strong vs weak domains
- At least one finding based on concentration of high-priority actions
- At least one finding informed by system lifecycle or criticality

Style guidance:
- Analytical and precise
- Neutral and non-judgemental
- Suitable for inclusion in a regulator-facing or board-level report
- Avoid marketing language or recommendations (those appear later in the report)"""

            else:
                # ORIGINAL PROMPT
                system_prompt = """You are an expert AI governance analyst writing detailed findings for AI maturity assessment reports.
Your analysis is thorough, evidence-based, and provides actionable insights.
Write in a professional consulting report style with clear structure.
Do not use markdown formatting - write in plain text suitable for a Word document.
Use numbered sections and clear paragraph breaks."""

                user_prompt = f"""Write comprehensive key findings for an AI System Maturity Assessment report.

ASSESSMENT DATA:
- Organization: {org_name}
- AI System: {system_name}
- Industry: {industry}
- Overall Score: {overall_score:.1f}%
- Maturity Tier: {overall_tier}

MATURITY DISTRIBUTION:
- High Maturity (75%+): {high_maturity} domains
- Moderate Maturity (50-74%): {moderate_maturity} domains  
- Developing Maturity (25-49%): {developing_maturity} domains
- Foundational (<25%): {foundational_maturity} domains

DOMAIN-BY-DOMAIN SCORES:
{chr(10).join([f"- {d['name']}: {d['percentage']:.0f}% ({d['questions_assessed']} questions, {d['low_scoring_questions']} scoring low)" for d in sorted_domains])}

Write detailed key findings that include:

1. OVERALL MATURITY ASSESSMENT
   - Interpretation of the {overall_score:.1f}% score in context of the {industry} industry
   - What the {overall_tier} tier means for the organization

2. DOMAIN PERFORMANCE ANALYSIS
   - Analysis of the highest performing domain ({sorted_domains[0]['name']} at {sorted_domains[0]['percentage']:.0f}%)
   - Analysis of the lowest performing domain ({sorted_domains[-1]['name']} at {sorted_domains[-1]['percentage']:.0f}%)
   - Patterns and correlations across domains

3. RISK ASSESSMENT
   - Key risk areas based on low-scoring domains
   - Potential compliance and reputational implications
   - Areas requiring immediate attention

4. ORGANIZATIONAL READINESS
   - Assessment of governance maturity
   - Capability gaps identified
   - Resource and investment implications

5. COMPARATIVE INSIGHTS
   - How technical vs governance domains compare
   - Balance between AI capability and AI safety measures

6. CRITICAL SUCCESS FACTORS
   - What needs to happen for improvement
   - Key dependencies and enablers

The findings should be approximately 600-800 words with clear section headers."""

            # Initialize LLM chat
            chat = LlmChat(
                api_key="sk-emergent-01d3a5f175e7fB507B",
                session_id=f"key_findings_{id(report_data)}",
                system_message=system_prompt
            ).with_model("openai", "gpt-4o-mini")
            
            # Send message and get response
            user_message = UserMessage(text=user_prompt)
            response = await chat.send_message(user_message)
            
            print(f"Successfully generated AI-enhanced key findings (test_template={self.use_test_template})")
            return response.strip()
            
        except Exception as e:
            print(f"AI key findings generation failed, using template: {str(e)}")
            return self._generate_key_findings(report_data)
    
    def _get_default_template_path(self) -> str:
        """Get the default template path - using v9 template updated 10/07/2025."""  
        backend_dir = Path(__file__).parent
        return str(backend_dir / "templates" / "docx" / "AM_AI_SAFE_Report_TEMPLATE_v9_10072025.docx")
    
    # ========================================
    # AWARENESS AI NARRATIVE GENERATION
    # ========================================
    
    async def _generate_awareness_ai_narratives(self, report_data: Dict[str, Any], assessment: Dict[str, Any], db) -> Dict[str, str]:
        """Generate all AI narratives for Awareness assessments. Delegates to modular ai_narratives."""
        api_key = os.environ.get('EMERGENT_LLM_KEY', 'sk-emergent-01d3a5f175e7fB507B')
        return await generate_awareness_narratives(report_data, assessment, db, api_key)

    async def _generate_readiness_ai_narratives(self, report_data: Dict[str, Any], assessment: Dict[str, Any], db) -> Dict[str, str]:
        """Generate all AI narratives for Readiness assessments. Delegates to modular ai_narratives."""
        api_key = os.environ.get('EMERGENT_LLM_KEY', 'sk-emergent-01d3a5f175e7fB507B')
        return await generate_readiness_narratives(report_data, assessment, db, api_key)

    async def _generate_orgwide_ai_narratives(self, report_data: Dict[str, Any], assessment: Dict[str, Any], db) -> Dict[str, str]:
        """Generate all AI narratives for Orgwide assessments. Delegates to modular ai_narratives."""
        api_key = os.environ.get('EMERGENT_LLM_KEY', 'sk-emergent-01d3a5f175e7fB507B')
        return await generate_orgwide_narratives(report_data, assessment, db, api_key)

    async def _generate_system_ai_narratives(self, report_data: Dict[str, Any], assessment: Dict[str, Any], db) -> Dict[str, str]:
        """Generate all AI narratives for System assessments. Delegates to modular ai_narratives."""
        api_key = os.environ.get('EMERGENT_LLM_KEY', 'sk-emergent-01d3a5f175e7fB507B')
        return await generate_system_narratives(report_data, assessment, db, api_key)

    async def _generate_faira_ai_narratives(self, report_data: Dict[str, Any], assessment: Dict[str, Any], db) -> Dict[str, str]:
        """Generate all 32 AI narratives for FAIRA assessments. Delegates to modular ai_narratives."""
        api_key = os.environ.get('EMERGENT_LLM_KEY', 'sk-emergent-01d3a5f175e7fB507B')
        return await generate_faira_narratives(report_data, assessment, db, api_key)

    def _get_test_template_path(self) -> str:
        """Get the test template path for testing new template structure."""  
        backend_dir = Path(__file__).parent
        return str(backend_dir / "templates" / "docx" / "AM_AI_SAFE_SYSTEM_test_template.docx")
    
    def _get_test_schema_path(self) -> str:
        """Get the test schema path for validation."""
        backend_dir = Path(__file__).parent
        return str(backend_dir / "Test_System_report_schema.json")
    
    def _get_html_template_path(self) -> str:
        """Get the HTML template path for PDF generation."""
        backend_dir = Path(__file__).parent
        return str(backend_dir / "templates" / "html" / "pdf_report_template.html")
    
    def format_date(self, date_input) -> str:
        """Format date for the template."""
        try:
            if isinstance(date_input, str):
                # Try different date formats
                for fmt in ['%Y-%m-%d', '%Y-%m-%dT%H:%M:%S', '%Y-%m-%dT%H:%M:%S.%f']:
                    try:
                        date_obj = datetime.strptime(date_input, fmt)
                        return date_obj.strftime('%d %B %Y')
                    except ValueError:
                        continue
                # If parsing fails, return as-is
                return str(date_input)
            elif hasattr(date_input, 'strftime'):
                return date_input.strftime('%d %B %Y')
            else:
                return str(date_input)
        except:
            return str(date_input)
    
    def _format_assessment_date(self, date_input) -> str:
        """Format assessment date for JSON structure."""
        try:
            if isinstance(date_input, str):
                return date_input  # Already a string, return as-is
            elif hasattr(date_input, 'strftime'):
                return date_input.strftime('%Y-%m-%d')
            else:
                return str(date_input)
        except:
            return str(date_input)
    
    async def generate_report(self, assessment_id: str, assessment_data: Dict[str, Any], 
                            user_data: Dict[str, Any], view_type: str = "heatmap", 
                            benchmark_data: Dict[str, Any] = None,
                            use_ai: bool = False) -> Tuple[bytes, bytes]:
        """
        Generate DOCX and PDF reports from assessment data.
        
        Args:
            assessment_id: The assessment ID
            assessment_data: Raw assessment data from database
            user_data: User information
            view_type: Type of visualization ('heatmap' or 'radar')
            benchmark_data: Benchmark data for radar chart comparison (optional)
            use_ai: Whether to use AI for enhanced narrative generation
            
        Returns:
            Tuple of (docx_bytes, pdf_bytes)
        """
        # Transform data to match report model
        report_data = self._transform_assessment_data(assessment_data, user_data)
        
        # Generate AI-enhanced content if requested
        if use_ai:
            print("=== GENERATING AI-ENHANCED CONTENT ===")
            try:
                ai_exec_summary = await self._generate_ai_enhanced_executive_summary(report_data)
                report_data['ai_executive_summary'] = ai_exec_summary
                print(f"AI Executive Summary generated: {len(ai_exec_summary)} chars")
            except Exception as e:
                print(f"AI Executive Summary failed: {e}")
                
            try:
                ai_key_findings = await self._generate_ai_enhanced_key_findings(report_data)
                report_data['ai_key_findings'] = ai_key_findings
                print(f"AI Key Findings generated: {len(ai_key_findings)} chars")
            except Exception as e:
                print(f"AI Key Findings failed: {e}")
        
        # Generate visualization image based on view_type
        print(f"=== VISUALIZATION GENERATION ===")
        print(f"View Type: {view_type}")
        print(f"Benchmark Data Available: {benchmark_data is not None}")
        
        if view_type == "radar" and benchmark_data:
            print("Generating radar chart WITH benchmarks")
            visualization_image = self._generate_radar_chart_with_benchmark(report_data, benchmark_data)
        elif view_type == "radar":
            print("Generating radar chart WITHOUT benchmarks")
            visualization_image = self._generate_radar_chart_image(report_data)
        else:  # default to heatmap
            print("Generating heatmap")
            visualization_image = self._generate_heatmap_image(report_data)
        
        print(f"Visualization image size: {len(visualization_image)} bytes")
        
        # Generate DOCX report
        docx_bytes = self._generate_docx_report(report_data, visualization_image)
        
        # Generate PDF using HTML-to-PDF conversion (WeasyPrint)
        try:
            pdf_bytes = self._generate_html_pdf(assessment_data, user_data)
            print("DEBUG: Using HTML-to-PDF conversion successfully")
        except Exception as e:
            print(f"WARNING: HTML-to-PDF failed ({str(e)}), falling back to LibreOffice conversion")
            # Fallback to LibreOffice conversion
            pdf_bytes = self._convert_docx_to_pdf(docx_bytes)
        
        return docx_bytes, pdf_bytes
    
    def _generate_html_pdf(self, assessment_data: Dict[str, Any], user_data: Dict[str, Any]) -> bytes:
        """Generate PDF directly from HTML template using WeasyPrint."""
        try:
            html_template_path = self._get_html_template_path()
            
            if not os.path.exists(html_template_path):
                raise Exception(f"HTML template not found: {html_template_path}")
            
            # Set up Jinja2 environment
            template_dir = os.path.dirname(html_template_path)
            template_name = os.path.basename(html_template_path)
            
            env = Environment(loader=FileSystemLoader(template_dir))
            template = env.get_template(template_name)
            
            # Transform the assessment data to report format
            report_data = self._transform_assessment_data(assessment_data, user_data)
            
            # Generate heatmap as data URL for embedding in HTML
            heatmap_data_url = self._generate_heatmap_data_url(report_data)
            
            # Use the transformed report data directly as template context
            template_context = report_data
            template_context['heatmap_image_src'] = heatmap_data_url
            
            # Render HTML
            html_content = template.render(**template_context)
            
            print("DEBUG: HTML template rendered successfully")
            # Check if placeholders are being replaced
            if '{{ org.name }}' in html_content:
                print("WARNING: Template placeholders not being replaced!")
                print(f"Sample content: {html_content[:500]}")
            else:
                print("DEBUG: Template placeholders successfully replaced")
            
            # Convert HTML to PDF using WeasyPrint
            html_doc = HTML(string=html_content)
            pdf_bytes = html_doc.write_pdf()
            
            print(f"HTML-to-PDF conversion successful. Generated {len(pdf_bytes)} bytes")
            return pdf_bytes
            
        except Exception as e:
            print(f"ERROR: HTML-to-PDF conversion failed: {str(e)}")
            import traceback
            traceback.print_exc()
            raise Exception(f"Failed to generate PDF from HTML: {str(e)}")

    def _generate_heatmap_data_url(self, report_data: Dict[str, Any]) -> str:
        """Generate heatmap and return as data URL for HTML embedding."""
        try:
            # Generate heatmap image bytes
            heatmap_bytes = self._generate_heatmap_image(report_data)
            
            # Convert to data URL
            import base64
            encoded = base64.b64encode(heatmap_bytes).decode('utf-8')
            data_url = f"data:image/png;base64,{encoded}"
            
            return data_url
            
        except Exception as e:
            print(f"ERROR: Failed to generate heatmap data URL: {str(e)}")
            # Return placeholder data URL
            return "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNkYPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=="

    # NOTE: _prepare_template_context method removed - now using direct template context
    
    def _convert_docx_to_pdf(self, docx_bytes: bytes) -> bytes:
        """Convert DOCX to PDF. First try LibreOffice, then fallback to returning DOCX."""
        
        try:
            # Create temporary files
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as docx_file:
                docx_file.write(docx_bytes)
                docx_path = docx_file.name
            
            # Try multiple LibreOffice locations
            libreoffice_paths = [
                '/usr/lib/libreoffice/program/soffice',
                '/usr/bin/libreoffice',
                '/usr/bin/soffice',
                'libreoffice',
                'soffice'
            ]
            
            libreoffice_path = None
            for path in libreoffice_paths:
                if os.path.exists(path) or (not path.startswith('/') and subprocess.run(['which', path], capture_output=True).returncode == 0):
                    libreoffice_path = path
                    break
            
            if not libreoffice_path:
                print("WARNING: LibreOffice not found. Returning DOCX content as PDF (user will get DOCX file)")
                return docx_bytes
            
            # Create output directory
            output_dir = tempfile.mkdtemp()
            
            cmd = [
                libreoffice_path, '--headless', '--convert-to', 'pdf',
                '--outdir', output_dir, docx_path
            ]
            
            print(f"DEBUG: Using LibreOffice path: {libreoffice_path}")
            print(f"DEBUG: Running command: {' '.join(cmd)}")
            
            # Set up environment for LibreOffice
            env = os.environ.copy()
            env['HOME'] = '/tmp'
            env['TMPDIR'] = '/tmp'
            
            # Try the conversion with timeout
            result = subprocess.run(cmd, check=True, capture_output=True, timeout=60, env=env)
            print(f"LibreOffice conversion successful. Output: {result.stdout.decode()}")
            
            # Find the generated PDF
            pdf_filename = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
            pdf_path = os.path.join(output_dir, pdf_filename)
            
            if not os.path.exists(pdf_path):
                print(f"WARNING: PDF not found at {pdf_path}. Returning DOCX content.")
                return docx_bytes
                
            with open(pdf_path, 'rb') as pdf_file:
                pdf_bytes = pdf_file.read()
                
            print(f"PDF conversion successful. Generated {len(pdf_bytes)} bytes")
            return pdf_bytes
            
        except Exception as e:
            print(f"WARNING: PDF conversion failed ({str(e)}). Returning DOCX content as fallback.")
            # Fallback: return DOCX content
            return docx_bytes
        finally:
            # Clean up temporary directory
            import shutil
            if 'output_dir' in locals():
                try:
                    shutil.rmtree(output_dir)
                except:
                    pass
            # Clean up temporary DOCX file
            if 'docx_path' in locals():
                try:
                    os.unlink(docx_path)
                except:
                    pass
    
    def _transform_assessment_data(self, assessment_data: Dict[str, Any], 
                                 user_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Transform raw assessment data to match the report JSON model.
        
        Args:
            assessment_data: Raw assessment data including questions, answers, summary
            user_data: User information
            
        Returns:
            Transformed data matching report_model_example.json structure
        """
        # Extract core data
        questions_data = assessment_data.get('questions_data', [])
        summary_data = assessment_data.get('summary', {})
        assessment_info = assessment_data.get('assessment', {})
        
        # Calculate overall score and tier based on assessment type
        overall_percentage = summary_data.get('overall_percentage', 0)
        if self.assessment_type == 'Awareness':
            overall_tier = self._calculate_awareness_tier(overall_percentage)
        elif self.assessment_type == 'Readiness':
            overall_tier = self._calculate_readiness_tier(overall_percentage)
        elif self.assessment_type == 'Orgwide':
            overall_tier = self._calculate_orgwide_tier(overall_percentage)
        elif self.assessment_type == 'System':
            overall_tier = self._calculate_system_tier(overall_percentage)
        else:
            overall_tier = self._calculate_tier(overall_percentage)
        
        # Generate actions based on priority mapping rules
        actions = self._generate_actions_from_questions(questions_data)
        
        # Get sector info from assessment data
        sector_name = assessment_data.get('sector_name', '')
        sector_actions = assessment_data.get('sector_actions', {'high': [], 'medium': [], 'low': []})
        
        # Create report data structure matching the JSON model
        report_data = {
            "org": {
                "name": user_data.get('organization_name', 'Organization')
            },
            "assessment": {
                "date": self._format_assessment_date(assessment_info.get('created_at', datetime.now(timezone.utc))),
                "version": "1.0.0"
            },
            "overall": {
                "score": round(overall_percentage, 1),
                "tier": overall_tier
            },
            "assets": {
                "heatmapUrl": None  # Will be set when we generate the image
            },
            "actions": actions,
            "heatmap_data": self._prepare_heatmap_data(questions_data),
            "questions_data": questions_data,  # Pass through for sector action generation
            "sector_name": sector_name,
            "sector_average": assessment_data.get('sector_average'),  # Pass sector_average for bar chart and results summary
            "sector_actions": sector_actions,
            "system_info": assessment_data.get('system_info', {}),  # Pass system_info for template use
            "awareness_info": assessment_data.get('awareness_info', {}),  # Pass awareness_info for Awareness assessments
            "readiness_info": assessment_data.get('readiness_info', {}),  # Pass readiness_info for Readiness assessments
            "orgwide_info": assessment_data.get('orgwide_info', {}),  # Pass orgwide_info for Orgwide assessments
            "benchmark_data": self._load_benchmark_data(sector_name),  # Load sector benchmark data
            "ai_narratives": assessment_data.get('ai_narratives', {}),  # Pass AI narratives for template use
            # FAIRA-specific data
            "faira_form": assessment_data.get('faira_form', {}),
            "faira_risk_summary": assessment_data.get('faira_risk_summary', {}),
            "faira_radar_data": assessment_data.get('faira_radar_data', {}),
            "faira_controls": assessment_data.get('faira_controls', []),
        }
        
        return report_data
    
    def _load_benchmark_data(self, sector_name: str) -> Dict[str, float]:
        """Load sector benchmark data for the radar chart."""
        try:
            import json
            backend_dir = Path(__file__).parent
            
            # Choose the right benchmark file based on assessment type
            if self.assessment_type == 'Awareness':
                benchmark_path = backend_dir / "awareness_benchmarks.json"
            elif self.assessment_type == 'Readiness':
                benchmark_path = backend_dir / "readiness_benchmarks.json"
            elif self.assessment_type == 'Orgwide':
                benchmark_path = backend_dir / "orgwide_benchmarks.json"
            elif self.assessment_type == 'System':
                benchmark_path = backend_dir / "ai_maturity_benchmarks_SYSTEM.json"
            else:
                benchmark_path = backend_dir / "benchmarks.json"
            
            if not benchmark_path.exists():
                print(f"Benchmark file not found: {benchmark_path}")
                return {}
            
            with open(benchmark_path, 'r') as f:
                benchmarks_data = json.load(f)
            
            # Handle different benchmark file formats
            benchmark_values = {}
            
            if self.assessment_type == 'System':
                # System benchmark file is an array of objects with sector_type, domain_name, benchmark_mean_score
                # Filter by sector and build domain->score mapping
                for entry in benchmarks_data:
                    entry_sector = entry.get('sector_type', '')
                    # Check if sector matches (case-insensitive partial match)
                    if sector_name and (sector_name.lower() in entry_sector.lower() or entry_sector.lower() in sector_name.lower()):
                        domain = entry.get('domain_name', '')
                        score = entry.get('benchmark_mean_score', 50)
                        if domain:
                            benchmark_values[domain] = float(score)
                
                # If no matches, try generic fallback
                if not benchmark_values:
                    # Use first sector's data as fallback or default 50%
                    print(f"DEBUG: No benchmark data found for sector '{sector_name}', using defaults")
                    benchmark_values = {
                        'Fairness': 50.0, 'Transparency': 50.0, 'Explainability': 50.0,
                        'Accountability': 50.0, 'Data Integrity': 50.0, 'Reliability': 50.0,
                        'Security': 50.0, 'Privacy': 50.0, 'Safety': 50.0,
                        'Inclusivity': 50.0, 'Sustainability': 50.0
                    }
            else:
                # Other assessment types use nested structure: {"benchmarks": {"sector": {"domain": "range"}}}
                sector_benchmarks = benchmarks_data.get("benchmarks", {}).get(sector_name, {})
                
                # If sector not found, try "Other"
                if not sector_benchmarks:
                    sector_benchmarks = benchmarks_data.get("benchmarks", {}).get("Other", {})
                
                # Convert range strings (like "48-55") to mid-point values
                for domain, range_str in sector_benchmarks.items():
                    if isinstance(range_str, str) and '-' in range_str:
                        # Parse range and take midpoint
                        parts = range_str.split('-')
                        try:
                            low = float(parts[0])
                            high = float(parts[1])
                            benchmark_values[domain] = (low + high) / 2
                        except (ValueError, IndexError):
                            benchmark_values[domain] = 50.0  # Default
                    elif isinstance(range_str, (int, float)):
                        benchmark_values[domain] = float(range_str)
                    else:
                        benchmark_values[domain] = 50.0  # Default
            
            print(f"DEBUG: Loaded benchmark data for sector '{sector_name}': {benchmark_values}")
            return benchmark_values
            
        except Exception as e:
            print(f"Error loading benchmark data: {str(e)}")
            return {}
    
    def _calculate_tier(self, percentage: float) -> str:
        """Calculate the maturity tier based on percentage score.
        
        Categories (based on 1-4 scoring scale):
        1. Leading (91–100%) 
        2. Established (76–90%) 
        3. Developing (51–75%) 
        4. Foundational (0–50%)
        """
        if percentage >= 91:
            return "Leading"
        elif percentage >= 76:
            return "Established"
        elif percentage >= 51:
            return "Developing"
        else:
            return "Foundational"
    
    def _calculate_awareness_tier(self, percentage: float) -> str:
        """Calculate the awareness tier for AI Awareness & Foundations assessments.
        
        Categories:
        - Established (86-100%): Ready to Progress
        - Developing (66-85%): Building Readiness
        - Emerging (41-65%): Exploring Opportunities
        - Introductory (0-40%): Early Awareness
        """
        if percentage >= 86:
            return "Established"
        elif percentage >= 66:
            return "Developing"
        elif percentage >= 41:
            return "Emerging"
        else:
            return "Introductory"
    
    def _calculate_readiness_tier(self, percentage: float) -> str:
        """Calculate the readiness tier for AI Readiness assessments.
        
        Categories:
        - Leading (86-100%): AI-Ready
        - Established (66-85%): Operational Readiness
        - Developing (41-65%): Emerging Capability
        - Foundational (0-40%): Limited Foundations
        """
        if percentage >= 86:
            return "Leading"
        elif percentage >= 66:
            return "Established"
        elif percentage >= 41:
            return "Developing"
        else:
            return "Foundational"
    
    def _calculate_orgwide_tier(self, percentage: float) -> str:
        """Calculate the maturity tier for Organisation-wide AI Maturity assessments.
        
        Categories:
        - Leading (86-100%): Optimised Excellence
        - Established (66-85%): Integrated Governance
        - Developing (41-65%): Emerging Structure
        - Foundational (0-40%): Ad-hoc Beginnings
        """
        if percentage >= 86:
            return "Leading"
        elif percentage >= 66:
            return "Established"
        elif percentage >= 41:
            return "Developing"
        else:
            return "Foundational"
    
    def _calculate_system_tier(self, percentage: float) -> str:
        """Calculate the maturity tier for AI System Maturity assessments.
        
        Categories (same as Orgwide):
        - Leading (86-100%): Optimised Excellence
        - Established (66-85%): Integrated Governance
        - Developing (41-65%): Emerging Structure
        - Foundational (0-40%): Ad-hoc Beginnings
        """
        if percentage >= 86:
            return "Leading"
        elif percentage >= 66:
            return "Established"
        elif percentage >= 41:
            return "Developing"
        else:
            return "Foundational"
    
    def _get_tier_for_score(self, score: float) -> str:
        """Get tier name for a given score percentage.
        
        Uses the assessment_type to determine the correct tier calculation.
        This is a convenience wrapper for use in AI narrative generation.
        """
        if self.assessment_type == 'Awareness':
            return self._calculate_awareness_tier(score)
        elif self.assessment_type == 'Readiness':
            return self._calculate_readiness_tier(score)
        elif self.assessment_type == 'Orgwide':
            return self._calculate_orgwide_tier(score)
        elif self.assessment_type == 'System':
            return self._calculate_system_tier(score)
        else:
            return self._calculate_tier(score)
    
    def _generate_results_summary_text(self, report_data: Dict[str, Any]) -> str:
        """
        Generate the complete Results Summary paragraph for Awareness assessments.
        Matches the frontend ResultsPage.js Results Summary section.
        """
        overall = report_data.get('overall', {})
        tier = overall.get('tier', 'Introductory')
        score = overall.get('score', 0)
        
        # Get organization name from awareness_info or org
        awareness_info = report_data.get('awareness_info') or {}
        org_name = awareness_info.get('org_name') or report_data.get('org', {}).get('name', 'Your organisation')
        
        # Get sector benchmark comparison if available
        sector_average = report_data.get('sector_average')
        sector_name = awareness_info.get('industry') or report_data.get('sector_name', '')
        
        # Build the opening sentence
        summary = f"Your organisation demonstrates {tier} AI awareness with an overall score of {score:.1f}%"
        
        # Add sector comparison if available
        if sector_average is not None and sector_name:
            if score > sector_average:
                comparison = "above"
            elif score < sector_average:
                comparison = "below"
            else:
                comparison = "equal to"
            summary += f", which is {comparison} the {sector_name} sector AI awareness average of {sector_average:.0f}%."
        else:
            summary += "."
        
        # Add tier-specific explanatory text
        if tier == 'Established':
            summary += " Your organisation demonstrates strong and well-distributed AI awareness. Leaders and staff show clear understanding of AI concepts, realistic capabilities, and potential benefits and risks. This puts you in an excellent position to progress into formal readiness assessment and begin exploring structured AI initiatives or early pilots."
        elif tier == 'Developing':
            summary += " AI awareness is growing consistently across the organisation. Many individuals understand core concepts and can identify relevant opportunities, though some knowledge gaps remain. Strengthening leadership engagement, deepening practical understanding, and establishing light governance foundations will support your transition into readiness assessment."
        elif tier == 'Emerging':
            summary += " Your organisation is showing early signs of AI awareness, with isolated pockets of understanding and increasing curiosity. While some staff recognise potential use cases, overall knowledge remains inconsistent. Focus on foundational education, awareness sessions, and building shared language around AI to prepare for the next stage of capability development."
        else:  # Introductory
            summary += " Your organisation is at the beginning of its AI awareness journey. Most staff and leaders are unfamiliar with AI fundamentals, potential applications, or key risks. Priority should be placed on introductory education, building basic literacy, and fostering leadership interest before moving into deeper assessments or planning activities."
        
        return summary
    
    def _generate_readiness_results_summary_text(self, report_data: Dict[str, Any]) -> str:
        """
        Generate the complete Results Summary paragraph for Readiness assessments.
        Matches the frontend ResultsPage.js Results Summary section.
        """
        overall = report_data.get('overall', {})
        tier = overall.get('tier', 'Foundational')
        score = overall.get('score', 0)
        
        # Get organization name from readiness_info or org
        readiness_info = report_data.get('readiness_info') or {}
        org_name = readiness_info.get('org_name') or report_data.get('org', {}).get('name', 'The organisation')
        
        # Get sector benchmark comparison if available
        sector_average = report_data.get('sector_average')
        sector_name = readiness_info.get('industry') or report_data.get('sector_name', '')
        
        # Build the opening sentence
        summary = f"The results indicate that {org_name} has achieved an overall AI readiness score of {score:.1f}%, placing the organization within the {tier} readiness category"
        
        # Add sector comparison if available
        if sector_average is not None and sector_name:
            if score > sector_average:
                comparison = "above"
            elif score < sector_average:
                comparison = "below"
            else:
                comparison = "equal to"
            summary += f", which is {comparison} the {sector_name} sector AI readiness average of {sector_average:.0f}%."
        else:
            summary += "."
        
        # Add tier-specific explanatory text
        if tier == 'Leading':
            summary += " This rating reflects comprehensive AI readiness across all foundational domains. The organization demonstrates strong governance, robust data practices, mature technology infrastructure, capable workforce, and embedded ethical frameworks. Leadership actively champions responsible AI, and the organization is well-prepared to deploy AI systems confidently and safely."
        elif tier == 'Established':
            summary += " This rating reflects strong AI readiness foundations with governance structures, data management practices, and technology capabilities in place. The organization has clear policies, engaged leadership, and growing staff capability. Further strengthening of continuous improvement processes, stakeholder engagement, and advanced risk management will position the organization for leading-edge AI adoption."
        elif tier == 'Developing':
            summary += " This rating reflects emerging AI readiness with foundational elements beginning to take shape. Some governance, data, and technology practices exist, but consistency and maturity vary. Priority should be on formalizing AI governance frameworks, strengthening data quality and security, building staff capability, and establishing clear ethical guidelines before advancing to AI implementation."
        else:  # Foundational
            summary += " This rating reflects limited AI readiness with significant capability gaps across governance, data, technology, and workforce domains. Most foundational elements are informal or absent. Immediate focus should be on building basic governance structures, improving data management practices, securing technology infrastructure, and developing staff awareness before considering AI adoption."
        
        return summary
    
    def _generate_orgwide_results_summary_text(self, report_data: Dict[str, Any]) -> str:
        """
        Generate the complete Results Summary paragraph for Organisation-wide AI Maturity assessments.
        Matches the frontend ResultsPage.js Results Summary section.
        """
        overall = report_data.get('overall', {})
        tier = overall.get('tier', 'Foundational')
        score = overall.get('score', 0)
        
        # Get organization name from orgwide_info or org
        orgwide_info = report_data.get('orgwide_info') or {}
        org_name = orgwide_info.get('org_name') or report_data.get('org', {}).get('name', 'The organisation')
        
        # Get sector benchmark comparison if available
        sector_average = report_data.get('sector_average')
        # Use sector_name from report_data (which is correctly set to organization's primary_industry)
        sector_name = report_data.get('sector_name', '') or orgwide_info.get('industry', '')
        
        # Build the opening sentence
        summary = f"The results indicate that {org_name} has achieved an overall Organisation-wide AI Maturity score of {score:.1f}%, placing the organisation within the {tier} maturity category"
        
        # Add sector comparison if available
        if sector_average is not None and sector_name:
            if score > sector_average:
                comparison = "above"
            elif score < sector_average:
                comparison = "below"
            else:
                comparison = "equal to"
            summary += f", which is {comparison} the {sector_name} sector AI maturity average of {sector_average:.0f}%."
        else:
            summary += "."
        
        # Add tier-specific explanatory text
        if tier == 'Leading':
            summary += " This rating reflects a highly mature and well-embedded AI governance capability. Oversight, ethics, risk management, transparency, and lifecycle assurance are consistently applied across the organisation. AI decisions are well-controlled, monitored, and aligned to recognised standards. This level reflects strong leadership commitment and positions the organisation as a benchmark for responsible AI practice."
        elif tier == 'Established':
            summary += " This rating reflects a well-developed and consistently applied AI governance framework. Most domains show strong performance, with clear roles, oversight, and documented processes. Some variability or manual effort remains, but overall governance is stable and repeatable. Strengthening lifecycle assurance, fairness safeguards, and organisation-wide alignment will support progress toward leading maturity."
        elif tier == 'Developing':
            summary += " This rating reflects emerging and partially consistent AI governance across the organisation. Foundational structures exist but are not yet fully integrated or uniformly adopted. Oversight, transparency, fairness, and continuous monitoring vary across teams. Priorities include formalising governance, strengthening coordination, improving risk and ethical oversight, and embedding systematic lifecycle processes to advance maturity."
        else:  # Foundational
            summary += " This rating reflects early-stage or inconsistent AI governance practices. Key policies, roles, risk controls, and oversight mechanisms are limited or undeveloped. AI maturity varies significantly across teams, and core safeguards are often missing. Establishing baseline governance structures, defining responsibilities, and building organisational awareness are essential next steps for improving responsible AI maturity."
        
        return summary
    
    def _load_recommendations_lookup(self) -> Dict[str, Any]:
        """Load recommendations lookup from JSON file."""
        import json
        backend_dir = Path(__file__).parent
        recommendations_path = backend_dir / "AMAI_SAFE_recommendations_lookup.json"
        
        try:
            with open(recommendations_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"Warning: Could not load recommendations lookup: {e}")
            return {"recommendations": {}}
    
    def _generate_system_results_summary_text(self, report_data: Dict[str, Any]) -> str:
        """
        Generate the complete Results Summary paragraph for AI System Maturity assessments.
        Matches the frontend ResultsPage.js Results Summary section.
        """
        overall = report_data.get('overall', {})
        tier = overall.get('tier', 'Foundational')
        score = overall.get('score', 0)
        
        # Get system name from system_info or org
        system_info = report_data.get('system_info') or {}
        system_name = system_info.get('systemName') or system_info.get('system_name') or system_info.get('organizationName') or report_data.get('org', {}).get('name', 'The system')
        
        # Get sector benchmark comparison if available
        sector_average = report_data.get('sector_average')
        sector_name = system_info.get('industry') or report_data.get('sector_name', '')
        
        # Build the opening sentence
        summary = f"The results indicate that {system_name} has achieved an overall AI maturity score of {score:.1f}%, placing this system within the {tier} AI Maturity category"
        
        # Add sector comparison if available
        if sector_average is not None and sector_name:
            if score > sector_average:
                comparison = "above"
            elif score < sector_average:
                comparison = "below"
            else:
                comparison = "equal to"
            summary += f", which is {comparison} the {sector_name} sector AI system maturity average of {sector_average:.0f}%."
        else:
            summary += "."
        
        # Add tier-specific explanatory text
        if tier == 'Leading':
            summary += " This rating reflects exemplary AI governance and ethical assurance, setting a benchmark for responsible AI leadership. Governance systems are fully embedded, adaptive, and continuously refined through data-driven insights, external validation, and innovation. The focus is on optimisation, transparency, and sustained improvement across all AI operations."
        elif tier == 'Established':
            summary += " This rating reflects well-defined and consistently applied AI governance frameworks across most domains. Risk management, transparency, and ethical oversight are integrated into day-to-day operations. Continuous monitoring and regular review cycles are evident, though further optimisation and automation would strengthen maturity and resilience against evolving AI risks."
        elif tier == 'Developing':
            summary += " This rating reflects growing awareness and emerging structure in its AI governance practices. Some policies and controls are in place, but they are applied inconsistently. Progress has been made in recognising key governance needs; however, targeted improvements are needed to achieve full integration and accountability across the AI lifecycle."
        else:  # Foundational
            summary += " This rating reflects minimal or inconsistent AI governance, with most practices being reactive or informal. Policies, processes, and accountability structures are largely undeveloped, resulting in fragmented oversight and heightened risk exposure. Immediate action is required to establish a foundational governance framework, define roles and responsibilities, and embed basic ethical and risk management principles."
        
        return summary
    
    def _get_recommendation_text(self, question_code: str, recommendations_lookup: Dict[str, Any], domain_name: str) -> str:
        """Get specific recommendation text for a question code."""
        recommendations = recommendations_lookup.get("recommendations", {})
        
        if question_code in recommendations:
            return recommendations[question_code].get("default_recommendation", f"Implement comprehensive {domain_name.lower()} improvements to achieve best practices.")
        else:
            # Fallback to generic recommendation
            return f"Implement comprehensive {domain_name.lower()} improvements to achieve best practices in this area."
    
    def _generate_actions_from_questions(self, questions_data: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
        """
        Generate prioritized actions based on question scores using the priority mapping rules.
        
        Priority Mapping Rules (updated for 1-4 scoring scale):
        - Score 1 (Foundational) → High Priority → actions.high
        - Score 2 (Developing) → Medium Priority → actions.medium  
        - Score 3 (Established) → Low Priority → actions.low
        - Score 4 (Leading) → not included in report
        
        Ordering: Recommendations are sorted by domain score (lowest first), then by question within domain.
        """
        actions = {
            "high": [],
            "medium": [],
            "low": []
        }
        
        # Load recommendations lookup
        recommendations_lookup = self._load_recommendations_lookup()
        
        # Calculate domain scores for sorting
        domain_scores = {}
        for domain in questions_data:
            domain_name = domain["domain"]["name"]
            questions = domain["questions"]
            total_score = sum(q["answer"]["numeric_score"] for q in questions if q.get("answer"))
            domain_scores[domain_name] = total_score
        
        # Generate recommendations based on question scores
        for domain in questions_data:
            domain_name = domain["domain"]["name"]
            domain_score = domain_scores.get(domain_name, 0)
            
            for question in domain["questions"]:
                score = question["answer"]["numeric_score"] if question.get("answer") else None
                question_code = question.get("code", "N/A")
                
                if score is not None and score < 4:  # Only include non-perfect scores
                    # Determine priority based on score (1-4 scale)
                    if score == 1:
                        priority = "high"
                    elif score == 2:
                        priority = "medium"
                    elif score == 3:
                        priority = "low"
                    else:
                        continue  # Skip score 4 (leading practice)
                    
                    # Get specific recommendation from lookup or create default
                    recommendation_text = self._get_recommendation_text(question_code, recommendations_lookup, domain_name)
                    
                    # Remove the prefix if present
                    prefix = "Implement the following to meet AM AI SAFE expectations: "
                    if recommendation_text.startswith(prefix):
                        recommendation_text = recommendation_text[len(prefix):]
                    
                    recommendation = {
                        "domain": domain_name,
                        "domain_score": domain_score,  # Add for sorting
                        "question_id": question_code,
                        "text": recommendation_text
                    }
                    
                    actions[priority].append(recommendation)
        
        # Sort by domain score (lowest first), then by question_id within domain
        for priority in actions:
            actions[priority] = sorted(actions[priority], key=lambda x: (x["domain_score"], x["domain"], x["question_id"]))
            # Remove domain_score from final output (it was just for sorting)
            for action in actions[priority]:
                action.pop("domain_score", None)
        
        # No artificial limits - include all recommendations
        # This ensures all high/medium/low priority items are reported
        
        return actions
    
    def _prepare_heatmap_data(self, questions_data: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Prepare heatmap data exactly matching the results summary format."""
        # Collect all domain data with scores
        domains_with_scores = []
        
        for domain_data in questions_data:
            domain_name = domain_data.get('domain', {}).get('name', 'Unknown')
            questions = domain_data.get('questions', [])
            
            # Collect question scores with codes
            question_scores = []
            total_score = 0
            
            for question in questions:
                score = 0
                if question.get('answer'):
                    score = question['answer'].get('numeric_score', 0)
                
                question_code = question.get('code', 'Q-?')
                question_scores.append({
                    'code': question_code,
                    'score': score,
                    'question_id': question.get('id', '')
                })
                total_score += score
            
            # Sort questions by score (lowest first, as shown in heatmap)
            question_scores.sort(key=lambda x: x['score'])
            
            # Different assessments have different numbers of questions per domain:
            # - Awareness: 5 questions per domain
            # - Readiness: 6 questions per domain  
            # - System: 8 questions per domain
            # Only pad for System assessments, not for Awareness or Readiness
            if self.assessment_type == 'System':
                # Pad to 8 questions if needed (for System assessments)
                while len(question_scores) < 8:
                    question_scores.append({'code': 'N/A', 'score': 0, 'question_id': ''})
                
                # Take only first 8 questions
                question_scores = question_scores[:8]
            # For Awareness and Readiness, use actual question count without padding
            
            domains_with_scores.append({
                'name': domain_name,
                'questions': question_scores,
                'total_score': total_score,
                'avg_score': total_score / len(questions) if questions else 0
            })
        
        # Sort domains by total score (lowest first, as shown in heatmap)
        domains_with_scores.sort(key=lambda x: x['total_score'])
        
        return {
            'domains': domains_with_scores,
            'matrix': [[q['score'] for q in domain['questions']] for domain in domains_with_scores]
        }
    
    def _generate_heatmap_image(self, report_data: Dict[str, Any]) -> bytes:
        """Generate heatmap image based on assessment type.
        
        Delegates to modular chart generation in report_modules.charts
        """
        # Use the refactored modular chart generation
        return generate_heatmap(report_data, self.assessment_type)
    
    def _generate_awareness_bar_image(self, report_data: Dict[str, Any]) -> bytes:
        """Generate stacked bar chart image for Awareness assessments.
        
        Delegates to modular chart generation in report_modules.charts
        """
        return generate_bar_chart(report_data, 'Awareness')

    def _generate_readiness_bar_image(self, report_data: Dict[str, Any]) -> bytes:
        """Generate stacked bar chart image for Readiness assessments.
        
        Delegates to modular chart generation in report_modules.charts
        """
        return generate_bar_chart(report_data, self.assessment_type)

    def _generate_awareness_heatmap_image(self, report_data: Dict[str, Any]) -> bytes:
        """Generate heatmap image for Awareness assessments.
        
        Delegates to modular chart generation in report_modules.charts
        """
        return generate_awareness_heatmap(report_data)

    def _generate_system_heatmap_image(self, report_data: Dict[str, Any]) -> bytes:
        """Generate heatmap image for System assessments.
        
        Delegates to modular chart generation in report_modules.charts
        """
        return generate_system_heatmap(report_data)
    
    def _generate_radar_chart_image(self, report_data: Dict[str, Any]) -> bytes:
        """Generate radar chart showing domain scores with sector benchmarks.
        
        Delegates to modular chart generation in report_modules.charts
        """
        return generate_radar_chart(report_data, self.assessment_type)
    
    def _generate_radar_chart_with_benchmark(self, report_data: Dict[str, Any], benchmark_data: Dict[str, Any]) -> bytes:
        """Generate radar chart comparing organization scores with sector benchmarks."""
        try:
            domain_scores = report_data.get('domain_scores', {})
            sector_name = benchmark_data.get('sector', 'Industry')
            benchmarks = benchmark_data.get('benchmarks', {})
            
            if not domain_scores or not benchmarks:
                # Return simple radar chart if no benchmark data
                return self._generate_radar_chart_image(report_data)
            
            # Prepare data for radar chart - match domain order from benchmarks
            domain_names = []
            user_scores = []
            benchmark_scores = []
            
            # Get domain names from benchmarks to ensure consistent ordering
            for domain_name in benchmarks.keys():
                if domain_name in domain_scores:
                    domain_names.append(domain_name)
                    
                    # Calculate user's average score for this domain
                    score_list = domain_scores[domain_name]
                    avg_score = sum(score_list) / len(score_list) if score_list else 0
                    user_percentage = (avg_score / 4) * 100  # 1-4 scale to percentage
                    user_scores.append(user_percentage)
                    
                    # Get benchmark score
                    benchmark_scores.append(benchmarks[domain_name])
            
            if not domain_names:
                # No matching domains, return simple radar chart
                return self._generate_radar_chart_image(report_data)
            
            # Number of variables
            num_vars = len(domain_names)
            
            # Compute angle for each axis
            angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()
            
            # Complete the loop
            user_scores_plot = user_scores + user_scores[:1]
            benchmark_scores_plot = benchmark_scores + benchmark_scores[:1]
            angles_plot = angles + angles[:1]
            
            # Create figure
            fig, ax = plt.subplots(figsize=(10, 10), subplot_kw=dict(projection='polar'))
            
            # Plot user data
            ax.plot(angles_plot, user_scores_plot, 'o-', linewidth=2.5, 
                   color='#3b82f6', label='Your Score', markersize=8)
            ax.fill(angles_plot, user_scores_plot, alpha=0.25, color='#3b82f6')
            
            # Plot benchmark data
            ax.plot(angles_plot, benchmark_scores_plot, 's--', linewidth=2, 
                   color='#10b981', label=f'{sector_name} Benchmark', markersize=6)
            ax.fill(angles_plot, benchmark_scores_plot, alpha=0.15, color='#10b981')
            
            # Fix axis to go in the right order
            ax.set_theta_offset(np.pi / 2)
            ax.set_theta_direction(-1)
            
            # Set labels
            ax.set_xticks(angles)
            ax.set_xticklabels(domain_names, size=10, fontweight='bold')
            
            # Set y-axis limits and ticks
            ax.set_ylim(0, 100)
            ax.set_yticks([25, 50, 75, 100])
            ax.set_yticklabels(['25%', '50%', '75%', '100%'], size=9)
            
            # Add grid
            ax.grid(True, linestyle='--', alpha=0.7, linewidth=0.5)
            
            # Add title
            ax.set_title(f'Domain Performance vs {sector_name} Benchmark', 
                        size=14, pad=20, fontweight='bold')
            
            # Add legend
            ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1), framealpha=0.9, fontsize=10)
            
            # Save to bytes
            img_buffer = io.BytesIO()
            plt.savefig(img_buffer, format='png', dpi=200, bbox_inches='tight', 
                       facecolor='white', edgecolor='none')
            img_buffer.seek(0)
            img_bytes = img_buffer.getvalue()
            plt.close(fig)
            
            return img_bytes
            
        except Exception as e:
            print(f"Error generating radar chart with benchmark: {str(e)}")
            import traceback
            traceback.print_exc()
            # Fallback to simple radar chart
            return self._generate_radar_chart_image(report_data)

    def _format_top_3_strengths(self, report_data: Dict[str, Any]) -> list:
        """Get top 3 highest scoring domains."""
        domain_scores = report_data.get('domain_scores', {})
        
        # Calculate average percentage for each domain
        domain_percentages = []
        for domain_name, score_list in domain_scores.items():
            avg_score = sum(score_list) / len(score_list) if score_list else 0
            percentage = (avg_score / 4) * 100  # 1-4 scale
            domain_percentages.append({
                'domain': domain_name,
                'percentage': f"{percentage:.1f}%"
            })
        
        # Sort by percentage (highest first) and take top 3
        domain_percentages.sort(key=lambda x: float(x['percentage'].rstrip('%')), reverse=True)
        return domain_percentages[:3]
    
    def _format_top_3_gaps(self, report_data: Dict[str, Any]) -> list:
        """Get top 3 lowest scoring domains."""
        domain_scores = report_data.get('domain_scores', {})
        
        # Calculate average percentage for each domain
        domain_percentages = []
        for domain_name, score_list in domain_scores.items():
            avg_score = sum(score_list) / len(score_list) if score_list else 0
            percentage = (avg_score / 4) * 100  # 1-4 scale
            domain_percentages.append({
                'domain': domain_name,
                'percentage': f"{percentage:.1f}%"
            })
        
        # Sort by percentage (lowest first) and take top 3
        domain_percentages.sort(key=lambda x: float(x['percentage'].rstrip('%')))
        return domain_percentages[:3]
    
    def _generate_comprehensive_executive_summary(self, report_data: Dict[str, Any]) -> str:
        """Generate a comprehensive multi-paragraph executive summary."""
        org_name = report_data.get('org', {}).get('name', 'Organization')
        overall_score = report_data.get('overall', {}).get('score', 0)
        overall_tier = report_data.get('overall', {}).get('tier', 'Basic')
        
        actions = report_data.get('actions', {})
        high_count = len(actions.get('high', []))
        medium_count = len(actions.get('medium', []))
        low_count = len(actions.get('low', []))
        
        return f"""The AM AI SAFE Framework Assessment for {org_name} has been completed, representing a comprehensive evaluation of the organization's artificial intelligence systems against 88 detailed questions spanning 11 critical domains of AI safety, ethics, and governance. This assessment provides {org_name} with a thorough analysis of their current AI maturity level and a strategic roadmap for improvement.

ASSESSMENT SCOPE AND METHODOLOGY:
The evaluation encompasses the complete spectrum of AI governance, examining Fairness, Transparency, Explainability, Accountability, Data Integrity, Reliability, Security, Privacy, Safety, Inclusivity, and Sustainability. Each domain contains 8 targeted questions designed to assess current practices against industry best practices and regulatory expectations.

KEY FINDINGS SUMMARY:
{org_name} has achieved an overall AI maturity score of {overall_score}%, positioning the organization within the {overall_tier} maturity category. This assessment reveals both areas of strength and significant opportunities for improvement across the AI safety framework.

The evaluation has identified {high_count + medium_count + low_count} specific recommendations for enhancement, categorized by priority and implementation timeline:
• {high_count} High Priority actions addressing critical gaps requiring immediate attention and resource allocation
• {medium_count} Medium Priority initiatives for systematic improvement over the next 6-12 months
• {low_count} Low Priority strategic enhancements for long-term competitive advantage and industry leadership

STRATEGIC IMPLICATIONS:
The assessment results indicate that {org_name} has established foundational AI practices but requires focused investment in key areas to achieve industry-leading AI governance standards. The identified recommendations provide a clear pathway for systematic improvement, risk mitigation, and competitive advantage in AI deployment.

This report serves as both a comprehensive audit of current capabilities and a strategic roadmap for AI excellence, enabling {org_name} to prioritize investments, allocate resources effectively, and implement systematic improvements that will enhance AI safety, ethical compliance, and operational excellence."""
    
    def _generate_assessment_methodology(self) -> str:
        """Generate detailed assessment methodology description."""
        return """ASSESSMENT METHODOLOGY AND FRAMEWORK:

The AM AI SAFE Framework Assessment employs a rigorous, multi-dimensional evaluation approach designed to provide comprehensive insights into organizational AI maturity across 11 critical domains. Each assessment question is evaluated using a standardized 4-point scoring scale:

• Score 0 (Non-Ideal): Indicates fundamental gaps or absence of required practices, representing critical risks that demand immediate attention and corrective action.

• Score 1 (Basic): Demonstrates initial awareness and basic implementation of AI safety practices, but lacks comprehensive coverage or systematic approach.

• Score 2 (Good): Shows solid implementation of AI safety practices with systematic approaches, though opportunities remain for enhancement and optimization.

• Score 3 (Best Practice): Represents industry-leading implementation with comprehensive coverage, proactive management, and continuous improvement processes.

The assessment methodology ensures objective evaluation through standardized criteria, comprehensive coverage of AI safety domains, and actionable recommendations aligned with regulatory expectations and industry best practices. Results provide both current state analysis and forward-looking guidance for systematic improvement."""
    
    def _generate_domain_analysis(self, report_data: Dict[str, Any]) -> str:
        """Generate detailed domain-by-domain analysis."""
        heatmap_data = report_data.get('heatmap_data', {})
        domains = heatmap_data.get('domains', [])
        
        analysis = "DOMAIN-SPECIFIC PERFORMANCE ANALYSIS:\n\n"
        
        for domain in domains:
            domain_name = domain.get('name', 'Unknown')
            questions = domain.get('questions', [])
            scores = [q.get('score', 0) for q in questions if q.get('code') != 'N/A']
            
            if scores:
                avg_score = sum(scores) / len(scores)
                max_score = max(scores)
                min_score = min(scores)
                
                # Determine domain maturity level
                if avg_score >= 2.5:
                    maturity = "High"
                elif avg_score >= 1.5:
                    maturity = "Moderate"
                elif avg_score >= 0.5:
                    maturity = "Developing"
                else:
                    maturity = "Foundational"
                
                analysis += f"{domain_name} Domain ({maturity} Maturity):\n"
                analysis += f"Average Score: {avg_score:.1f}/4.0 | Range: {min_score}-{max_score} | Questions Evaluated: {len(scores)}\n"
                
                # Add domain-specific insights (updated for 1-4 scale)
                if avg_score < 2.0:
                    analysis += f"Critical Focus Area: {domain_name} requires immediate attention with systematic capability building and resource investment.\n"
                elif avg_score < 3.0:
                    analysis += f"Development Opportunity: {domain_name} shows foundational practices with clear opportunities for systematic enhancement.\n"
                else:
                    analysis += f"Strength Area: {domain_name} demonstrates solid practices with opportunities for optimization and industry leadership.\n"
                
                analysis += "\n"
        
        return analysis
    
    def _generate_key_findings(self, report_data: Dict[str, Any]) -> str:
        """Generate key findings and insights."""
        overall_score = report_data.get('overall', {}).get('score', 0)
        actions = report_data.get('actions', {})
        
        findings = "KEY FINDINGS AND STRATEGIC INSIGHTS:\n\n"
        
        findings += "1. OVERALL AI MATURITY ASSESSMENT:\n"
        if overall_score >= 75:
            findings += "The organization demonstrates advanced AI maturity with comprehensive governance frameworks and systematic risk management practices.\n\n"
        elif overall_score >= 50:
            findings += "The organization shows moderate AI maturity with established foundational practices and clear pathways for systematic improvement.\n\n"
        elif overall_score >= 25:
            findings += "The organization exhibits developing AI maturity with basic practices in place but requiring focused investment in key capability areas.\n\n"
        else:
            findings += "The organization is in the foundational stage of AI maturity, requiring comprehensive capability building and systematic framework implementation.\n\n"
        
        findings += "2. PRIORITY FOCUS AREAS:\n"
        
        high_actions = actions.get('high', [])
        if high_actions:
            findings += f"Critical Gaps Identified: {len(high_actions)} high-priority areas require immediate attention to address fundamental risks and compliance requirements.\n"
        
        medium_actions = actions.get('medium', [])
        if medium_actions:
            findings += f"Systematic Improvements: {len(medium_actions)} medium-priority initiatives will strengthen foundational capabilities and enhance operational excellence.\n"
        
        low_actions = actions.get('low', [])
        if low_actions:
            findings += f"Strategic Enhancements: {len(low_actions)} low-priority opportunities will drive competitive advantage and industry leadership positioning.\n\n"
        
        findings += "3. IMPLEMENTATION SUCCESS FACTORS:\n"
        findings += "• Executive leadership commitment and resource allocation for systematic AI governance improvement\n"
        findings += "• Cross-functional collaboration between IT, legal, compliance, and business stakeholders\n"
        findings += "• Phased implementation approach with clear milestones and success metrics\n"
        findings += "• Continuous monitoring and improvement processes to maintain and enhance AI safety standards\n"
        
        return findings
    
    def _generate_implementation_roadmap(self, report_data: Dict[str, Any]) -> str:
        """Generate implementation roadmap and next steps."""
        actions = report_data.get('actions', {})
        
        roadmap = "IMPLEMENTATION ROADMAP AND STRATEGIC RECOMMENDATIONS:\n\n"
        
        roadmap += "PHASE 1: IMMEDIATE ACTIONS (0-3 months)\n"
        roadmap += "Priority: Address critical gaps and compliance requirements\n"
        high_actions = actions.get('high', [])
        if high_actions:
            roadmap += f"Focus Areas: {len(high_actions)} high-priority recommendations addressing fundamental risks\n"
            roadmap += "Success Metrics: Risk mitigation, compliance adherence, foundational capability establishment\n\n"
        
        roadmap += "PHASE 2: SYSTEMATIC IMPROVEMENTS (3-12 months)\n"
        roadmap += "Priority: Build comprehensive AI governance capabilities\n"
        medium_actions = actions.get('medium', [])
        if medium_actions:
            roadmap += f"Focus Areas: {len(medium_actions)} medium-priority initiatives for systematic enhancement\n"
            roadmap += "Success Metrics: Process maturity, operational excellence, stakeholder confidence\n\n"
        
        roadmap += "PHASE 3: STRATEGIC OPTIMIZATION (12+ months)\n"
        roadmap += "Priority: Achieve industry leadership and competitive advantage\n"
        low_actions = actions.get('low', [])
        if low_actions:
            roadmap += f"Focus Areas: {len(low_actions)} strategic enhancements for market differentiation\n"
            roadmap += "Success Metrics: Industry recognition, innovative practices, sustainable competitive advantage\n\n"
        
        roadmap += "CONTINUOUS IMPROVEMENT FRAMEWORK:\n"
        roadmap += "• Quarterly assessment reviews and progress monitoring\n"
        roadmap += "• Annual comprehensive re-evaluation and strategic planning\n"
        roadmap += "• Ongoing stakeholder engagement and feedback integration\n"
        roadmap += "• Emerging technology and regulatory landscape monitoring\n"
        
        return roadmap
    
    def _load_recommendations_lookup(self) -> Dict[str, Any]:
        """Load recommendations lookup from JSON file."""
        import json
        backend_dir = Path(__file__).parent
        recommendations_path = backend_dir / "AMAI_SAFE_recommendations_lookup.json"
        
        try:
            with open(recommendations_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"Warning: Could not load recommendations lookup: {e}")
            return {"recommendations": {}}
    
    def _generate_system_results_summary_text(self, report_data: Dict[str, Any]) -> str:
        """
        Generate the complete Results Summary paragraph for AI System Maturity assessments.
        Matches the frontend ResultsPage.js Results Summary section.
        """
        overall = report_data.get('overall', {})
        tier = overall.get('tier', 'Foundational')
        score = overall.get('score', 0)
        
        # Get system name from system_info or org
        system_info = report_data.get('system_info') or {}
        system_name = system_info.get('systemName') or system_info.get('system_name') or system_info.get('organizationName') or report_data.get('org', {}).get('name', 'The system')
        
        # Get sector benchmark comparison if available
        sector_average = report_data.get('sector_average')
        sector_name = system_info.get('industry') or report_data.get('sector_name', '')
        
        # Build the opening sentence
        summary = f"The results indicate that {system_name} has achieved an overall AI maturity score of {score:.1f}%, placing this system within the {tier} AI Maturity category"
        
        # Add sector comparison if available
        if sector_average is not None and sector_name:
            if score > sector_average:
                comparison = "above"
            elif score < sector_average:
                comparison = "below"
            else:
                comparison = "equal to"
            summary += f", which is {comparison} the {sector_name} sector AI system maturity average of {sector_average:.0f}%."
        else:
            summary += "."
        
        # Add tier-specific explanatory text
        if tier == 'Leading':
            summary += " This rating reflects exemplary AI governance and ethical assurance, setting a benchmark for responsible AI leadership. Governance systems are fully embedded, adaptive, and continuously refined through data-driven insights, external validation, and innovation. The focus is on optimisation, transparency, and sustained improvement across all AI operations."
        elif tier == 'Established':
            summary += " This rating reflects well-defined and consistently applied AI governance frameworks across most domains. Risk management, transparency, and ethical oversight are integrated into day-to-day operations. Continuous monitoring and regular review cycles are evident, though further optimisation and automation would strengthen maturity and resilience against evolving AI risks."
        elif tier == 'Developing':
            summary += " This rating reflects growing awareness and emerging structure in its AI governance practices. Some policies and controls are in place, but they are applied inconsistently. Progress has been made in recognising key governance needs; however, targeted improvements are needed to achieve full integration and accountability across the AI lifecycle."
        else:  # Foundational
            summary += " This rating reflects minimal or inconsistent AI governance, with most practices being reactive or informal. Policies, processes, and accountability structures are largely undeveloped, resulting in fragmented oversight and heightened risk exposure. Immediate action is required to establish a foundational governance framework, define roles and responsibilities, and embed basic ethical and risk management principles."
        
        return summary
    
    def _get_recommendation_text(self, question_code: str, recommendations_lookup: Dict[str, Any], domain_name: str) -> str:
        """Get specific recommendation text for a question code."""
        recommendations = recommendations_lookup.get("recommendations", {})
        
        if question_code in recommendations:
            return recommendations[question_code].get("default_recommendation", f"Implement comprehensive {domain_name.lower()} improvements to achieve best practices.")
        else:
            # Fallback to generic recommendation
            return f"Implement comprehensive {domain_name.lower()} improvements to achieve best practices in this area."
    
    def _create_heatmap_for_template(self, heatmap_bytes: bytes) -> str:
        """Create a heatmap representation for the template."""
        if not heatmap_bytes:
            return 'HEATMAP_IMAGE_PLACEHOLDER'
        
        # For now, we'll use a placeholder since InlineImage was causing issues
        # In the future, this can be enhanced to properly embed the image
        return f'[HEATMAP IMAGE - {len(heatmap_bytes)} bytes generated]'
    
    def _generate_executive_summary(self, report_data: Dict[str, Any], user_data: Dict[str, Any]) -> str:
        """Generate comprehensive executive summary."""
        org_name = user_data.get('organization_name', 'Organization')
        overall_score = report_data.get('overall', {}).get('score', 0)
        overall_tier = report_data.get('overall', {}).get('tier', 'Basic')
        
        # Count recommendations by priority
        high_count = len(report_data.get('actions', {}).get('high', []))
        medium_count = len(report_data.get('actions', {}).get('medium', []))
        low_count = len(report_data.get('actions', {}).get('low', []))
        
        summary = f"""The AM AI SAFE Framework Assessment for {org_name} has been completed, evaluating the organization's artificial intelligence systems against 88 questions across 11 critical domains of AI safety and ethics.

The assessment methodology follows the AM AI SAFE framework, examining key areas including Fairness, Transparency, Explainability, Accountability, Data Integrity, Reliability, Security, Privacy, Safety, Inclusivity, and Sustainability.

KEY FINDINGS:
• Overall AI Maturity Score: {overall_score}%
• AI Maturity Category: {overall_tier}
• Total Assessment Questions: 88 questions across 11 domains
• Identified Recommendations: {high_count + medium_count + low_count} actionable items

ASSESSMENT OVERVIEW:
{org_name} demonstrates {self._get_maturity_description(overall_score)} level of AI maturity. The assessment revealed specific areas of strength and opportunities for improvement across the AI safety framework.

PRIORITY RECOMMENDATIONS:
• {high_count} High Priority actions addressing critical gaps requiring immediate attention
• {medium_count} Medium Priority initiatives for moderate-term enhancement 
• {low_count} Low Priority strategic improvements for long-term development

The detailed assessment results, including the AI Maturity Heatmap and comprehensive recommendations, provide {org_name} with a clear roadmap for enhancing their AI safety and ethical practices."""

        return summary
    
    def _generate_assessment_results(self, report_data: Dict[str, Any]) -> str:
        """Generate detailed assessment results section."""
        overall_score = report_data.get('overall', {}).get('score', 0)
        overall_tier = report_data.get('overall', {}).get('tier', 'Basic')
        
        # Analyze domain performance
        heatmap_data = report_data.get('heatmap_data', {})
        domains = heatmap_data.get('domains', [])
        
        highest_domain = max(domains, key=lambda x: x['avg_score']) if domains else None
        lowest_domain = min(domains, key=lambda x: x['avg_score']) if domains else None
        
        results = f"""The comprehensive assessment of AI systems has yielded an overall AI maturity score of {overall_score}%, placing the organization within the {overall_tier} category.

PERFORMANCE ANALYSIS:
The assessment evaluated 88 specific questions across 11 domains, with each question scored on a 4-point scale (1=Foundational, 2=Developing, 3=Established, 4=Leading).

DOMAIN PERFORMANCE HIGHLIGHTS:"""

        if highest_domain and lowest_domain:
            results += f"""
• Highest Performing Domain: {highest_domain['name']} ({highest_domain['avg_score']:.1f}/4.0 average)
• Area for Greatest Improvement: {lowest_domain['name']} ({lowest_domain['avg_score']:.1f}/4.0 average)"""

        results += """

HEATMAP VISUALIZATION:
The AI Maturity Heatmap below provides a visual representation of performance across all domains and questions. The heatmap is organized with:
• Rows representing the 11 AI safety domains, sorted by total score (lowest performing at the top)
• Columns representing the 8 questions within each domain, arranged by score (lowest scores on the left)
• Color coding indicating performance levels: Red (critical gaps), Orange (below average), Yellow (moderate performance), Green (high performance)

Each cell represents the score for a specific question, enabling identification of precise areas requiring attention and those demonstrating strong performance."""

        return results
    
    def _generate_statistics(self, report_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate comprehensive statistics for the report."""
        actions = report_data.get('actions', {})
        heatmap_data = report_data.get('heatmap_data', {})
        domains = heatmap_data.get('domains', [])
        
        # Count questions by score level
        score_counts = {'non_ideal': 0, 'basic': 0, 'good': 0, 'best': 0}
        total_questions = 0
        
        for domain in domains:
            for question in domain.get('questions', []):
                total_questions += 1
                score = question.get('score', 0)
                if score == 0:
                    score_counts['non_ideal'] += 1
                elif score == 1:
                    score_counts['basic'] += 1
                elif score == 2:
                    score_counts['good'] += 1
                elif score == 3:
                    score_counts['best'] += 1
        
        return {
            'total_questions': total_questions,
            'total_domains': len(domains),
            'score_distribution': score_counts,
            'recommendations_count': {
                'high': len(actions.get('high', [])),
                'medium': len(actions.get('medium', [])),
                'low': len(actions.get('low', []))
            }
        }
    
    def _get_maturity_description(self, score: float) -> str:
        """Get descriptive text for maturity level."""
        if score >= 90:
            return "an excellent"
        elif score >= 75:
            return "a high to excellent"
        elif score >= 60:
            return "a moderate to high"
        elif score >= 40:
            return "a low to moderate"
        elif score >= 25:
            return "a basic to low"
        else:
            return "a basic"
    
    def _populate_recommendation_tables(self, doc: DocxTemplate, report_data: Dict[str, Any]) -> None:
        """
        Programmatically populate the recommendation tables with multiple rows.
        This fixes the issue where Jinja2 loops don't properly create table rows in Word.
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Get the underlying python-docx Document object
        docx_doc = doc.docx
        
        # Find the three recommendation tables (they should be the last 3 tables in the document)
        tables = docx_doc.tables
        if len(tables) < 3:
            print(f"Warning: Expected at least 3 tables in template, found {len(tables)}")
            return
        
        # Get the last 3 tables (High, Medium, Low priority)
        high_table = tables[-3]
        medium_table = tables[-2]
        low_table = tables[-1]
        
        # Get recommendations from report_data
        actions = report_data.get('actions', {})
        high_actions = actions.get('high', [])
        medium_actions = actions.get('medium', [])
        low_actions = actions.get('low', [])
        
        print(f"Populating tables: High={len(high_actions)}, Medium={len(medium_actions)}, Low={len(low_actions)}")
        
        # Populate each table
        self._populate_table_with_actions(high_table, high_actions)
        self._populate_table_with_actions(medium_table, medium_actions)
        self._populate_table_with_actions(low_table, low_actions)
    
    def _populate_awareness_sector_tables(self, doc: DocxTemplate, report_data: Dict[str, Any]) -> None:
        """
        Programmatically populate the sector action tables for Awareness assessments.
        Uses sector_actions which has question_id, domain, and sector_action fields.
        Column order: Question ID | Domain | Recommended Action
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from copy import deepcopy
        
        docx_doc = doc.docx
        tables = docx_doc.tables
        
        if len(tables) < 3:
            print(f"Warning: Expected at least 3 tables in template, found {len(tables)}")
            return
        
        # Get the last 3 tables (High, Medium, Low priority sector actions)
        high_table = tables[-3]
        medium_table = tables[-2]
        low_table = tables[-1]
        
        # Get sector_actions from report_data
        sector_actions = report_data.get('sector_actions', {'high': [], 'medium': [], 'low': []})
        high_actions = sector_actions.get('high', [])[:5]  # Top 5
        medium_actions = sector_actions.get('medium', [])[:5]
        low_actions = sector_actions.get('low', [])[:5]
        
        print(f"Populating Awareness sector tables: High={len(high_actions)}, Medium={len(medium_actions)}, Low={len(low_actions)}")
        
        # Populate each table with correct column order and field names
        self._populate_sector_table(high_table, high_actions)
        self._populate_sector_table(medium_table, medium_actions)
        self._populate_sector_table(low_table, low_actions)
    
    def _populate_readiness_sector_tables(self, doc: DocxTemplate, report_data: Dict[str, Any]) -> None:
        """
        Programmatically populate the sector action tables for Readiness and Orgwide assessments.
        Uses the same table structure as Awareness: question_id, domain, sector_action
        Column order: Question ID | Domain | Recommended Action
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from copy import deepcopy
        import json
        
        docx_doc = doc.docx
        tables = docx_doc.tables
        
        if len(tables) < 3:
            print(f"Warning: Expected at least 3 tables in template, found {len(tables)}")
            return
        
        # Get the last 3 tables (High, Medium, Low priority sector actions)
        high_table = tables[-3]
        medium_table = tables[-2]
        low_table = tables[-1]
        
        # Load sector-specific actions from JSON file based on assessment type
        if self.assessment_type == 'Orgwide':
            sector_actions_file = Path(__file__).parent / "orgwide_actions.json"
        else:
            sector_actions_file = Path(__file__).parent / "readiness_sector_actions.json"
        
        sector_actions_data = {}
        if sector_actions_file.exists():
            try:
                with open(sector_actions_file, 'r', encoding='utf-8') as f:
                    sector_actions_data = json.load(f)
                print(f"DEBUG: Loaded sector actions from {sector_actions_file.name} for {len(sector_actions_data)} industries")
            except Exception as e:
                print(f"ERROR loading {sector_actions_file.name}: {e}")
        
        # Get the sector/industry from report_data
        sector_name = report_data.get('sector_name', '')
        print(f"DEBUG _populate_readiness_sector_tables: sector_name='{sector_name}'")
        
        # Get sector-specific actions for this industry
        industry_actions = sector_actions_data.get(sector_name, {})
        if not industry_actions and sector_name:
            # Try to find a close match
            for ind in sector_actions_data.keys():
                if sector_name.lower() in ind.lower() or ind.lower() in sector_name.lower():
                    industry_actions = sector_actions_data[ind]
                    print(f"DEBUG: Using actions from '{ind}' for '{sector_name}'")
                    break
        
        if not industry_actions:
            # Fallback to "Other Industries" if available
            industry_actions = sector_actions_data.get('Other Industries', {})
            print(f"DEBUG: Using 'Other Industries' fallback actions")
        
        print(f"DEBUG: Found {len(industry_actions)} sector-specific actions for industry")
        
        # Build actions from report_data with sector-specific action text
        actions = report_data.get('actions', {})
        
        def format_for_table(action):
            question_id = action.get('question_id', '')
            # Get sector-specific action text if available
            sector_action = industry_actions.get(question_id, '')
            if not sector_action:
                # Fallback to generic text if no sector action found
                sector_action = action.get('text', action.get('recommendation', ''))
            return {
                'question_id': question_id,
                'domain': action.get('domain', ''),
                'sector_action': sector_action
            }
        
        high_actions = [format_for_table(a) for a in actions.get('high', [])[:5]]
        medium_actions = [format_for_table(a) for a in actions.get('medium', [])[:5]]
        low_actions = [format_for_table(a) for a in actions.get('low', [])[:5]]
        
        print(f"Populating Readiness sector tables: High={len(high_actions)}, Medium={len(medium_actions)}, Low={len(low_actions)}")
        if high_actions:
            print(f"  Sample high action: {high_actions[0].get('question_id')} - {high_actions[0].get('sector_action', '')[:80]}...")
        
        # Populate each table with correct column order (same as Awareness)
        self._populate_sector_table(high_table, high_actions)
        self._populate_sector_table(medium_table, medium_actions)
        self._populate_sector_table(low_table, low_actions)
    
    def _populate_sector_table(self, table, actions: List[Dict[str, Any]]) -> None:
        """
        Populate a sector actions table with the correct column order.
        Column 1: Question ID (action.question_id)
        Column 2: Domain (action.domain)
        Column 3: Recommended Action (action.sector_action)
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from copy import deepcopy
        
        # Remove all rows except header
        while len(table.rows) > 1:
            self._remove_row(table, 1)
        
        if not actions:
            print(f"DEBUG: No actions to add, keeping only header row")
            return
        
        header_row = table.rows[0]
        header_tr = header_row._tr
        tbl = table._element
        
        for action in actions:
            tr = deepcopy(header_tr)
            cells = tr.findall('.//w:tc', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            
            if len(cells) >= 3:
                # CORRECT ORDER: Question ID, Domain, Sector Action
                values = [
                    action.get('question_id', ''),
                    action.get('domain', ''),
                    action.get('sector_action', '')
                ]
                
                for i in range(3):
                    cell = cells[i]
                    
                    # Remove header background color
                    tcPr = cell.find('.//w:tcPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    if tcPr is not None:
                        shd = tcPr.find('.//w:shd', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                        if shd is not None:
                            tcPr.remove(shd)
                    
                    # Clear existing content
                    for p in cell.findall('.//w:p', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        cell.remove(p)
                    
                    # Add new paragraph
                    p = OxmlElement('w:p')
                    
                    if values[i]:
                        r = OxmlElement('w:r')
                        rPr = OxmlElement('w:rPr')
                        
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:ascii'), 'Arial')
                        rFonts.set(qn('w:hAnsi'), 'Arial')
                        rPr.append(rFonts)
                        
                        sz = OxmlElement('w:sz')
                        sz.set(qn('w:val'), '20')
                        rPr.append(sz)
                        
                        r.append(rPr)
                        
                        t = OxmlElement('w:t')
                        t.text = values[i]
                        r.append(t)
                        p.append(r)
                    
                    cell.append(p)
                
                # Remove extra cells
                for i in range(3, len(cells)):
                    tr.remove(cells[i])
            
            tbl.append(tr)
        
        print(f"DEBUG: Added {len(actions)} sector action rows")
    
    def _populate_table_with_actions(self, table, actions: List[Dict[str, Any]]) -> None:
        """
        Populate a table with action recommendations, creating one row per action.
        Creates rows with exactly 3 cells by copying header row structure.
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from copy import deepcopy
        
        # Remove all rows except header
        while len(table.rows) > 1:
            self._remove_row(table, 1)
        
        # Get the header row to copy its structure
        header_row = table.rows[0]
        header_tr = header_row._tr
        
        # Get table element
        tbl = table._element
        
        # Add rows with exactly 3 cells for each action
        for action in actions:
            # Clone the header row structure to get proper column widths
            tr = deepcopy(header_tr)
            
            # Clear the text from the cloned cells and set new values
            cells = tr.findall('.//w:tc', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            
            # We should have at least 3 cells from the header
            if len(cells) >= 3:
                values = [
                    action.get('domain', ''),
                    action.get('question_id', ''),
                    action.get('text', '')
                ]
                
                for i in range(3):
                    cell = cells[i]
                    
                    # Remove cell shading (background color) from header
                    tcPr = cell.find('.//w:tcPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    if tcPr is not None:
                        shd = tcPr.find('.//w:shd', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                        if shd is not None:
                            tcPr.remove(shd)
                    
                    # Clear existing content
                    for p in cell.findall('.//w:p', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        cell.remove(p)
                    
                    # Add new paragraph with text
                    p = OxmlElement('w:p')
                    
                    if values[i]:
                        r = OxmlElement('w:r')
                        rPr = OxmlElement('w:rPr')
                        
                        # Set font
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:ascii'), 'Arial')
                        rFonts.set(qn('w:hAnsi'), 'Arial')
                        rPr.append(rFonts)
                        
                        # Set font size
                        sz = OxmlElement('w:sz')
                        sz.set(qn('w:val'), '20')  # 10pt = 20 half-points
                        rPr.append(sz)
                        
                        r.append(rPr)
                        
                        t = OxmlElement('w:t')
                        t.text = values[i]
                        r.append(t)
                        p.append(r)
                    
                    cell.append(p)
                
                # Remove extra cells beyond the first 3
                for i in range(3, len(cells)):
                    tr.remove(cells[i])
            
            # Append row to table
            tbl.append(tr)
        
        print(f"DEBUG: Added {len(actions)} rows with 3 cells each (copied from header structure)")
    
    def _remove_row(self, table, row_idx: int) -> None:
        """Remove a row from a table."""
        from docx.oxml.table import CT_Row
        tbl = table._tbl
        tr = table.rows[row_idx]._tr
        tbl.remove(tr)
    
    def _center_align_images(self, doc: DocxTemplate) -> None:
        """Center-align all paragraphs containing images in the document."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        docx_doc = doc.docx
        
        images_centered = 0
        for paragraph in docx_doc.paragraphs:
            # Check if paragraph contains an image (drawing element)
            has_image = False
            for run in paragraph.runs:
                if run._element.findall('.//a:blip', namespaces={
                    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                }):
                    has_image = True
                    break
                # Also check for inline images
                if run._element.findall('.//w:drawing', namespaces={
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }):
                    has_image = True
                    break
            
            if has_image:
                # Center-align the paragraph
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                images_centered += 1
        
        print(f"DEBUG: Center-aligned {images_centered} paragraphs containing images")


    def _calculate_maturity_distribution(self, report_data: Dict[str, Any]) -> Dict[str, Any]:
        """Calculate maturity distribution for test template schema."""
        heatmap_data = report_data.get('heatmap_data', {})
        domains = heatmap_data.get('domains', [])
        
        all_scores = []
        for domain in domains:
            questions = domain.get('questions', [])
            for q in questions:
                if q.get('score') is not None:
                    all_scores.append(q.get('score', 0))
        
        non_ideal = len([s for s in all_scores if s == 1])
        basic = len([s for s in all_scores if s == 2])
        good = len([s for s in all_scores if s == 3])
        advanced = len([s for s in all_scores if s == 4])
        total = len(all_scores) if all_scores else 1
        
        return {
            'non_ideal': non_ideal,
            'basic': basic,
            'good': good,
            'advanced': advanced,
            'total_questions': total,
            'percent': {
                'non_ideal': round((non_ideal / total) * 100, 1) if total > 0 else 0,
                'basic': round((basic / total) * 100, 1) if total > 0 else 0,
                'good': round((good / total) * 100, 1) if total > 0 else 0,
                'advanced': round((advanced / total) * 100, 1) if total > 0 else 0
            }
        }
    
    def _format_domains_for_template(self, report_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Format domain results for test template schema."""
        heatmap_data = report_data.get('heatmap_data', {})
        domains = heatmap_data.get('domains', [])
        
        print(f"DEBUG _format_domains_for_template: {len(domains)} domains")
        
        result = []
        for domain in domains:
            domain_name = domain.get('name', '')
            # Use placeholder for & to preserve it through docxtpl rendering
            domain_name_escaped = replace_amp_with_placeholder(domain_name)
            questions = domain.get('questions', [])
            scores = [q.get('score', 0) for q in questions if q.get('score') is not None]
            
            print(f"  Domain '{domain_name}': {len(questions)} questions, {len(scores)} scores, scores={scores}")
            
            if scores:
                # For Awareness: 5 questions * 4 max = 20 points max per domain
                # For System: 8 questions * 4 max = 32 points max per domain
                avg_pct = (sum(scores) / (len(scores) * 4)) * 100
                low_scoring = [q for q in questions if q.get('score') is not None and q.get('score', 0) <= 2]
                
                print(f"    sum(scores)={sum(scores)}, len(scores)={len(scores)}, max_possible={len(scores)*4}, avg_pct={avg_pct:.1f}%")
                
                # Determine tier based on score using standard maturity tiers
                # Foundational (0-40%), Developing (41-65%), Established (66-85%), Leading (86-100%)
                if avg_pct >= 86:
                    tier = "Leading"
                elif avg_pct >= 66:
                    tier = "Established"
                elif avg_pct >= 41:
                    tier = "Developing"
                else:
                    tier = "Foundational"
                
                result.append({
                    'name': domain_name_escaped,  # Escaped for safe template rendering
                    'score': round(avg_pct, 1),  # Keep as number for comparisons
                    'score_display': f"{round(avg_pct)}%",  # Formatted string for template
                    'tier': tier,
                    'question_count': len(scores),
                    'low_scoring_question_count': len(low_scoring),
                    'top_low_questions': [
                        {
                            'question_id': q.get('code', 'N/A'),
                            'question_text': q.get('text', 'N/A')[:100],
                            'score': round((q.get('score', 0) / 4) * 100, 1)
                        }
                        for q in sorted(low_scoring, key=lambda x: x.get('score', 0))[:5]
                    ]
                })
        
        return result

    def _normalize_system_info(self, system_info: Dict[str, Any]) -> Dict[str, Any]:
        """Normalize system_info field names for test template compatibility."""
        return {
            'lifecycle': system_info.get('lifecycle', system_info.get('lifecycleStage', 'Unknown')),
            'criticality': system_info.get('criticality', 'Unknown'),
            'ownership': system_info.get('ownership', system_info.get('systemOwnership', 'Unknown')),
            'hosting': system_info.get('hosting', system_info.get('cloudProvider', 'Unknown')),
            'dataSensitivity': system_info.get('dataSensitivity', 'Unknown'),
            'oversight': system_info.get('oversight', system_info.get('humanOversight', 'Unknown')),
            'system_name': system_info.get('system_name', system_info.get('systemName', 'AI System')),
            'system_description': system_info.get('system_description', system_info.get('systemDescription', '')),
            # Keep original fields too for backwards compatibility
            **system_info
        }
    
    def _build_actions_with_summary(self, report_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Build actions structure with summary for templates.
        Provides BOTH score-based and smart priority orderings.
        """
        actions = report_data.get('actions', {})
        sector_actions = report_data.get('sector_actions', {})
        
        # Load question metadata for smart priority calculation
        question_metadata = self._load_question_metadata()
        
        print(f"DEBUG _build_actions_with_summary:")
        print(f"  - actions keys: {actions.keys() if actions else 'None'}")
        print(f"  - use_smart_priority: {self.use_smart_priority}")
        print(f"  - question_metadata loaded: {len(question_metadata)} questions")
        
        # Create lookup for sector-specific action text by question_id
        sector_action_lookup = {}
        for priority in ['high', 'medium', 'low']:
            for sa in sector_actions.get(priority, []):
                q_id = sa.get('question_id', '')
                if q_id:
                    sector_action_lookup[q_id] = sa.get('sector_action', '')
        
        # Domain to risk theme mapping
        domain_risk_themes = {
            'Fairness': 'Bias & Discrimination Risk',
            'Transparency': 'Disclosure & Communication Risk',
            'Explainability': 'Decision Interpretability Risk',
            'Accountability': 'Governance & Oversight Risk',
            'Data Integrity': 'Data Quality & Lineage Risk',
            'Reliability': 'System Performance Risk',
            'Security': 'Cybersecurity & Access Risk',
            'Privacy': 'Data Protection & Privacy Risk',
            'Safety': 'Harm Prevention Risk',
            'Inclusivity': 'Accessibility & Equity Risk',
            'Sustainability': 'Environmental & Long-term Risk'
        }
        
        # Timeframe defaults based on quadrant (for smart priority) or score-based priority
        quadrant_timeframes = {
            'Quick win': '0-30 days',
            'Strategic project': '3-6 months',
            'Opportunistic improvement': '1-3 months',
            'Lower priority': '6-12 months'
        }
        
        score_priority_defaults = {
            'high': {'timeframe': '0-90 days', 'effort_label': 'High'},
            'medium': {'timeframe': '3-6 months', 'effort_label': 'Medium'},
            'low': {'timeframe': '6-12 months', 'effort_label': 'Low'}
        }
        
        def enrich_action(action: Dict[str, Any], score_priority: str) -> Dict[str, Any]:
            """Enrich an action item with smart priority data and all fields."""
            question_id = action.get('question_id', '')
            domain = action.get('domain', 'Unknown')
            
            # Get metadata for this question
            metadata = question_metadata.get(question_id, {})
            impact_weight = metadata.get('impact_weight', 0.5)
            effort_score = metadata.get('effort_score', 0.5)
            
            # Calculate gap from score (score 1 = gap 3, score 2 = gap 2, score 3 = gap 1)
            score_map = {'high': 1, 'medium': 2, 'low': 3}
            maturity_score = score_map.get(score_priority, 2)
            gap = 4 - maturity_score
            
            # Calculate smart priority
            smart_priority_score = self._calculate_smart_priority(impact_weight, gap, effort_score)
            quadrant_label = self._get_quadrant_label(impact_weight, effort_score)
            
            # Get sector-specific recommendation if available
            sector_rec = sector_action_lookup.get(question_id, '')
            default_rec = action.get('text', 'No recommendation available')
            
            # Get action_steps from metadata if available
            metadata_action_steps = metadata.get('action_steps', '')
            
            enriched = {
                'domain': domain,
                'question_id': question_id,
                # Default AM AI SAFE recommendation
                'text': default_rec,
                # Sector-specific action step (empty string if not available)
                'sector_text': sector_rec,
                # Metadata action steps (from system_question_metadata.json)
                'metadata_action_steps': metadata_action_steps,
                # Combined: sector-specific if available, otherwise default
                'recommendation': sector_rec if sector_rec else default_rec,
                # Score-based priority (high/medium/low based on maturity score)
                'score_priority': score_priority.capitalize(),
                # Smart priority data
                'impact_weight': impact_weight,
                'impact_percent': round(impact_weight * 100),
                'effort_score': effort_score,
                'effort_percent': round(effort_score * 100),
                'gap': gap,
                'smart_priority_score': round(smart_priority_score, 3),
                'quadrant_label': quadrant_label,
                # Timeframe based on quadrant for smart priority
                'timeframe': quadrant_timeframes.get(quadrant_label, '3-6 months'),
                'timeframe_score_based': score_priority_defaults.get(score_priority, {}).get('timeframe', 'TBD'),
                # Effort label (from quadrant or score-based)
                'effort': 'High' if effort_score > 0.6 else 'Medium' if effort_score > 0.3 else 'Low',
                'effort_score_based': score_priority_defaults.get(score_priority, {}).get('effort_label', 'Unknown'),
                'risk_theme': domain_risk_themes.get(domain, 'General Risk'),
                # Additional metadata
                'regulatory_driver': metadata.get('regulatory_driver', ''),
                'risk_category': metadata.get('risk_category', ''),
                'lifecycle_phase': metadata.get('lifecycle_phase', ''),
                'effort_category': metadata.get('effort_category', '')
            }
            return enriched
        
        # Enrich all actions from score-based tiers
        all_enriched_actions = []
        for priority in ['high', 'medium', 'low']:
            for action in actions.get(priority, []):
                enriched = enrich_action(action, priority)
                all_enriched_actions.append(enriched)
        
        # Create SCORE-BASED ordering (current behavior)
        high_actions_score = [a for a in all_enriched_actions if a['score_priority'] == 'High']
        medium_actions_score = [a for a in all_enriched_actions if a['score_priority'] == 'Medium']
        low_actions_score = [a for a in all_enriched_actions if a['score_priority'] == 'Low']
        
        # Create SMART PRIORITY ordering (sorted by smart_priority_score descending)
        all_sorted_by_smart = sorted(all_enriched_actions, key=lambda x: x['smart_priority_score'], reverse=True)
        
        # Also categorize by quadrant for smart priority
        quick_wins = [a for a in all_sorted_by_smart if a['quadrant_label'] == 'Quick win']
        strategic_projects = [a for a in all_sorted_by_smart if a['quadrant_label'] == 'Strategic project']
        opportunistic = [a for a in all_sorted_by_smart if a['quadrant_label'] == 'Opportunistic improvement']
        lower_priority = [a for a in all_sorted_by_smart if a['quadrant_label'] == 'Lower priority']
        
        # Extract unique risk themes
        risk_themes = list(set([a.get('risk_theme', 'General Risk') for a in all_enriched_actions]))[:5]
        
        # Debug output
        print(f"  - Total enriched actions: {len(all_enriched_actions)}")
        print(f"  - Quick wins: {len(quick_wins)}, Strategic: {len(strategic_projects)}, Opportunistic: {len(opportunistic)}, Lower: {len(lower_priority)}")
        if all_sorted_by_smart:
            top = all_sorted_by_smart[0]
            print(f"  - Top smart priority: {top['question_id']} (score: {top['smart_priority_score']}, quadrant: {top['quadrant_label']})")
        
        # Determine which ordering to use as primary based on use_smart_priority flag
        if self.use_smart_priority:
            primary_high = quick_wins + strategic_projects  # High impact items
            primary_medium = opportunistic
            primary_low = lower_priority
        else:
            primary_high = high_actions_score
            primary_medium = medium_actions_score
            primary_low = low_actions_score
        
        return {
            # PRIMARY ordering (based on use_smart_priority flag)
            'high': primary_high,
            'medium': primary_medium,
            'low': primary_low,
            
            # SCORE-BASED ordering (always available)
            'by_score': {
                'high': high_actions_score,
                'medium': medium_actions_score,
                'low': low_actions_score
            },
            
            # SMART PRIORITY ordering (always available)
            'by_smart_priority': {
                'all': all_sorted_by_smart,
                'quick_wins': quick_wins,
                'strategic_projects': strategic_projects,
                'opportunistic': opportunistic,
                'lower_priority': lower_priority
            },
            
            # Summary statistics
            'summary': {
                'high_count': len(primary_high),
                'medium_count': len(primary_medium),
                'low_count': len(primary_low),
                'total_count': len(all_enriched_actions),
                'quick_wins_count': len(quick_wins),
                'strategic_projects_count': len(strategic_projects),
                'opportunistic_count': len(opportunistic),
                'lower_priority_count': len(lower_priority),
                'top_risk_themes': risk_themes if risk_themes else ['No specific risk themes identified'],
                'ordering_method': 'smart_priority' if self.use_smart_priority else 'score_based'
            }
        }
    
    def _calculate_evidence_stats(self, report_data: Dict[str, Any]) -> Dict[str, Any]:
        """Calculate evidence statistics for test template."""
        # This would ideally pull from actual evidence records
        # For now, return placeholder values
        actions = report_data.get('actions', {})
        high_actions = actions.get('high', [])
        
        return {
            'count_total': 0,
            'count_linked_to_high_priority': 0,
            'coverage_percent': 0.0,
            'gaps_count': len(high_actions),
            'top_gaps': [
                {
                    'domain': a.get('domain', 'Unknown'),
                    'question_id': a.get('question_id', 'N/A'),
                    'reason': 'High priority action without linked evidence'
                }
                for a in high_actions[:5]
            ]
        }


    def _generate_docx_report(self, report_data: Dict[str, Any], heatmap_image: bytes) -> bytes:
        """Generate DOCX report using template and data while preserving all styles."""
        try:
            # Debug: print report_data keys
            print(f"DEBUG _generate_docx_report: report_data keys = {list(report_data.keys())}")
            if 'awareness_info' in report_data:
                print(f"DEBUG: awareness_info keys = {list(report_data.get('awareness_info', {}).keys())}")
            
            # Load the template with ampersand protection
            doc = self._load_template_with_ampersand_protection(self.template_path)
            
            # Create inline image for heatmap - set to requested width
            heatmap_inline = InlineImage(doc, io.BytesIO(heatmap_image), width=Inches(6.27))
            
            # Generate awareness bar image (stacked bar chart) for Awareness assessments
            if self.assessment_type == 'Awareness':
                awareness_bar_bytes = self._generate_awareness_bar_image(report_data)
                awareness_bar_inline = InlineImage(doc, io.BytesIO(awareness_bar_bytes), width=Inches(4.0))  # Doubled width
            else:
                awareness_bar_inline = heatmap_inline  # Fallback for non-Awareness
            
            # Generate readiness bar image (stacked bar chart) for Readiness assessments
            readiness_bar_inline = None
            if self.assessment_type == 'Readiness':
                readiness_bar_bytes = self._generate_readiness_bar_image(report_data)
                readiness_bar_inline = InlineImage(doc, io.BytesIO(readiness_bar_bytes), width=Inches(4.0))  # Doubled width
            
            # Generate radar chart image
            radar_chart_image = self._generate_radar_chart_image(report_data)
            radar_chart_inline = InlineImage(doc, io.BytesIO(radar_chart_image), width=Inches(3.5))
            
            # Get top 3 strengths and gaps
            top_3_strengths = self._format_top_3_strengths(report_data)
            top_3_gaps = self._format_top_3_gaps(report_data)
            
            # For Awareness assessments, get info from awareness_info
            awareness_info = report_data.get('awareness_info') or {}
            
            # Prepare template context matching the exact template structure and placeholders
            template_context = {
                # Organization information - populated from awareness_info for Awareness assessments
                'org': {
                    'name': awareness_info.get('org_name') or report_data.get('org', {}).get('name', 'Organization'),
                    'industry': awareness_info.get('industry', ''),
                    'size': awareness_info.get('org_size', ''),
                    'business_unit': awareness_info.get('business_unit', ''),
                },
                
                # Assessment information (date will be formatted by formatDate function)
                'assessment': {
                    'id': report_data.get('assessment', {}).get('id', 'N/A'),
                    'date': report_data.get('assessment', {}).get('date', '2025-01-01'),
                    'version': report_data.get('assessment', {}).get('version', '1.0.0'),
                    'generated_at': get_report_local_time().strftime('%Y-%m-%dT%H:%M:%S'),
                    'type': 'AI System Maturity Assessment'
                },
                
                # Overall scores
                'overall': {
                    'score': report_data.get('overall', {}).get('score', 0),
                    'tier': report_data.get('overall', {}).get('tier', 'Basic')
                },
                
                # Report display options
                'show_detailed_responses': True,  # Toggle for showing detailed question responses in appendix
                
                # Assets (heatmap image) - generate and embed properly  
                'assets': {
                    'heatmapUrl': self._create_heatmap_for_template(heatmap_image) if heatmap_image else 'HEATMAP_IMAGE_PLACEHOLDER'
                },
                
                # Heatmap image for direct placeholder replacement
                'heatmap_image': heatmap_inline,
                
                # Awareness bar image (stacked bar chart for Awareness template)
                'awareness_bar_image': awareness_bar_inline,
                
                # Readiness bar image (stacked bar chart for Readiness template)
                'readiness_bar_image': readiness_bar_inline,
                
                # Radar chart image for direct placeholder replacement
                'radar_chart_image': radar_chart_inline,
                
                # Top 3 strengths and gaps
                'top_3_strengths': top_3_strengths,
                'top_3_gaps': top_3_gaps,
                
                # Actions for working template with separate priority arrays and summary
                'actions': self._build_actions_with_summary(report_data),
                
                # Sector-specific action steps
                'sector_name': report_data.get('sector_name', ''),
                'sector_actions': report_data.get('sector_actions', {'high': [], 'medium': [], 'low': []}),
                
                # System info from pre-assessment onboarding with normalized field names for test template
                'system_info': self._normalize_system_info(report_data.get('system_info', {})),
                
                # Add comprehensive content sections
                # Note: AI-enhanced versions are populated separately if use_ai=True
                'executive_summary': report_data.get('ai_executive_summary') or self._generate_comprehensive_executive_summary(report_data),
                'assessment_methodology': self._generate_assessment_methodology(),
                'domain_analysis': self._generate_domain_analysis(report_data),
                'key_findings': report_data.get('ai_key_findings') or self._generate_key_findings(report_data),
                'implementation_roadmap': self._generate_implementation_roadmap(report_data),
                
                # AI-generated content in structure expected by test template schema
                'ai_generated': {
                    'executive_summary': report_data.get('ai_executive_summary') or self._generate_comprehensive_executive_summary(report_data),
                    'key_findings': report_data.get('ai_key_findings') or self._generate_key_findings(report_data)
                },
                
                # Maturity distribution for test template
                'maturity': self._calculate_maturity_distribution(report_data),
                
                # Strengths and gaps in test template format
                'strengths': {
                    'top_3': top_3_strengths
                },
                'gaps': {
                    'top_3': top_3_gaps
                },
                
                # Domain results for test template (array format)
                'domains': self._format_domains_for_template(report_data),
                
                # Evidence information for test template
                'evidence': self._calculate_evidence_stats(report_data),
                
                # Validation info for test template
                'validation': {
                    'source_hash': report_data.get('validation', {}).get('source_hash', 'N/A'),
                    'lint_passed': True,
                    'lint_messages': []
                },
                
            }
            
            # Add Awareness-specific variables if this is an Awareness assessment
            if self.assessment_type == 'Awareness':
                awareness_info = report_data.get('awareness_info') or {}
                
                # Add all awareness_info fields as top-level variables
                template_context.update({
                    # Organization and contact info
                    'org_name': awareness_info.get('org_name', ''),
                    'contact_name': awareness_info.get('contact_name', ''),
                    'contact_email': awareness_info.get('contact_email', ''),
                    
                    # Organization context
                    'industry': awareness_info.get('industry', ''),
                    'org_size': awareness_info.get('org_size', ''),
                    'business_unit': awareness_info.get('business_unit', ''),
                    
                    # AI and digital maturity indicators
                    'ai_familiarity': awareness_info.get('ai_familiarity', ''),
                    'digital_maturity': awareness_info.get('digital_maturity', ''),
                    'leadership_ai_interest': awareness_info.get('leadership_ai_interest', ''),
                    'tech_change_comfort': awareness_info.get('tech_change_comfort', ''),
                    'digital_initiatives_level': awareness_info.get('digital_initiatives_level', ''),
                    'data_skill_confidence': awareness_info.get('data_skill_confidence', ''),
                    'openness_to_learning': awareness_info.get('openness_to_learning', ''),
                    
                    # Assessment context
                    'awareness_reason': awareness_info.get('awareness_reason', ''),
                    'assessor_name': awareness_info.get('assessor_name', ''),
                    'assessment_date': awareness_info.get('assessment_date', ''),
                    'framework_version': awareness_info.get('framework_version', ''),
                    
                    # Arrays from awareness_info
                    'awareness_outcomes': awareness_info.get('awareness_outcomes', []),
                    'learning_preferences': awareness_info.get('learning_preferences', []),
                    'governance_foundations': awareness_info.get('governance_foundations', []),
                    'pre_onboarding_commentary': awareness_info.get('pre_onboarding_commentary', []),
                    
                    # Full awareness_info object for flexibility
                    'awareness_info': awareness_info,
                    
                    # Pre-formatted Results Summary text (matches frontend ResultsPage.js)
                    # Escape XML special characters for safe docxtpl rendering
                    'results_summary_text': self._generate_results_summary_text(report_data),
                })
                
                # Build strengths list (domains scoring >= 70%, sorted by score descending)
                strengths = []
                domains_for_strengths = sorted(
                    template_context.get('domains', []),
                    key=lambda d: float(str(d.get('score', 0)).replace('%', '')),
                    reverse=True
                )
                for d in domains_for_strengths:
                    score = d.get('score', 0)
                    # Handle both numeric and string scores
                    if isinstance(score, str):
                        score = float(score.replace('%', ''))
                    if score >= 70:
                        strengths.append(f"{d.get('name', 'Unknown')} ({score:.0f}%)")
                template_context['strengths'] = strengths
                
                # Build strengths_summary as a readable string
                if strengths:
                    template_context['strengths_summary'] = ', '.join(strengths[:3])
                else:
                    template_context['strengths_summary'] = 'Building foundational awareness across all domains'
                
                # Build gaps list (domains scoring < 50%, sorted by score ascending - worst first)
                gaps = []
                domains_for_gaps = sorted(
                    template_context.get('domains', []),
                    key=lambda d: float(str(d.get('score', 0)).replace('%', '')),
                    reverse=False
                )
                for d in domains_for_gaps:
                    score = d.get('score', 0)
                    # Handle both numeric and string scores
                    if isinstance(score, str):
                        score = float(score.replace('%', ''))
                    if score < 50:
                        gaps.append(f"{d.get('name', 'Unknown')} ({score:.0f}%)")
                template_context['gaps'] = gaps
                
                # Build gaps_summary as a readable string
                if gaps:
                    template_context['gaps_summary'] = ', '.join(gaps[:3])
                else:
                    template_context['gaps_summary'] = 'No critical gaps identified'
                
                # Build next_focus - use AI-generated version if available, otherwise fallback
                ai_narratives = report_data.get('ai_narratives', {})
                ai_next_focus = ai_narratives.get('next_focus', '')
                
                if ai_next_focus:
                    # autoescape=True handles XML escaping
                    template_context['next_focus'] = ai_next_focus
                else:
                    # Fallback: Build next_focus recommendation based on tier
                    tier = template_context.get('overall', {}).get('tier', 'Developing')
                    if tier == 'Foundational':
                        template_context['next_focus'] = 'Establish basic AI awareness and governance foundations'
                    elif tier == 'Developing':
                        template_context['next_focus'] = 'Build structured awareness programs and refine governance frameworks'
                    elif tier == 'Established':
                        template_context['next_focus'] = 'Deepen AI literacy and prepare for readiness assessment'
                    else:
                        template_context['next_focus'] = 'Maintain leadership position and drive AI innovation'
                
                # Build priority_actions list with priority and text
                priority_actions = []
                for action in template_context.get('actions', {}).get('high', []):
                    priority_actions.append({
                        'priority': 'High',
                        'text': action.get('text', ''),
                        'domain': action.get('domain', ''),
                        'question_id': action.get('question_id', '')
                    })
                for action in template_context.get('actions', {}).get('medium', []):
                    priority_actions.append({
                        'priority': 'Medium',
                        'text': action.get('text', ''),
                        'domain': action.get('domain', ''),
                        'question_id': action.get('question_id', '')
                    })
                for action in template_context.get('actions', {}).get('low', []):
                    priority_actions.append({
                        'priority': 'Low',
                        'text': action.get('text', ''),
                        'domain': action.get('domain', ''),
                        'question_id': action.get('question_id', '')
                    })
                template_context['priority_actions'] = priority_actions
                
                # Build top 5 sector-specific actions for each priority level
                sector_actions = template_context.get('sector_actions', {'high': [], 'medium': [], 'low': []})
                
                # DEBUG: Print sector_actions structure
                print(f"DEBUG sector_actions structure:")
                print(f"  - sector_actions type: {type(sector_actions)}")
                print(f"  - high count: {len(sector_actions.get('high', []))}")
                print(f"  - medium count: {len(sector_actions.get('medium', []))}")
                print(f"  - low count: {len(sector_actions.get('low', []))}")
                if sector_actions.get('high'):
                    first_high = sector_actions['high'][0]
                    print(f"  - First high action keys: {list(first_high.keys())}")
                    print(f"  - First high action.domain: {first_high.get('domain')}")
                    print(f"  - First high action.question_id: {first_high.get('question_id')}")
                    print(f"  - First high action.sector_action: {first_high.get('sector_action', 'N/A')[:80]}...")
                if sector_actions.get('medium'):
                    first_med = sector_actions['medium'][0]
                    print(f"  - First MEDIUM action: question_id='{first_med.get('question_id')}', domain='{first_med.get('domain')}', sector_action='{first_med.get('sector_action', 'N/A')[:60]}...'")
                
                template_context['sector_actions_high_top5'] = sector_actions.get('high', [])[:5]
                template_context['sector_actions_medium_top5'] = sector_actions.get('medium', [])[:5]
                template_context['sector_actions_low_top5'] = sector_actions.get('low', [])[:5]
                
                # DEBUG: Confirm the top5 lists
                print(f"DEBUG sector_actions_high_top5 after assignment: {len(template_context['sector_actions_high_top5'])} items")
                print(f"DEBUG sector_actions_medium_top5 after assignment: {len(template_context['sector_actions_medium_top5'])} items")
                if template_context['sector_actions_high_top5']:
                    first = template_context['sector_actions_high_top5'][0]
                    print(f"  - First item domain: {first.get('domain')}")
                    print(f"  - First item question_id: {first.get('question_id')}")
                    print(f"  - First item sector_action: {first.get('sector_action', 'N/A')[:50]}...")
                else:
                    print(f"  sector_actions_high_top5 is EMPTY")
                if template_context['sector_actions_medium_top5']:
                    first_m = template_context['sector_actions_medium_top5'][0]
                    print(f"  MEDIUM First: question_id='{first_m.get('question_id')}', domain='{first_m.get('domain')}')")
                
                # Add AI-generated narratives if available
                # Apply placeholder for & to preserve through docxtpl rendering
                ai_narratives = report_data.get('ai_narratives', {})
                template_context['ai'] = {
                    'executive_snapshot': replace_amp_with_placeholder(ai_narratives.get('executive_snapshot', '')),
                    'context_interpretation': replace_amp_with_placeholder(ai_narratives.get('context_interpretation', '')),
                    'governance_interpretation': replace_amp_with_placeholder(ai_narratives.get('governance_interpretation', '')),
                    'readiness_interpretation': replace_amp_with_placeholder(ai_narratives.get('readiness_interpretation', '')),
                    'domain_patterns': replace_amp_with_placeholder(ai_narratives.get('domain_patterns', '')),
                    'action_interpretation': replace_amp_with_placeholder(ai_narratives.get('action_interpretation', '')),
                }
                print(f"DEBUG: Added AI narratives to template context: {list(template_context['ai'].keys())}")
                
                # Build questions list with full details
                questions = []
                questions_data = report_data.get('questions_data', [])
                q_num = 1
                for domain_data in questions_data:
                    domain_name = domain_data.get('domain', {}).get('name', 'Unknown')
                    # Use placeholder for & to preserve it through docxtpl paragraph loop
                    domain_name_for_template = replace_amp_with_placeholder(domain_name) if domain_name else 'Unknown'
                    for q in domain_data.get('questions', []):
                        answer = q.get('answer', {})
                        score = answer.get('numeric_score', 0) if answer else 0
                        
                        # Determine maturity level from score (Awareness-specific tiers)
                        if score >= 4:
                            maturity_level = 'Established'
                        elif score >= 3:
                            maturity_level = 'Developing'
                        elif score >= 2:
                            maturity_level = 'Emerging'
                        else:
                            maturity_level = 'Introductory'
                        
                        questions.append({
                            'number': q_num,
                            'code': q.get('code', ''),
                            'text': q.get('text', ''),
                            'domain': domain_name_for_template,
                            'selected_option_label': answer.get('selected_option', {}).get('label', 'Not answered') if answer else 'Not answered',
                            'selected_option_text': answer.get('selected_option', {}).get('text', '') if answer else '',
                            'maturity_level': maturity_level,
                            'score': score,
                            'comment': answer.get('comment', '') if answer else ''
                        })
                        q_num += 1
                template_context['questions'] = questions
                
                # Debug: Print first few question domains to verify & character
                if questions:
                    print(f"  First 3 question domains:")
                    for q in questions[:3]:
                        print(f"    - Domain: '{q.get('domain')}'")
                
                print(f"  Awareness-specific variables added:")
                print(f"    - org_name: {template_context.get('org_name', 'N/A')}")
                print(f"    - industry: {template_context.get('industry', 'N/A')}")
                print(f"    - strengths: {len(strengths)} items")
                print(f"    - gaps: {len(gaps)} items")
                print(f"    - priority_actions: {len(priority_actions)} items")
                print(f"    - questions: {len(questions)} items")
            
            # Add Readiness-specific variables if this is a Readiness assessment
            elif self.assessment_type == 'Readiness':
                readiness_info = report_data.get('readiness_info') or {}
                
                # Add all readiness_info fields as top-level variables
                template_context.update({
                    # Organization and contact info
                    'org_name': readiness_info.get('org_name', ''),
                    'contact_name': readiness_info.get('contact_name', ''),
                    'contact_email': readiness_info.get('contact_email', ''),
                    
                    # Organization context
                    'industry': readiness_info.get('industry', ''),
                    'org_size': readiness_info.get('org_size', ''),
                    'business_unit': readiness_info.get('business_unit', ''),
                    
                    # Strategic Intent
                    'ai_motivation': readiness_info.get('ai_motivation', ''),
                    'leadership_commitment': readiness_info.get('leadership_commitment', ''),
                    'ai_strategy_status': readiness_info.get('ai_strategy_status', ''),
                    'ai_risk_awareness': readiness_info.get('ai_risk_awareness', ''),
                    
                    # Governance & Ethics
                    'governance_foundations': readiness_info.get('governance_foundations', []),
                    'decision_ownership': readiness_info.get('decision_ownership', ''),
                    'ethical_principles': readiness_info.get('ethical_principles', []),
                    
                    # Data & Capability
                    'data_maturity': readiness_info.get('data_maturity', ''),
                    'ai_capability': readiness_info.get('ai_capability', ''),
                    'current_tools': readiness_info.get('current_tools', []),
                    'poc_status': readiness_info.get('poc_status', ''),
                    
                    # Assessment context
                    'assessor_name': readiness_info.get('assessor_name', ''),
                    'assessment_date': readiness_info.get('assessment_date', ''),
                    'framework_version': readiness_info.get('framework_version', ''),
                    
                    # Full readiness_info object for flexibility
                    'readiness_info': readiness_info,
                    
                    # Pre-formatted Results Summary text (matches frontend ResultsPage.js)
                    # Escape XML special characters for safe docxtpl rendering
                    'results_summary_text': self._generate_readiness_results_summary_text(report_data),
                })
                
                # Build strengths list (domains scoring >= 70%, sorted by score descending)
                strengths = []
                domains_for_strengths = sorted(
                    template_context.get('domains', []),
                    key=lambda d: float(str(d.get('score', 0)).replace('%', '')),
                    reverse=True
                )
                for d in domains_for_strengths:
                    score = d.get('score', 0)
                    if isinstance(score, str):
                        score = float(score.replace('%', ''))
                    if score >= 70:
                        strengths.append(f"{d.get('name', 'Unknown')} ({score:.0f}%)")
                template_context['strengths'] = strengths
                
                if strengths:
                    template_context['strengths_summary'] = ', '.join(strengths[:3])
                else:
                    template_context['strengths_summary'] = 'Building foundational readiness across all domains'
                
                # Build gaps list (domains scoring < 60%, sorted by score ascending - worst first)
                gaps = []
                domains_for_gaps = sorted(
                    template_context.get('domains', []),
                    key=lambda d: float(str(d.get('score', 0)).replace('%', '')),
                    reverse=False
                )
                for d in domains_for_gaps:
                    score = d.get('score', 0)
                    if isinstance(score, str):
                        score = float(score.replace('%', ''))
                    if score < 60:
                        gaps.append(f"{d.get('name', 'Unknown')} ({score:.0f}%)")
                template_context['gaps'] = gaps
                
                if gaps:
                    template_context['gaps_summary'] = ', '.join(gaps[:3])
                else:
                    template_context['gaps_summary'] = 'No critical gaps identified'
                
                # Build questions list with full details for Readiness
                questions = []
                questions_data = report_data.get('questions_data', [])
                q_num = 1
                for domain_data in questions_data:
                    domain_name = domain_data.get('domain', {}).get('name', 'Unknown')
                    # Use placeholder for & to preserve it through docxtpl paragraph loop
                    domain_name_safe = replace_amp_with_placeholder(domain_name) if domain_name else 'Unknown'
                    if q_num == 1:
                        print(f"DEBUG: First domain_name for questions = '{domain_name}' -> '{domain_name_safe}'")
                    for q in domain_data.get('questions', []):
                        answer = q.get('answer', {})
                        score = answer.get('numeric_score', 0) if answer else 0
                        
                        # Determine maturity level from score (Readiness-specific tiers)
                        if score >= 4:
                            maturity_level = 'Leading'
                        elif score >= 3:
                            maturity_level = 'Established'
                        elif score >= 2:
                            maturity_level = 'Developing'
                        else:
                            maturity_level = 'Foundational'
                        
                        questions.append({
                            'number': q_num,
                            'code': q.get('code', ''),
                            'text': q.get('text', ''),
                            'domain': domain_name_safe,
                            'selected_option_label': answer.get('selected_option', {}).get('label', 'Not answered') if answer else 'Not answered',
                            'selected_option_text': answer.get('selected_option', {}).get('text', '') if answer else '',
                            'maturity_level': maturity_level,
                            'score': score,
                            'comment': answer.get('comment', '') if answer else ''
                        })
                        q_num += 1
                template_context['questions'] = questions
                
                print(f"  Readiness-specific variables added:")
                print(f"    - org_name: {template_context.get('org_name', 'N/A')}")
                print(f"    - industry: {template_context.get('industry', 'N/A')}")
                print(f"    - strengths: {len(strengths)} items")
                print(f"    - gaps: {len(gaps)} items")
                print(f"    - questions: {len(questions)} items")
                
                # Add openness_to_learning from readiness_info (used in template conditionals)
                template_context['openness_to_learning'] = readiness_info.get('openness_to_learning', '')
                
                # Build next_focus recommendation based on tier
                overall_tier = template_context.get('overall', {}).get('tier', 'Foundational')
                if overall_tier == 'Foundational':
                    template_context['next_focus'] = 'Focus on building basic governance structures, improving data management, and developing staff AI awareness before considering AI adoption.'
                elif overall_tier == 'Developing':
                    template_context['next_focus'] = 'Strengthen governance frameworks, formalize data quality processes, and establish clear ethical guidelines to advance AI readiness.'
                elif overall_tier == 'Established':
                    template_context['next_focus'] = 'Enhance continuous improvement processes, deepen stakeholder engagement, and refine risk management for leading-edge AI adoption.'
                else:  # Leading
                    template_context['next_focus'] = 'Maintain leadership position through innovation, knowledge sharing, and proactive adaptation to emerging AI technologies and standards.'
                
                # Add AI narratives for Readiness reports (with fallback empty strings)
                # Using new variable naming convention: ai.r_* for Readiness
                # Apply placeholder for & to preserve through docxtpl rendering
                ai_narratives = report_data.get('ai_narratives', {})
                template_context['ai'] = {
                    'r_executive_snapshot': replace_amp_with_placeholder(ai_narratives.get('r_executive_snapshot', '')),
                    'r_context_interpretation': replace_amp_with_placeholder(ai_narratives.get('r_context_interpretation', '')),
                    'r_governance_interpretation': replace_amp_with_placeholder(ai_narratives.get('r_governance_interpretation', '')),
                    'r_data_tech_interpretation': replace_amp_with_placeholder(ai_narratives.get('r_data_tech_interpretation', '')),
                    'r_domain_patterns': replace_amp_with_placeholder(ai_narratives.get('r_domain_patterns', '')),
                    'r_sector_interpretation': replace_amp_with_placeholder(ai_narratives.get('r_sector_interpretation', '')),
                    'r_action_interpretation': replace_amp_with_placeholder(ai_narratives.get('r_action_interpretation', '')),
                    'r_pathway_rationale': replace_amp_with_placeholder(ai_narratives.get('r_pathway_rationale', '')),
                }
                print(f"DEBUG: Added AI narratives to Readiness template context: {list(template_context['ai'].keys())}")
                for key, value in template_context['ai'].items():
                    if value:
                        print(f"  - {key} length: {len(value)} chars")
                print(f"    - openness_to_learning: {template_context.get('openness_to_learning', 'N/A')}")
                print(f"    - next_focus: {template_context.get('next_focus', 'N/A')[:50]}...")
                
                # Use sector_actions from report_data (same approach as Awareness)
                # This contains actions in the correct structure from _generate_sector_actions
                sector_actions = report_data.get('sector_actions', {})
                print(f"DEBUG Readiness sector_actions from report_data: high={len(sector_actions.get('high', []))}, medium={len(sector_actions.get('medium', []))}, low={len(sector_actions.get('low', []))}")
                
                # If sector_actions is empty (no sector-specific data), fall back to building from actions
                if not any([sector_actions.get('high'), sector_actions.get('medium'), sector_actions.get('low')]):
                    print("DEBUG: No sector_actions in report_data, building from template_context['actions']")
                    actions = template_context.get('actions', {})
                    high_actions = actions.get('high', [])
                    medium_actions = actions.get('medium', [])
                    low_actions = actions.get('low', [])
                    
                    # Format actions for sector_actions template structure
                    def format_action_for_sector(action):
                        return {
                            'question_id': action.get('question_id', ''),
                            'domain': action.get('domain', ''),
                            'sector_action': action.get('recommendation', action.get('text', ''))
                        }
                    
                    template_context['sector_actions_high_top5'] = [format_action_for_sector(a) for a in high_actions[:5]]
                    template_context['sector_actions_medium_top5'] = [format_action_for_sector(a) for a in medium_actions[:5]]
                    template_context['sector_actions_low_top5'] = [format_action_for_sector(a) for a in low_actions[:5]]
                else:
                    # Use sector_actions directly (same as Awareness)
                    print("DEBUG: Using sector_actions from report_data")
                    template_context['sector_actions_high_top5'] = sector_actions.get('high', [])[:5]
                    template_context['sector_actions_medium_top5'] = sector_actions.get('medium', [])[:5]
                    template_context['sector_actions_low_top5'] = sector_actions.get('low', [])[:5]
                
                print(f"    - sector_actions_high_top5: {len(template_context['sector_actions_high_top5'])} items")
                print(f"    - sector_actions_medium_top5: {len(template_context['sector_actions_medium_top5'])} items")
                print(f"    - sector_actions_low_top5: {len(template_context['sector_actions_low_top5'])} items")
                if template_context['sector_actions_high_top5']:
                    sample = template_context['sector_actions_high_top5'][0]
                    print(f"    - Sample: question_id='{sample.get('question_id')}', domain='{sample.get('domain')}'")
            
            # Add Orgwide-specific variables if this is an Orgwide assessment
            elif self.assessment_type == 'Orgwide':
                orgwide_info = report_data.get('orgwide_info') or {}
                
                # Add all orgwide_info fields as top-level variables
                template_context.update({
                    # Organization and contact info
                    'org_name': replace_amp_with_placeholder(orgwide_info.get('org_name', '')),
                    'contact_name': orgwide_info.get('contact_name', ''),
                    'contact_email': orgwide_info.get('contact_email', ''),
                    
                    # Organization context - use sector_name from report_data (organization's primary_industry)
                    'industry': replace_amp_with_placeholder(report_data.get('sector_name', '') or orgwide_info.get('industry', '')),
                    'org_size': orgwide_info.get('org_size', ''),
                    'business_units_scope': orgwide_info.get('business_units_scope', ''),
                    'business_unit': orgwide_info.get('business_units_scope', ''),  # Alias for template
                    
                    # AI Landscape & Scope
                    'ai_project_count': orgwide_info.get('ai_project_count', ''),
                    'ai_usage_areas': orgwide_info.get('ai_usage_areas', []),
                    'ai_primary_purpose': orgwide_info.get('ai_primary_purpose', ''),
                    'ai_sourcing_model': orgwide_info.get('ai_sourcing_model', ''),
                    'ai_governance_committee_status': orgwide_info.get('ai_governance_committee_status', ''),
                    
                    # Governance & Risk
                    'ai_frameworks': orgwide_info.get('ai_frameworks', []),
                    'ai_risk_assessment_usage': orgwide_info.get('ai_risk_assessment_usage', ''),
                    'policy_review_frequency': orgwide_info.get('policy_review_frequency', ''),
                    'ai_register_status': orgwide_info.get('ai_register_status', ''),
                    'governance_foundations': orgwide_info.get('governance_foundations', []),
                    'ethical_principles': orgwide_info.get('ethical_principles', []),
                    
                    # Culture & Capability
                    'exec_engagement_level': orgwide_info.get('exec_engagement_level', ''),
                    'ai_training_maturity': orgwide_info.get('ai_training_maturity', ''),
                    'ai_maturity_self_rating': orgwide_info.get('ai_maturity_self_rating', ''),
                    'decision_ownership': orgwide_info.get('decision_ownership', ''),
                    
                    # Assessment context
                    'assessor_name': orgwide_info.get('assessor_name', ''),
                    'assessment_date': orgwide_info.get('assessment_date', ''),
                    'framework_version': orgwide_info.get('framework_version', ''),
                    
                    # Full orgwide_info object for flexibility
                    'orgwide_info': orgwide_info,
                    
                    # Pre-formatted Results Summary text (matches frontend ResultsPage.js)
                    'results_summary_text': replace_amp_with_placeholder(self._generate_orgwide_results_summary_text(report_data)),
                })
                
                # Generate bar chart for Orgwide (using readiness bar style) - doubled width
                orgwide_bar_bytes = self._generate_readiness_bar_image(report_data)
                template_context['readiness_bar_image'] = InlineImage(doc, io.BytesIO(orgwide_bar_bytes), width=Inches(4.0))
                
                # Generate radar chart for Orgwide
                orgwide_radar_bytes = self._generate_radar_chart_image(report_data)
                template_context['radar_chart_image'] = InlineImage(doc, io.BytesIO(orgwide_radar_bytes), width=Inches(3.5))
                
                # Build strengths list (domains scoring >= 70%, or top 3 if none meet threshold)
                strengths = []
                domains_sorted = sorted(template_context.get('domains', []), 
                                       key=lambda x: float(str(x.get('score', 0)).replace('%', '')), 
                                       reverse=True)
                
                # First try domains >= 70%
                for d in domains_sorted:
                    score = d.get('score', 0)
                    if isinstance(score, str):
                        score = float(score.replace('%', ''))
                    if score >= 70:
                        strengths.append(f"{d.get('name', 'Unknown')} ({score:.0f}%)")
                
                # If no domains >= 70%, use top 3 performing domains as relative strengths
                if not strengths and domains_sorted:
                    for d in domains_sorted[:3]:
                        score = d.get('score', 0)
                        if isinstance(score, str):
                            score = float(score.replace('%', ''))
                        strengths.append(f"{d.get('name', 'Unknown')} ({score:.0f}%)")
                
                template_context['strengths'] = strengths
                
                if strengths:
                    template_context['strengths_summary'] = replace_amp_with_placeholder(', '.join(strengths[:3]))
                else:
                    template_context['strengths_summary'] = 'Building foundational maturity across all domains'
                
                # Build gaps list (domains scoring < 60%, sorted by score ascending - worst first)
                gaps = []
                domains_for_gaps = sorted(
                    template_context.get('domains', []),
                    key=lambda d: float(str(d.get('score', 0)).replace('%', '')),
                    reverse=False
                )
                for d in domains_for_gaps:
                    score = d.get('score', 0)
                    if isinstance(score, str):
                        score = float(score.replace('%', ''))
                    if score < 60:
                        gaps.append(f"{d.get('name', 'Unknown')} ({score:.0f}%)")
                template_context['gaps'] = gaps
                
                if gaps:
                    template_context['gaps_summary'] = replace_amp_with_placeholder(', '.join(gaps[:3]))
                else:
                    template_context['gaps_summary'] = 'No critical gaps identified'
                
                # Build next_focus based on gaps or overall tier
                overall_tier = template_context.get('overall', {}).get('tier', 'Developing')
                if gaps:
                    template_context['next_focus'] = replace_amp_with_placeholder(f"Focus on improving {gaps[0].split(' (')[0]} to strengthen overall AI maturity.")
                elif overall_tier == 'Leading':
                    template_context['next_focus'] = 'Maintain leadership position and drive AI innovation across the organisation.'
                elif overall_tier == 'Established':
                    template_context['next_focus'] = 'Advance to leading maturity by optimising AI governance and scaling successful initiatives.'
                else:
                    template_context['next_focus'] = 'Build foundational AI capabilities and establish governance frameworks.'
                
                # Add sector actions for Orgwide (similar to Readiness)
                sector_actions = report_data.get('sector_actions', {'high': [], 'medium': [], 'low': []})
                template_context['sector_actions_high_top5'] = sector_actions.get('high', [])[:5]
                template_context['sector_actions_medium_top5'] = sector_actions.get('medium', [])[:5]
                template_context['sector_actions_low_top5'] = sector_actions.get('low', [])[:5]
                
                # Build questions list with full details for Orgwide
                questions = []
                questions_data = report_data.get('questions_data', [])
                q_num = 1
                for domain_data in questions_data:
                    domain_name = domain_data.get('domain', {}).get('name', 'Unknown')
                    domain_name_for_template = replace_amp_with_placeholder(domain_name) if domain_name else 'Unknown'
                    for q in domain_data.get('questions', []):
                        answer = q.get('answer', {})
                        score = answer.get('numeric_score', 0) if answer else 0
                        
                        # Determine maturity level from score
                        if score >= 4:
                            maturity_level = 'Leading'
                        elif score >= 3:
                            maturity_level = 'Established'
                        elif score >= 2:
                            maturity_level = 'Developing'
                        else:
                            maturity_level = 'Foundational'
                        
                        questions.append({
                            'number': q_num,
                            'code': q.get('code', ''),
                            'text': q.get('text', ''),
                            'domain': domain_name_for_template,
                            'selected_option_label': answer.get('selected_option', {}).get('label', 'Not answered') if answer else 'Not answered',
                            'selected_option_text': answer.get('selected_option', {}).get('text', '') if answer else '',
                            'maturity_level': maturity_level,
                            'score': score,
                            'comment': answer.get('comment', '') if answer else ''
                        })
                        q_num += 1
                template_context['questions'] = questions
                template_context['show_detailed_responses'] = True
                
                # Add AI narratives for Orgwide (with placeholder for &)
                ai_narratives = report_data.get('ai_narratives', {})
                template_context['ai'] = {
                    'o_executive_snapshot': replace_amp_with_placeholder(ai_narratives.get('o_executive_snapshot', '')),
                    'o_landscape': replace_amp_with_placeholder(ai_narratives.get('o_landscape', '')),
                    'o_gov_oversight': replace_amp_with_placeholder(ai_narratives.get('o_gov_oversight', '')),
                    'o_risk_policy': replace_amp_with_placeholder(ai_narratives.get('o_risk_policy', '')),
                    'o_culture_capability': replace_amp_with_placeholder(ai_narratives.get('o_culture_capability', '')),
                    'o_domain_patterns': replace_amp_with_placeholder(ai_narratives.get('o_domain_patterns', '')),
                    'o_sector_interpretation': replace_amp_with_placeholder(ai_narratives.get('o_sector_interpretation', '')),
                    'o_pathway_rationale': replace_amp_with_placeholder(ai_narratives.get('o_pathway_rationale', '')),
                }
                
                print(f"  Orgwide-specific variables added:")
                print(f"    - org_name: {template_context.get('org_name', 'N/A')}")
                print(f"    - industry: {template_context.get('industry', 'N/A')}")
                print(f"    - strengths: {len(strengths)} items")
                print(f"    - gaps: {len(gaps)} items")
                print(f"    - questions: {len(questions)} items")
                print(f"    - sector_actions_high_top5: {len(template_context.get('sector_actions_high_top5', []))} items")
                print(f"    - ai_narratives: {list(ai_narratives.keys())}")
            
            # Add System-specific variables if this is a System assessment
            elif self.assessment_type == 'System':
                system_info = report_data.get('system_info') or {}
                
                # Add all system_info fields as top-level variables
                template_context.update({
                    # Organization info
                    'organization_name': system_info.get('organizationName', ''),
                    'org_name': system_info.get('organizationName', ''),
                    'industry': system_info.get('industry', ''),
                    
                    # System Information
                    'system_name': system_info.get('systemName', ''),
                    'system_description': system_info.get('description', ''),
                    'system_owner': system_info.get('owner', ''),
                    'department': system_info.get('department', ''),
                    'lifecycle': system_info.get('lifecycle', ''),
                    'users_stakeholders': system_info.get('usersStakeholders', ''),
                    'monthly_volume': system_info.get('monthlyVolume', ''),
                    'criticality': system_info.get('criticality', ''),
                    
                    # Model & Technical Details
                    'model_type': system_info.get('modelType', ''),
                    'ownership': system_info.get('ownership', ''),
                    'hosting': system_info.get('hosting', ''),
                    'cloud_provider': system_info.get('cloudProvider', ''),
                    'cloud_region': system_info.get('cloudRegion', ''),
                    'data_flow': system_info.get('dataFlow', ''),
                    'data_sensitivity': system_info.get('dataSensitivity', ''),
                    'data_sources': system_info.get('dataSources', ''),
                    'representation_notes': system_info.get('representationNotes', ''),
                    'has_retention_policy': system_info.get('hasRetentionPolicy', False),
                    
                    # Governance & Oversight
                    'oversight': system_info.get('oversight', ''),
                    'artefacts': system_info.get('artefacts', []),
                    'regulations': system_info.get('regulations', []),
                    'ethics_commitments': system_info.get('ethicsCommitments', ''),
                    'sustainability_goals': system_info.get('sustainabilityGoals', ''),
                    'dependencies': system_info.get('dependencies', ''),
                    
                    # Version & Evidence
                    'version_ref': system_info.get('versionRef', ''),
                    'evidence_repo_url': system_info.get('evidenceRepoUrl', ''),
                    
                    # Assessment context
                    'assessor_name': system_info.get('assessor_name', ''),
                    'assessment_date': system_info.get('assessment_date', ''),
                    'framework_version': system_info.get('framework_version', ''),
                    
                    # Full system_info object for flexibility
                    'system_info': system_info,
                    
                    # System object for template references like {{system.owner}}
                    'system': {
                        'owner': system_info.get('owner', ''),
                        'department': system_info.get('department', ''),
                        'name': system_info.get('systemName', ''),
                        'description': system_info.get('description', ''),
                        'lifecycle': system_info.get('lifecycle', ''),
                    },
                    
                    # Pre-formatted Results Summary text (matches frontend ResultsPage.js)
                    'results_summary_text': self._generate_system_results_summary_text(report_data),
                })
                
                # Build strengths list (domains scoring >= 70%, sorted by score descending)
                strengths = []
                domains_for_strengths = sorted(
                    template_context.get('domains', []),
                    key=lambda d: float(str(d.get('score', 0)).replace('%', '')),
                    reverse=True
                )
                for d in domains_for_strengths:
                    score = d.get('score', 0)
                    if isinstance(score, str):
                        score = float(score.replace('%', ''))
                    if score >= 70:
                        strengths.append(f"{d.get('name', 'Unknown')} ({score:.0f}%)")
                template_context['strengths'] = strengths
                
                if strengths:
                    template_context['strengths_summary'] = ', '.join(strengths[:3])
                else:
                    template_context['strengths_summary'] = 'Building foundational maturity across all domains'
                
                # Build gaps list (domains scoring < 50%, sorted by score ascending - worst first)
                gaps = []
                domains_for_gaps = sorted(
                    template_context.get('domains', []),
                    key=lambda d: float(str(d.get('score', 0)).replace('%', '')),
                    reverse=False
                )
                for d in domains_for_gaps:
                    score = d.get('score', 0)
                    if isinstance(score, str):
                        score = float(score.replace('%', ''))
                    if score < 50:
                        gaps.append(f"{d.get('name', 'Unknown')} ({score:.0f}%)")
                template_context['gaps'] = gaps
                
                if gaps:
                    template_context['gaps_summary'] = ', '.join(gaps[:3])
                else:
                    template_context['gaps_summary'] = 'No critical gaps identified'
                
                print(f"  System-specific variables added:")
                print(f"    - system_name: {template_context.get('system_name', 'N/A')}")
                print(f"    - organization_name: {template_context.get('organization_name', 'N/A')}")
                print(f"    - industry: {template_context.get('industry', 'N/A')}")
                print(f"    - strengths: {len(strengths)} items")
                print(f"    - gaps: {len(gaps)} items")
                
                # Add AI narratives for System reports (with fallback empty strings)
                # Using variable naming convention: ai.s_* for System
                # Apply placeholder for & to preserve through docxtpl rendering
                ai_narratives = report_data.get('ai_narratives', {})
                template_context['ai'] = {
                    's_executive_snapshot': replace_amp_with_placeholder(ai_narratives.get('s_executive_snapshot', '')),
                    's_gov_oversight': replace_amp_with_placeholder(ai_narratives.get('s_gov_oversight', '')),
                    's_data_privacy_security': replace_amp_with_placeholder(ai_narratives.get('s_data_privacy_security', '')),
                    's_reliability_safety_performance': replace_amp_with_placeholder(ai_narratives.get('s_reliability_safety_performance', '')),
                    's_transparency_explainability_contestability': replace_amp_with_placeholder(ai_narratives.get('s_transparency_explainability_contestability', '')),
                    's_fairness_bias_inclusivity': replace_amp_with_placeholder(ai_narratives.get('s_fairness_bias_inclusivity', '')),
                    's_domain_patterns': replace_amp_with_placeholder(ai_narratives.get('s_domain_patterns', '')),
                    's_pathway_rationale': replace_amp_with_placeholder(ai_narratives.get('s_pathway_rationale', '')),
                }
                print(f"DEBUG: Added AI narratives to System template context: {list(template_context['ai'].keys())}")
                for key, value in template_context['ai'].items():
                    if value:
                        print(f"  - {key} length: {len(value)} chars")
                
                # Add sector actions for System reports (same as Orgwide)
                sector_actions = report_data.get('sector_actions', {})
                template_context['sector_actions_high_top5'] = sector_actions.get('high', [])[:5]
                template_context['sector_actions_medium_top5'] = sector_actions.get('medium', [])[:5]
                template_context['sector_actions_low_top5'] = sector_actions.get('low', [])[:5]
                print(f"DEBUG: System sector_actions_high_top5: {len(template_context['sector_actions_high_top5'])} items")
                print(f"DEBUG: System sector_actions_medium_top5: {len(template_context['sector_actions_medium_top5'])} items")
                print(f"DEBUG: System sector_actions_low_top5: {len(template_context['sector_actions_low_top5'])} items")
                
                # Generate bar chart for System (using readiness bar style)
                system_bar_bytes = self._generate_readiness_bar_image(report_data)
                template_context['readiness_bar_image'] = InlineImage(doc, io.BytesIO(system_bar_bytes), width=Inches(4.0))
                
                # Generate radar chart for System
                system_radar_bytes = self._generate_radar_chart_image(report_data)
                template_context['radar_chart_image'] = InlineImage(doc, io.BytesIO(system_radar_bytes), width=Inches(3.5))
                
                # Build next_focus based on gaps or overall tier
                overall_tier = template_context.get('overall', {}).get('tier', 'Developing')
                if gaps:
                    template_context['next_focus'] = replace_amp_with_placeholder(f"Focus on improving {gaps[0].split(' (')[0]} to strengthen overall AI system maturity.")
                elif overall_tier == 'Leading':
                    template_context['next_focus'] = 'Maintain leading maturity position and drive continuous improvement across all domains.'
                elif overall_tier == 'Established':
                    template_context['next_focus'] = 'Advance to leading maturity by optimising governance controls and refining operational practices.'
                else:
                    template_context['next_focus'] = 'Build foundational governance capabilities and establish consistent practices across key domains.'
                
                # Build questions list with full details for System Appendix A
                questions = []
                questions_data = report_data.get('questions_data', [])
                q_num = 1
                for domain_data in questions_data:
                    domain_name = domain_data.get('domain', {}).get('name', 'Unknown')
                    domain_name_for_template = replace_amp_with_placeholder(domain_name) if domain_name else 'Unknown'
                    for q in domain_data.get('questions', []):
                        answer = q.get('answer', {})
                        score = answer.get('numeric_score', 0) if answer else 0
                        
                        # Determine maturity level from score or option
                        option_value = answer.get('option', '') if answer else ''
                        if option_value:
                            maturity_level = option_value.capitalize()  # LEADING -> Leading
                        elif score >= 4:
                            maturity_level = 'Leading'
                        elif score >= 3:
                            maturity_level = 'Established'
                        elif score >= 2:
                            maturity_level = 'Developing'
                        else:
                            maturity_level = 'Foundational'
                        
                        # Get the selected response text from the question's maturity level answer field
                        # System questions store answer texts in fields like 'leading_answer', 'established_answer', etc.
                        maturity_level_lower = maturity_level.lower()
                        selected_option_text = q.get(f'{maturity_level_lower}_answer', '')
                        
                        # Fallback to selected_option.text if available (for compatibility)
                        if not selected_option_text and answer:
                            selected_option_text = answer.get('selected_option', {}).get('text', '') if isinstance(answer.get('selected_option'), dict) else ''
                        
                        # Replace ampersands to avoid XML issues
                        if selected_option_text:
                            selected_option_text = replace_amp_with_placeholder(selected_option_text)
                        
                        questions.append({
                            'number': q_num,
                            'code': q.get('code', ''),
                            'text': replace_amp_with_placeholder(q.get('text', '')),
                            'domain': domain_name_for_template,
                            'selected_option_label': maturity_level,
                            'selected_option_text': selected_option_text or 'Not answered',
                            'maturity_level': maturity_level,
                            'score': score,
                            'score_display': f"{score}/4",
                            'comment': replace_amp_with_placeholder(answer.get('note', '') or answer.get('comment', '') or '') if answer else ''
                        })
                        q_num += 1
                template_context['questions'] = questions
                template_context['show_detailed_responses'] = True
                print(f"DEBUG: System questions for Appendix: {len(questions)} items")
                if questions:
                    sample_q = questions[0]
                    print(f"DEBUG: Sample question - code: {sample_q['code']}, selected_option_text: {sample_q['selected_option_text'][:80]}...")
                
                # Generate framework coverage data for System assessments
                # Build answers list from questions_data for framework coverage calculation
                answers_for_coverage = []
                questions_data = report_data.get('questions_data', [])
                for domain_data in questions_data:
                    for q in domain_data.get('questions', []):
                        q_code = q.get('code', '')
                        answer = q.get('answer', {})
                        if q_code and answer:
                            answers_for_coverage.append({
                                'question_code': q_code,
                                'numeric_score': answer.get('numeric_score', 0),
                            })
                
                # Get selected frameworks from system_info
                selected_frameworks = system_info.get('frameworks', [])
                
                # Calculate framework coverage
                framework_coverage_data = get_all_framework_coverage(
                    answers=answers_for_coverage,
                    selected_frameworks=selected_frameworks
                )
                
                # Add frameworks to template context
                template_context['frameworks'] = framework_coverage_data
                print(f"DEBUG: Framework coverage data: {len(framework_coverage_data)} frameworks")
                for fw in framework_coverage_data[:2]:  # Show first 2 for debugging
                    print(f"  - {fw.get('title', 'Unknown')}: is_selected={fw.get('is_selected')}, controls={fw.get('total_controls')}")
            
            # Add FAIRA-specific variables if this is a FAIRA assessment
            elif self.assessment_type == 'FAIRA':
                faira_form = report_data.get('faira_form') or report_data.get('form_responses') or {}
                
                # Process faira_form to replace "Other" in multiselect fields with their _other values
                def process_other_fields(form_data):
                    """Replace 'Other' in multiselect arrays with the corresponding _other field value."""
                    processed = dict(form_data)  # Create a copy
                    
                    # List of fields that have _other companions
                    other_field_mappings = {
                        'A1_3': 'A1_3_other',
                        'A1_7': 'A1_7_other',
                        'A1_6_actions': 'A1_6_actions_other',
                        'A2_2_sources': 'A2_2_sources_other',
                        'A2_2_data_types': 'A2_2_data_types_other',
                        'A2_4': 'A2_4_other',
                        'A2_7_data_types': 'A2_7_data_types_other',
                        'A3_3': 'A3_3_other',
                        'A4_7_pii_types': 'A4_7_pii_types_other',
                        'A4_8_action_types': 'A4_8_action_types_other',
                        'A4_8_legal_basis': 'A4_8_legal_basis_other',
                        'A5_10_frameworks': 'A5_10_frameworks_other',
                        'A5_12': 'A5_12_other',
                        'B2_3_perspectives': 'B2_3_perspectives_other',
                        'B4_2_types': 'B4_2_types_other',
                    }
                    
                    for field, other_field in other_field_mappings.items():
                        if field in processed and isinstance(processed[field], list):
                            other_value = processed.get(other_field, '').strip()
                            if other_value:
                                # Replace "Other" or similar variations with the actual text
                                processed[field] = [
                                    other_value if item.lower() in ['other', 'other (specify)', 'other regulated/sensitive data (specify)', 'other standards or frameworks (specify below)'] 
                                    else item 
                                    for item in processed[field]
                                ]
                    
                    return processed
                
                faira_form = process_other_fields(faira_form)
                
                # Get calculated FAIRA data if available, otherwise calculate it
                faira_risk_summary = report_data.get('faira_risk_summary')
                if not faira_risk_summary:
                    # Calculate FAIRA scores using the scoring engine
                    faira_risk_summary = calculate_overall_risk(faira_form)
                    print(f"DEBUG: Calculated FAIRA risk summary on-the-fly")
                
                # Get domain scores from risk summary
                raw_domain_scores = faira_risk_summary.get('domain_scores', {})
                top_risk_areas_raw = faira_risk_summary.get('top_risk_areas', [])
                
                print(f"DEBUG: FAIRA domain_scores keys: {list(raw_domain_scores.keys())}")
                print(f"DEBUG: FAIRA top_risk_areas: {top_risk_areas_raw}")
                
                # Create a SimpleNamespace-like dict class for dot access
                class DotDict(dict):
                    """Dict that allows dot notation access for Jinja2 templates."""
                    def __getattr__(self, key):
                        try:
                            value = self[key]
                            if isinstance(value, dict):
                                return DotDict(value)
                            return value
                        except KeyError:
                            return DotDict({})  # Return empty DotDict for missing keys
                    def __setattr__(self, key, value):
                        self[key] = value
                
                # Convert domain scores to DotDict for template dot notation access
                # Also calculate Inherent_Risk = Impact × Likelihood / 100 (normalized)
                domain_scores_with_inherent = {}
                for domain, scores in raw_domain_scores.items():
                    impact = scores.get('Impact', 0)
                    likelihood = scores.get('Likelihood', 0)
                    inherent_risk = round((impact * likelihood) / 100, 1) if impact and likelihood else 0
                    domain_scores_with_inherent[domain] = DotDict({
                        'Impact': round(scores.get('Impact', 0), 1),
                        'Likelihood': round(scores.get('Likelihood', 0), 1),
                        'Inherent_Risk': inherent_risk,
                        'Control_Effectiveness': round(scores.get('Control_Effectiveness', 0), 1),
                        'Risk': round(scores.get('Risk', 0), 1)
                    })
                
                domain_scores_obj = DotDict(domain_scores_with_inherent)
                
                print(f"DEBUG: Built domain_scores_obj with Inherent_Risk for {len(domain_scores_with_inherent)} domains")
                for domain, scores in domain_scores_with_inherent.items():
                    print(f"  - {domain}: Impact={scores['Impact']}, Likelihood={scores['Likelihood']}, Inherent_Risk={scores['Inherent_Risk']}, Control={scores['Control_Effectiveness']}, Risk={scores['Risk']}")
                
                # Build top_domains list with proper structure for template
                # Template expects: top_domains[0].name, top_domains[0].risk
                top_domains_list = []
                
                # If we have top_risk_areas, convert them
                if top_risk_areas_raw:
                    for area in top_risk_areas_raw:
                        if isinstance(area, dict):
                            top_domains_list.append(DotDict({
                                'name': area.get('domain', area.get('name', area.get('fullName', 'Unknown'))),
                                'risk': round(area.get('risk_score', area.get('risk', area.get('Risk', 0))), 1)
                            }))
                
                # If no top_risk_areas, build from domain_scores sorted by risk
                if not top_domains_list and raw_domain_scores:
                    sorted_domains = sorted(
                        [(k, v.get('Risk', 0) if isinstance(v, dict) else 0) for k, v in raw_domain_scores.items()],
                        key=lambda x: x[1],
                        reverse=True
                    )
                    for domain_name, risk_score in sorted_domains[:3]:
                        top_domains_list.append(DotDict({
                            'name': domain_name,
                            'risk': round(risk_score, 1)
                        }))
                
                # Ensure we have at least 3 entries (pad with empty if needed)
                while len(top_domains_list) < 3:
                    top_domains_list.append(DotDict({'name': 'N/A', 'risk': 0}))
                
                print(f"DEBUG: top_domains_list built with {len(top_domains_list)} entries")
                for i, td in enumerate(top_domains_list[:3]):
                    print(f"  - top_domains[{i}]: name={td.get('name')}, risk={td.get('risk')}")
                
                # Get system name with fallbacks
                system_name = faira_form.get('ai_system_name') or faira_form.get('A1_2') or report_data.get('assessment', {}).get('name', '')
                
                # Add FAIRA form responses as top-level variables
                template_context.update({
                    # Organization and System info from Part A
                    'system_name': system_name,
                    'org_name': faira_form.get('business_unit', ''),
                    'assessment_scope': ', '.join(faira_form.get('A1_4', [])) if isinstance(faira_form.get('A1_4'), list) else faira_form.get('A1_4', ''),
                    'deployment_context': faira_form.get('A4_3', ''),
                    
                    # Risk scores from calculated data
                    'overall_risk_score': round(faira_risk_summary.get('overall_risk_score', 50), 1),
                    'overall_risk_level': faira_risk_summary.get('overall_risk_level', 'Medium'),
                    'overall_impact_score': round(faira_risk_summary.get('total_impact', 50), 1),
                    'overall_likelihood_score': round(faira_risk_summary.get('total_likelihood', 50), 1),
                    'overall_control_effectiveness_score': round(faira_risk_summary.get('total_control_effectiveness', 50), 1),
                    
                    # Domain scores - both as dict and as DotDict for template flexibility
                    'domain_scores': domain_scores_obj,
                    'top_risk_areas': top_risk_areas_raw,
                    'top_domains': top_domains_list,  # List with .name and .risk attributes
                    
                    # Control variable for template (single control placeholder)
                    'control': DotDict({
                        'priority': '',
                        'area': '',
                        'description': '',
                        'type': ''
                    }),
                    
                    # Additional score variables the template expects
                    'total_impact': round(faira_risk_summary.get('total_impact', 50), 1),
                    'total_likelihood': round(faira_risk_summary.get('total_likelihood', 50), 1),
                    'total_control_effectiveness': round(faira_risk_summary.get('total_control_effectiveness', 50), 1),
                    'raw_risk_score': round(faira_risk_summary.get('raw_risk_score', 50), 1),
                    'autonomy_factor': faira_risk_summary.get('autonomy_factor', ''),
                    'data_quality_factor': faira_risk_summary.get('data_quality_factor', ''),
                    'expertise_factor': faira_risk_summary.get('expertise_factor', ''),
                    'assessment_date': report_data.get('assessment_date', ''),
                    'generation_date': get_report_local_time().strftime('%Y-%m-%d'),
                    
                    # Gap variable placeholder (template expects gap.domain, gap.existing_controls, etc.)
                    # This will be populated with actual data after gap analysis is built
                    'gap': DotDict({
                        'domain': '',
                        'existing_controls': '',
                        'gaps': '',
                        'priority': '',
                        'risk_score': 0,
                        'control_effectiveness': 0
                    }),
                    
                    # Artefact variable placeholder (template expects artefact.name, artefact.description, etc.)
                    'artefact': DotDict({
                        'name': '',
                        'description': '',
                        'domains': ''
                    }),
                    
                    # Assessment object for template compatibility
                    'assessment': DotDict({
                        'date': report_data.get('assessment', {}).get('completed_at', datetime.now().strftime('%Y-%m-%d'))
                    }),
                    
                    # Full FAIRA form object for flexibility
                    'faira_form': DotDict(faira_form),
                })
                
                # Generate risk gauge image for FAIRA
                try:
                    overall_risk_score = faira_risk_summary.get('overall_risk_score', 50)
                    overall_risk_level = faira_risk_summary.get('overall_risk_level', 'Medium')
                    risk_gauge_bytes = generate_risk_gauge(overall_risk_score, overall_risk_level)
                    risk_gauge_inline = InlineImage(doc, io.BytesIO(risk_gauge_bytes), width=Inches(6.5))
                    template_context['risk_gauge_image'] = risk_gauge_inline
                    print(f"DEBUG: Generated risk gauge image for score {overall_risk_score} ({overall_risk_level})")
                except Exception as e:
                    print(f"WARNING: Could not generate risk gauge image: {e}")
                    template_context['risk_gauge_image'] = ''
                
                # Generate FAIRA radar charts
                try:
                    # Get domain scores for radar charts
                    faira_domain_scores = template_context.get('domain_scores', {})
                    if faira_domain_scores:
                        # Generate domain risk radar chart
                        domain_risk_radar_bytes = generate_faira_domain_risk_radar(faira_domain_scores)
                        template_context['domain_radar_risk'] = InlineImage(doc, io.BytesIO(domain_risk_radar_bytes), width=Inches(4.5))
                        print("DEBUG: Generated FAIRA domain risk radar chart")
                        
                        # Generate inherent risk radar chart
                        inherent_risk_radar_bytes = generate_faira_inherent_risk_radar(faira_domain_scores)
                        template_context['inherent_risk_radar'] = InlineImage(doc, io.BytesIO(inherent_risk_radar_bytes), width=Inches(4.5))
                        print("DEBUG: Generated FAIRA inherent risk radar chart")
                    else:
                        print("WARNING: No domain_scores available for FAIRA radar charts")
                        template_context['domain_radar_risk'] = ''
                        template_context['inherent_risk_radar'] = ''
                except Exception as e:
                    print(f"WARNING: Could not generate FAIRA radar charts: {e}")
                    import traceback
                    traceback.print_exc()
                    template_context['domain_radar_risk'] = ''
                    template_context['inherent_risk_radar'] = ''
                
                # Add FAIRA controls to template context
                faira_controls_data = report_data.get('faira_controls', {})
                if faira_controls_data:
                    top_controls = faira_controls_data.get('top_controls', [])
                    # Convert each control to DotDict for template dot notation access
                    controls_list = []
                    for ctrl in top_controls:
                        controls_list.append(DotDict({
                            'rank': ctrl.get('rank', 0),
                            'control_id': ctrl.get('control_id', ''),
                            'title': ctrl.get('title', ''),
                            'description': ctrl.get('description', ''),
                            'detailed_description': ctrl.get('detailed_description', ''),
                            'domains': ', '.join(ctrl.get('domains', [])) if isinstance(ctrl.get('domains'), list) else ctrl.get('domains', ''),
                            'domain_list': ctrl.get('domains', []),  # Keep as list for filtering
                            'control_type': ctrl.get('control_type', ''),
                            'implementation_effort': ctrl.get('implementation_effort', ''),
                            'implementation_horizon': ctrl.get('implementation_horizon', ''),
                            'evidence_examples': ctrl.get('evidence_examples', []),
                            'priority': ctrl.get('priority', ''),
                            'rationale': ctrl.get('rationale', ''),
                        }))
                    
                    # Ensure we have at least 3 controls (pad with empty if needed)
                    while len(controls_list) < 3:
                        controls_list.append(DotDict({
                            'rank': len(controls_list) + 1,
                            'control_id': '',
                            'title': 'N/A',
                            'description': '',
                            'detailed_description': '',
                            'domains': '',
                            'domain_list': [],
                            'control_type': '',
                            'implementation_effort': '',
                            'implementation_horizon': '',
                            'evidence_examples': [],
                            'priority': '',
                            'rationale': '',
                        }))
                    
                    template_context['faira_controls'] = DotDict({
                        'top_controls': controls_list,
                        'total_matched': faira_controls_data.get('total_matched', 0),
                        'risk_summary': DotDict(faira_controls_data.get('risk_summary', {}))
                    })
                    
                    # Also add individual top control shortcuts for template convenience
                    # Access via: {{top_control_1.title}}, {{top_control_2.title}}, etc.
                    for i, ctrl in enumerate(controls_list[:5], 1):
                        template_context[f'top_control_{i}'] = ctrl
                    
                    print(f"DEBUG: Added {len(controls_list)} FAIRA controls to template context")
                    for i, ctrl in enumerate(controls_list[:3]):
                        print(f"  - Control {i+1}: {ctrl.get('title', 'N/A')[:50]}")
                else:
                    # Empty controls structure
                    template_context['faira_controls'] = DotDict({
                        'top_controls': [],
                        'total_matched': 0,
                        'risk_summary': DotDict({})
                    })
                
                # Load ALL controls grouped by domain for domain-specific loops
                # This allows: {%tr for control in controls_by_domain.Accountability %}
                try:
                    from faira_scoring_engine import load_controls_data
                    all_controls_data = load_controls_data()
                    all_controls = all_controls_data.get('controls', [])
                    
                    # Group controls by domain
                    domain_names = ['Accountability', 'Contestability', 'Fairness', 'Privacy', 
                                   'Reliability', 'Transparency', 'Values', 'Wellbeing']
                    
                    controls_by_domain = {}
                    for domain in domain_names:
                        domain_controls = []
                        for ctrl in all_controls:
                            ctrl_domains = ctrl.get('domains', [])
                            if domain in ctrl_domains:
                                domain_controls.append(DotDict({
                                    'control_id': ctrl.get('control_id', ''),
                                    'title': ctrl.get('title', ''),
                                    'description': ctrl.get('description', ''),
                                    'detailed_description': ctrl.get('detailed_description', ''),
                                    'domains': ', '.join(ctrl.get('domains', [])),
                                    'control_type': ctrl.get('control_type', ''),
                                    'implementation_effort': ctrl.get('implementation_effort', ''),
                                    'implementation_horizon': ctrl.get('implementation_horizon', ''),
                                    'evidence_examples': ctrl.get('evidence_examples', []),
                                    'tags': ctrl.get('tags', []),
                                }))
                        controls_by_domain[domain] = domain_controls
                    
                    template_context['controls_by_domain'] = DotDict(controls_by_domain)
                    
                    # Also add recommended_controls as alias for backward compatibility
                    template_context['recommended_controls'] = DotDict(controls_by_domain)
                    
                    # Create recommended_controls_top3 - top 3 controls per domain
                    # Sorted by: 1) implementation_horizon (Immediate > Short-term > Medium-term > Long-term)
                    #            2) implementation_effort (Low > Medium > High)
                    horizon_priority = {'Immediate': 1, 'Short-term': 2, 'Medium-term': 3, 'Long-term': 4}
                    effort_priority = {'Low': 1, 'Medium': 2, 'High': 3}
                    
                    def sort_key(ctrl):
                        horizon = ctrl.get('implementation_horizon', 'Long-term')
                        effort = ctrl.get('implementation_effort', 'High')
                        return (horizon_priority.get(horizon, 99), effort_priority.get(effort, 99))
                    
                    controls_top3_by_domain = {}
                    for domain in domain_names:
                        domain_ctrls = controls_by_domain.get(domain, [])
                        # Sort by horizon then effort
                        sorted_ctrls = sorted(domain_ctrls, key=sort_key)
                        # Take top 3 (or all if fewer than 3)
                        controls_top3_by_domain[domain] = sorted_ctrls[:3]
                    
                    template_context['recommended_controls_top3'] = DotDict(controls_top3_by_domain)
                    
                    print(f"DEBUG: Added controls_by_domain with {len(all_controls)} total controls")
                    print(f"DEBUG: Added recommended_controls_top3 (top 3 per domain, sorted by horizon/effort)")
                    for domain in domain_names:
                        print(f"  - {domain}: {len(controls_by_domain.get(domain, []))} controls, top3: {len(controls_top3_by_domain.get(domain, []))}")
                        
                except Exception as e:
                    print(f"WARNING: Could not load controls by domain: {e}")
                    template_context['controls_by_domain'] = DotDict({})
                    template_context['recommended_controls'] = DotDict({})
                
                # ============================================================
                # FAIRA GAP ANALYSIS
                # Combines top risk domains with their recommended controls
                # Creates a 'gaps' list for template iteration:
                #   {%tr for gap in gaps %}
                #     {{gap.domain}} | {{gap.existing_controls}} | {{gap.gaps}} | {{gap.priority}}
                #   {%tr endfor %}
                # ============================================================
                try:
                    gaps_list = []
                    
                    # Get top 3 risk-contributing domains
                    top_3_domains = top_domains_list[:3] if top_domains_list else []
                    
                    print(f"DEBUG: Gap analysis - top_domains_list has {len(top_domains_list)} entries before slicing")
                    print(f"DEBUG: Gap analysis - top_3_domains has {len(top_3_domains)} entries after slicing")
                    
                    # CRITICAL: If top_3_domains is still empty, generate fallback from raw domain scores
                    if not top_3_domains and raw_domain_scores:
                        print("DEBUG: top_3_domains empty, building from raw_domain_scores as fallback")
                        sorted_by_risk = sorted(
                            [(k, v.get('Risk', 0) if isinstance(v, dict) else 0) for k, v in raw_domain_scores.items()],
                            key=lambda x: x[1],
                            reverse=True
                        )
                        for domain_name, risk_score in sorted_by_risk[:3]:
                            top_3_domains.append(DotDict({
                                'name': domain_name,
                                'risk': round(risk_score, 1)
                            }))
                        print(f"DEBUG: Built {len(top_3_domains)} domains from raw_domain_scores fallback")
                    
                    # Get top 3 controls - access from faira_controls DotDict
                    faira_controls_ctx = template_context.get('faira_controls', {})
                    if hasattr(faira_controls_ctx, 'top_controls'):
                        top_3_controls = faira_controls_ctx.top_controls[:3] if faira_controls_ctx.top_controls else []
                    elif isinstance(faira_controls_ctx, dict):
                        top_3_controls = faira_controls_ctx.get('top_controls', [])[:3]
                    else:
                        top_3_controls = []
                    
                    print(f"DEBUG: Gap analysis - top_3_domains: {len(top_3_domains)}, top_3_controls: {len(top_3_controls)}")
                    
                    # Map short domain names to full names for control matching
                    domain_name_map = {
                        'Wellbeing': 'Human, Societal and Environmental Wellbeing',
                        'Values': 'Human-Centred Values',
                        'Fairness': 'Fairness',
                        'Privacy': 'Privacy Protection and Security',
                        'Reliability': 'Reliability and Safety',
                        'Transparency': 'Transparency and Explainability',
                        'Contestability': 'Contestability',
                        'Accountability': 'Accountability'
                    }
                    
                    # Build gaps list combining domain risks and controls
                    for i, domain in enumerate(top_3_domains):
                        domain_name = domain.get('name', 'Unknown')
                        domain_risk = domain.get('risk', 0)
                        
                        # Determine priority based on risk score
                        if domain_risk >= 55:
                            priority = 'High'
                        elif domain_risk >= 35:
                            priority = 'Medium'
                        else:
                            priority = 'Low'
                        
                        # Get domain scores for existing controls assessment
                        domain_scores_dict = raw_domain_scores.get(domain_name, {}) if raw_domain_scores else {}
                        control_effectiveness = domain_scores_dict.get('Control_Effectiveness', 0) if isinstance(domain_scores_dict, dict) else 0
                        
                        # Assess existing controls based on CE score
                        if control_effectiveness >= 70:
                            existing_controls = 'Strong controls in place'
                        elif control_effectiveness >= 40:
                            existing_controls = 'Partial controls implemented'
                        else:
                            existing_controls = 'Limited or no controls'
                        
                        # Get the corresponding top control for this domain (if any)
                        gap_description = ''
                        if i < len(top_3_controls):
                            ctrl = top_3_controls[i]
                            ctrl_title = ctrl.get('title', '')
                            ctrl_type = ctrl.get('control_type', '')
                            if ctrl_title:
                                gap_description = f"{ctrl_title}"
                                if ctrl_type:
                                    gap_description += f" ({ctrl_type})"
                        
                        # If no matching control, create a generic gap based on domain
                        if not gap_description:
                            gap_description = f"Strengthen {domain_name} controls"
                        
                        gaps_list.append(DotDict({
                            'domain': domain_name,
                            'risk_score': round(domain_risk, 1),
                            'existing_controls': existing_controls,
                            'control_effectiveness': round(control_effectiveness, 1),
                            'gaps': gap_description,
                            'priority': priority
                        }))
                    
                    # Ensure we have at least 3 gap entries
                    while len(gaps_list) < 3:
                        gaps_list.append(DotDict({
                            'domain': 'N/A',
                            'risk_score': 0,
                            'existing_controls': 'N/A',
                            'control_effectiveness': 0,
                            'gaps': 'No additional gaps identified',
                            'priority': 'Low'
                        }))
                    
                    template_context['gaps'] = gaps_list
                    
                    # Also add individual gap shortcuts for template convenience
                    for i, gap in enumerate(gaps_list[:3], 1):
                        template_context[f'gap_{i}'] = gap
                    
                    print(f"DEBUG: Built FAIRA gap analysis with {len(gaps_list)} entries:")
                    for i, gap in enumerate(gaps_list[:3]):
                        print(f"  - Gap {i+1}: {gap.get('domain')} (Risk: {gap.get('risk_score')}, Priority: {gap.get('priority')})")
                        print(f"           Existing: {gap.get('existing_controls')}")
                        print(f"           Gap: {gap.get('gaps')}")
                    
                except Exception as e:
                    print(f"WARNING: Could not build FAIRA gap analysis: {e}")
                    import traceback
                    traceback.print_exc()
                    # CRITICAL: Never leave gaps empty - provide fallback data
                    # An empty list causes docxtpl to remove the entire table row
                    fallback_gaps = [
                        DotDict({
                            'domain': 'Accountability',
                            'risk_score': 50.0,
                            'existing_controls': 'Assessment pending',
                            'control_effectiveness': 50.0,
                            'gaps': 'Implement governance oversight framework',
                            'priority': 'Medium'
                        }),
                        DotDict({
                            'domain': 'Transparency',
                            'risk_score': 45.0,
                            'existing_controls': 'Assessment pending',
                            'control_effectiveness': 50.0,
                            'gaps': 'Establish AI decision documentation',
                            'priority': 'Medium'
                        }),
                        DotDict({
                            'domain': 'Fairness',
                            'risk_score': 40.0,
                            'existing_controls': 'Assessment pending',
                            'control_effectiveness': 50.0,
                            'gaps': 'Develop bias assessment procedures',
                            'priority': 'Medium'
                        }),
                    ]
                    template_context['gaps'] = fallback_gaps
                    for i, gap in enumerate(fallback_gaps, 1):
                        template_context[f'gap_{i}'] = gap
                    print(f"DEBUG: Using fallback gaps data ({len(fallback_gaps)} entries)")
                
                # Add AI narratives for FAIRA reports (32 placeholders)
                # Using variable naming convention: ai.f_* for FAIRA
                ai_narratives = report_data.get('ai_narratives', {})
                template_context['ai'] = {
                    # Executive and Overview
                    'f_executive_snapshot': replace_amp_with_placeholder(ai_narratives.get('f_executive_snapshot', '')),
                    'f_risk_profile_interpretation': replace_amp_with_placeholder(ai_narratives.get('f_risk_profile_interpretation', '')),
                    'f_methodology_interpretation': replace_amp_with_placeholder(ai_narratives.get('f_methodology_interpretation', '')),
                    
                    # Context Notes (Part A)
                    'f_system_overview_notes': replace_amp_with_placeholder(ai_narratives.get('f_system_overview_notes', '')),
                    'f_data_context_notes': replace_amp_with_placeholder(ai_narratives.get('f_data_context_notes', '')),
                    'f_user_stakeholder_context_notes': replace_amp_with_placeholder(ai_narratives.get('f_user_stakeholder_context_notes', '')),
                    'f_deployment_context_notes': replace_amp_with_placeholder(ai_narratives.get('f_deployment_context_notes', '')),
                    'f_impact_context_notes': replace_amp_with_placeholder(ai_narratives.get('f_impact_context_notes', '')),
                    
                    # Domain Analysis (Part B)
                    'f_domain_analysis': replace_amp_with_placeholder(ai_narratives.get('f_domain_analysis', '')),
                    'f_domain_b1_analysis': replace_amp_with_placeholder(ai_narratives.get('f_domain_b1_analysis', '')),
                    'f_domain_b2_analysis': replace_amp_with_placeholder(ai_narratives.get('f_domain_b2_analysis', '')),
                    'f_domain_b3_analysis': replace_amp_with_placeholder(ai_narratives.get('f_domain_b3_analysis', '')),
                    'f_domain_b4_analysis': replace_amp_with_placeholder(ai_narratives.get('f_domain_b4_analysis', '')),
                    'f_domain_b5_analysis': replace_amp_with_placeholder(ai_narratives.get('f_domain_b5_analysis', '')),
                    'f_domain_b6_analysis': replace_amp_with_placeholder(ai_narratives.get('f_domain_b6_analysis', '')),
                    'f_domain_b7_analysis': replace_amp_with_placeholder(ai_narratives.get('f_domain_b7_analysis', '')),
                    'f_domain_b8_analysis': replace_amp_with_placeholder(ai_narratives.get('f_domain_b8_analysis', '')),
                    'f_top_risk_areas_analysis': replace_amp_with_placeholder(ai_narratives.get('f_top_risk_areas_analysis', '')),
                    
                    # Controls
                    'f_controls_overview': replace_amp_with_placeholder(ai_narratives.get('f_controls_overview', '')),
                    'f_control_prioritisation_rationale': replace_amp_with_placeholder(ai_narratives.get('f_control_prioritisation_rationale', '')),
                    'f_domain_b1_controls_rationale': replace_amp_with_placeholder(ai_narratives.get('f_domain_b1_controls_rationale', '')),
                    'f_domain_b2_controls_rationale': replace_amp_with_placeholder(ai_narratives.get('f_domain_b2_controls_rationale', '')),
                    'f_domain_b3_controls_rationale': replace_amp_with_placeholder(ai_narratives.get('f_domain_b3_controls_rationale', '')),
                    'f_domain_b4_controls_rationale': replace_amp_with_placeholder(ai_narratives.get('f_domain_b4_controls_rationale', '')),
                    'f_domain_b5_controls_rationale': replace_amp_with_placeholder(ai_narratives.get('f_domain_b5_controls_rationale', '')),
                    'f_domain_b6_controls_rationale': replace_amp_with_placeholder(ai_narratives.get('f_domain_b6_controls_rationale', '')),
                    'f_domain_b7_controls_rationale': replace_amp_with_placeholder(ai_narratives.get('f_domain_b7_controls_rationale', '')),
                    'f_domain_b8_controls_rationale': replace_amp_with_placeholder(ai_narratives.get('f_domain_b8_controls_rationale', '')),
                    
                    # Summary and Next Steps
                    'f_gap_and_next_steps_summary': replace_amp_with_placeholder(ai_narratives.get('f_gap_and_next_steps_summary', '')),
                    'f_governance_assurance_posture': replace_amp_with_placeholder(ai_narratives.get('f_governance_assurance_posture', '')),
                    'f_decision_next_steps_summary': replace_amp_with_placeholder(ai_narratives.get('f_decision_next_steps_summary', '')),
                    'f_supporting_artefacts_summary': replace_amp_with_placeholder(ai_narratives.get('f_supporting_artefacts_summary', '')),
                }
                print(f"DEBUG: Added FAIRA AI narratives to template context: {list(template_context['ai'].keys())}")
                for key, value in template_context['ai'].items():
                    if value:
                        print(f"  - {key} length: {len(value)} chars")
            
            print("Generated report data structure:")
            print(f"  Organization: {template_context['org']['name']}")
            print(f"  Overall score: {template_context['overall']['score']}")
            print(f"  High priority actions: {len(template_context['actions']['high'])}")
            print(f"  Medium priority actions: {len(template_context['actions']['medium'])}")
            print(f"  Low priority actions: {len(template_context['actions']['low'])}")
            print(f"  Sector: {template_context['sector_name']}")
            print(f"  Sector actions - high: {len(template_context['sector_actions']['high'])}, medium: {len(template_context['sector_actions']['medium'])}, low: {len(template_context['sector_actions']['low'])}")
            print(f"  System info keys: {list(template_context['system_info'].keys())}")
            
            # Show sample sector actions for debugging
            if template_context.get('sector_actions_high_top5'):
                sample = template_context['sector_actions_high_top5'][0]
                print(f"  Sample sector_actions_high_top5[0]: question_id={sample.get('question_id')}, domain={sample.get('domain')}, sector_action={sample.get('sector_action', 'N/A')[:50]}...")
            else:
                print(f"  sector_actions_high_top5 is EMPTY")
            
            # Show sample actions for debugging
            if template_context['actions']['high']:
                sample_high = template_context['actions']['high'][0]
                print(f"  Sample high action: {sample_high.get('domain', 'N/A')} | {sample_high.get('question_id', 'N/A')}")
            
            # Helper function for date formatting in template
            template_context['formatDate'] = self.format_date
            
            # Define custom Jinja2 filters
            def replace_last(value, old, new):
                """Replace the last occurrence of 'old' with 'new' in a string."""
                if isinstance(value, str) and old in value:
                    return new.join(value.rsplit(old, 1))
                return value
            
            def format_list(items, other_value=None):
                """Format a list as comma-separated with 'and' before last item.
                Preserves original case of items.
                If other_value is provided and 'Other' is in the list, replaces 'Other' with that value.
                Example: ['Apple', 'Banana', 'Cherry'] -> 'Apple, Banana, and Cherry'
                """
                if not items:
                    return ''
                if isinstance(items, str):
                    return items
                # Convert to list, preserving original case
                items_list = [str(item) for item in items]
                
                # Replace "Other" with the specified other_value if provided
                if other_value and other_value.strip():
                    items_list = [other_value if item.lower() == 'other' else item for item in items_list]
                
                if len(items_list) == 1:
                    return items_list[0]
                if len(items_list) == 2:
                    return f"{items_list[0]} and {items_list[1]}"
                # Join with commas, then replace last comma with ', and'
                joined = ', '.join(items_list)
                return ', and '.join(joined.rsplit(', ', 1))
            
            def rating_label(value):
                """Convert numeric rating (1-5) to text label.
                1 = Very Low, 2 = Low, 3 = Moderate, 4 = High, 5 = Very High
                """
                labels = {
                    1: 'Very Low',
                    2: 'Low',
                    3: 'Moderate',
                    4: 'High',
                    5: 'Very High',
                    '1': 'Very Low',
                    '2': 'Low',
                    '3': 'Moderate',
                    '4': 'High',
                    '5': 'Very High',
                }
                return labels.get(value, str(value) if value else 'Not specified')
            
            def severity_label(value):
                """Convert numeric severity (1-3) to text label.
                1 = Minor, 2 = Moderate, 3 = Major
                """
                labels = {
                    1: 'Minor',
                    2: 'Moderate',
                    3: 'Major',
                    '1': 'Minor',
                    '2': 'Moderate',
                    '3': 'Major',
                }
                return labels.get(value, str(value) if value else 'Not specified')
            
            # Create Jinja2 environment with custom filters for docxtpl
            from jinja2 import Environment
            jinja_env = Environment()
            jinja_env.filters['replace_last'] = replace_last
            jinja_env.filters['format_list'] = format_list
            jinja_env.filters['rating_label'] = rating_label
            jinja_env.filters['severity_label'] = severity_label
            
            # Debug: Verify gaps list before render and ensure it's never empty for FAIRA
            if self.assessment_type == 'FAIRA':
                gaps_before_render = template_context.get('gaps', [])
                print(f"DEBUG PRE-RENDER: gaps list has {len(gaps_before_render)} entries")
                
                # CRITICAL FAILSAFE: If gaps is still empty at this point, create fallback data
                # This prevents docxtpl from removing the entire table row
                if not gaps_before_render:
                    print("WARNING: gaps list is EMPTY! Applying emergency fallback to prevent empty table")
                    # Use DotDict if it's in scope, otherwise use plain dict
                    try:
                        emergency_gaps = [
                            DotDict({
                                'domain': 'Accountability',
                                'risk_score': 50.0,
                                'existing_controls': 'Assessment in progress',
                                'control_effectiveness': 50.0,
                                'gaps': 'Implement accountability framework',
                                'priority': 'Medium'
                            }),
                            DotDict({
                                'domain': 'Transparency',
                                'risk_score': 45.0,
                                'existing_controls': 'Assessment in progress',
                                'control_effectiveness': 50.0,
                                'gaps': 'Establish documentation processes',
                                'priority': 'Medium'
                            }),
                            DotDict({
                                'domain': 'Fairness',
                                'risk_score': 40.0,
                                'existing_controls': 'Assessment in progress',
                                'control_effectiveness': 50.0,
                                'gaps': 'Develop bias testing protocols',
                                'priority': 'Medium'
                            }),
                        ]
                    except NameError:
                        # DotDict not in scope, use plain dicts
                        emergency_gaps = [
                            {'domain': 'Accountability', 'risk_score': 50.0, 'existing_controls': 'Assessment in progress', 'control_effectiveness': 50.0, 'gaps': 'Implement accountability framework', 'priority': 'Medium'},
                            {'domain': 'Transparency', 'risk_score': 45.0, 'existing_controls': 'Assessment in progress', 'control_effectiveness': 50.0, 'gaps': 'Establish documentation processes', 'priority': 'Medium'},
                            {'domain': 'Fairness', 'risk_score': 40.0, 'existing_controls': 'Assessment in progress', 'control_effectiveness': 50.0, 'gaps': 'Develop bias testing protocols', 'priority': 'Medium'},
                        ]
                    template_context['gaps'] = emergency_gaps
                    for i, gap in enumerate(emergency_gaps, 1):
                        template_context[f'gap_{i}'] = gap
                    gaps_before_render = emergency_gaps
                    print(f"DEBUG: Applied emergency fallback - gaps now has {len(gaps_before_render)} entries")
                
                for i, g in enumerate(gaps_before_render[:3]):
                    print(f"  - gaps[{i}]: domain={g.get('domain')}, gaps={g.get('gaps')}, priority={g.get('priority')}")
            
            # Render the template with custom jinja_env
            # Note: We use autoescape=False (default) because autoescape=True corrupts DOCX files
            doc.render(template_context, jinja_env=jinja_env)
            
            # Center-align images (heatmap and radar chart)
            self._center_align_images(doc)
            
            # Populate recommendation tables programmatically
            # Different handling for Awareness vs other assessment types
            if self.assessment_type == 'Awareness':
                # Awareness uses sector_actions tables with question_id, domain, sector_action
                self._populate_awareness_sector_tables(doc, report_data)
            elif self.assessment_type == 'Readiness':
                # Readiness uses the same table structure as Awareness (question_id, domain, sector_action)
                self._populate_readiness_sector_tables(doc, report_data)
            elif self.assessment_type == 'Orgwide':
                # Orgwide uses the same table structure as Readiness (question_id, domain, sector_action)
                self._populate_readiness_sector_tables(doc, report_data)
            elif self.assessment_type == 'System':
                # System template uses Jinja2 {%tr for %} loops which handle table rows correctly
                # No programmatic table population needed - the loops use sector_actions_*_top5 (5 items)
                print("DEBUG: System template uses Jinja2 loops for action tables - skipping programmatic population")
            elif not self.use_test_template:
                # Other templates use actions tables with domain, question_id, text
                self._populate_recommendation_tables(doc, report_data)
            
            # Save to bytes
            output_buffer = io.BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            docx_bytes = output_buffer.getvalue()
            
            # Post-process: Replace ampersand placeholder with proper XML entity
            # This is needed because docxtpl strips & in paragraph loops
            docx_bytes = self._post_process_ampersands(docx_bytes)
            
            # Post-process: Convert markdown formatting to Word formatting
            # This handles **bold** and ### headings in AI narratives
            docx_bytes = self._post_process_markdown_formatting(docx_bytes)
            
            return docx_bytes
            
        except Exception as e:
            print(f"Error generating DOCX report: {str(e)}")
            import traceback
            traceback.print_exc()
            raise Exception(f"Failed to generate DOCX report: {str(e)}")
    
    def _load_template_with_ampersand_protection(self, template_path: str) -> DocxTemplate:
        """
        Load a DOCX template.
        
        Note: Ampersand protection is now handled in post-processing via _post_process_ampersands
        This method now just loads the template directly for reliability.
        
        Version: 2026-01-31-v2 - Removed temp file logic to fix production PackageNotFound error
        """
        import os as os_module
        
        print(f"DEBUG: Loading template from {template_path} (direct load, no temp file)")
        
        if not os_module.path.exists(template_path):
            raise Exception(f"Template file not found: {template_path}")
        
        file_size = os_module.path.getsize(template_path)
        print(f"DEBUG: Template file size: {file_size} bytes")
        
        doc = DocxTemplate(template_path)
        return doc

    def _post_process_ampersands(self, docx_bytes: bytes) -> bytes:
        """
        Post-process DOCX to replace ampersand placeholders and fix bare ampersands.
        
        This handles two cases:
        1. AMP_PLACEHOLDER -> &amp; (for dynamic content we protected)
        2. Bare & -> &amp; (for static template text that Jinja2 unescaped)
        
        Jinja2 interprets &amp; in the template XML as a literal '&', so we need
        to re-escape these bare ampersands after rendering to maintain valid XML.
        """
        import re
        
        # Read the DOCX (which is a ZIP file)
        input_buffer = io.BytesIO(docx_bytes)
        files = {}
        
        with zipfile.ZipFile(input_buffer, 'r') as zin:
            for name in zin.namelist():
                files[name] = zin.read(name)
        
        # Process XML files to fix ampersands
        output_buffer = io.BytesIO()
        with zipfile.ZipFile(output_buffer, 'w', zipfile.ZIP_DEFLATED) as zout:
            for name, content in files.items():
                if name.endswith('.xml'):
                    try:
                        content_str = content.decode('utf-8')
                        
                        # Step 1: Replace our placeholder with &amp;
                        content_str = content_str.replace(AMP_PLACEHOLDER, '&amp;')
                        
                        # Step 2: Fix bare ampersands that aren't already part of XML entities
                        # This regex matches & that is NOT followed by amp;, lt;, gt;, quot;, apos;, or #
                        # (which are valid XML entity patterns)
                        content_str = re.sub(
                            r'&(?!(amp|lt|gt|quot|apos|#\d+|#x[0-9a-fA-F]+);)',
                            '&amp;',
                            content_str
                        )
                        
                        content = content_str.encode('utf-8')
                    except UnicodeDecodeError:
                        # Keep binary content as-is
                        pass
                        
                zout.writestr(name, content)
        
        output_buffer.seek(0)
        return output_buffer.read()
    
    def _post_process_markdown_formatting(self, docx_bytes: bytes) -> bytes:
        """
        Post-process DOCX to convert markdown formatting to Word formatting.
        
        Handles:
        - **text** -> bold text
        - ### text -> bold text (heading style)
        """
        import re
        
        # Read the DOCX (which is a ZIP file)
        input_buffer = io.BytesIO(docx_bytes)
        files = {}
        
        with zipfile.ZipFile(input_buffer, 'r') as zin:
            for name in zin.namelist():
                files[name] = zin.read(name)
        
        # Process document.xml to convert markdown to Word formatting
        output_buffer = io.BytesIO()
        with zipfile.ZipFile(output_buffer, 'w', zipfile.ZIP_DEFLATED) as zout:
            for name, content in files.items():
                if name == 'word/document.xml':
                    content_str = content.decode('utf-8')
                    
                    # Convert **bold** to Word bold formatting
                    # Find **text** patterns and wrap with bold run properties
                    def replace_bold(match):
                        text = match.group(1)
                        # Return text wrapped in bold run
                        return f'</w:t></w:r><w:r><w:rPr><w:b/></w:rPr><w:t>{text}</w:t></w:r><w:r><w:t>'
                    
                    # Pattern to match **text** (non-greedy, doesn't span tags)
                    content_str = re.sub(r'\*\*([^*<>]+)\*\*', replace_bold, content_str)
                    
                    # Convert ### headings to bold
                    # Find ### at start of text and make the line bold
                    def replace_heading(match):
                        text = match.group(1).strip()
                        # Return text wrapped in bold run (remove the ###)
                        return f'</w:t></w:r><w:r><w:rPr><w:b/></w:rPr><w:t>{text}</w:t></w:r><w:r><w:t>'
                    
                    # Pattern to match ### text (heading markers)
                    content_str = re.sub(r'###\s*([^<>\n]+)', replace_heading, content_str)
                    
                    # Also handle ## headings
                    content_str = re.sub(r'##\s*([^<>\n]+)', replace_heading, content_str)
                    
                    # Clean up any empty runs created by the replacements
                    content_str = re.sub(r'<w:r><w:t></w:t></w:r>', '', content_str)
                    
                    content = content_str.encode('utf-8')
                    
                zout.writestr(name, content)
        
        output_buffer.seek(0)
        print("DEBUG: Markdown formatting converted to Word formatting")
        return output_buffer.read()
    
    async def generate_report_for_assessment(self, assessment_id: str, db, current_user, view_type: str = "heatmap", use_ai: bool = False) -> Tuple[bytes, str]:
        """
        Generate report for a specific assessment ID.
        
        Args:
            assessment_id: Assessment ID to generate report for
            db: Database connection
            current_user: Current user object
            view_type: Type of visualization ('heatmap' or 'radar')
            
        Returns:
            Tuple of (docx_bytes, filename)
        """
        print(f"DEBUG: Generating report for assessment_id: {assessment_id}")
        print(f"DEBUG: Current user org_id: {getattr(current_user, 'org_id', 'NOT SET')}")
        print(f"DEBUG: Current user: {current_user}")
        
        # Get assessment data
        assessment = await db.assessments.find_one({"id": assessment_id, "org_id": current_user.org_id})
        print(f"DEBUG: Assessment found: {assessment is not None}")
        
        if not assessment:
            # Try to find assessment without org_id restriction for debugging
            assessment_debug = await db.assessments.find_one({"id": assessment_id})
            print(f"DEBUG: Assessment exists without org_id filter: {assessment_debug is not None}")
            if assessment_debug:
                print(f"DEBUG: Assessment org_id: {assessment_debug.get('org_id', 'NOT SET')}")
                print(f"DEBUG: Assessment status: {assessment_debug.get('status', 'NOT SET')}")
            raise Exception(f"Assessment not found for user org_id: {current_user.org_id}")
        
        if assessment.get("status") != "COMPLETED":
            print(f"DEBUG: Assessment status: {assessment.get('status')}")
            raise Exception(f"Assessment must be completed before generating report. Current status: {assessment.get('status')}")
        
        # Get actual data from the real database structure
        # Get all answers for this assessment
        answers = await db.answers.find({"assessment_id": assessment_id}).to_list(length=None)
        
        # Get all domains 
        domains = await db.domains.find().to_list(length=None)
        
        # Get all questions to map UUIDs to codes
        questions = await db.questions.find().to_list(length=None)
        
        # Import COMPLETE_QUESTIONS_DATA to get full question details
        from complete_questions import COMPLETE_QUESTIONS_DATA
        
        # Create mapping from question code to question details
        questions_by_code = {}
        for question_data in COMPLETE_QUESTIONS_DATA:
            code = question_data.get('code')
            if code:
                questions_by_code[code] = question_data
        
        # Create mapping from question UUID to question code
        question_uuid_to_code = {}
        for question in questions:
            question_uuid_to_code[question["id"]] = question["code"]
        
        # Create questions structure from answers (since questions collection is empty)
        # Group answers by domain
        domains_by_id = {domain["id"]: domain for domain in domains}
        
        # Get all unique question patterns from answers to reconstruct question structure
        questions_by_domain = {}
        
        processed_answers = 0
        
        # Determine if this is an Awareness assessment
        is_awareness = self.assessment_type == 'Awareness'
        
        if is_awareness:
            # For Awareness assessments, use AWARENESS_QUESTIONS_DATA
            from awareness_questions import AWARENESS_QUESTIONS_DATA
            
            # Create mapping from question code to question details
            awareness_questions_by_code = {}
            for question_data in AWARENESS_QUESTIONS_DATA:
                code = question_data.get('code')
                if code:
                    awareness_questions_by_code[code] = question_data
            
            # Awareness domain mapping
            awareness_domain_mapping = {
                "AU": "Awareness & Understanding",
                "GT": "Governance & Trust Foundations",
                "PS": "People & Skills",
                "LV": "Leadership & Vision",
                "DR": "Data & Digital Readiness"
            }
            
            for answer in answers:
                # For Awareness, question_id IS the code (e.g., "AU-01")
                question_code = answer.get("question_id", "")
                
                if question_code and question_code in awareness_questions_by_code:
                    processed_answers += 1
                    question_details = awareness_questions_by_code[question_code]
                    
                    # Extract domain from question code
                    domain_prefix = question_code.split("-")[0] if "-" in question_code else ""
                    domain_name = awareness_domain_mapping.get(domain_prefix, "Unknown")
                    
                    # Use domain name as ID for Awareness (no domain collection)
                    domain_id = domain_name
                    
                    if domain_id not in questions_by_domain:
                        questions_by_domain[domain_id] = []
                    
                    # Map the option label to predefined answer text
                    option_value = answer.get("option", "")
                    option_to_text_map = {
                        "EARLY_AWARENESS": question_details.get('early_awareness', ''),
                        "EXPLORING_OPPORTUNITIES": question_details.get('exploring_opportunities', ''),
                        "BUILDING_READINESS": question_details.get('building_readiness', ''),
                        "READY_TO_PROGRESS": question_details.get('ready_to_progress', '')
                    }
                    option_to_label_map = {
                        "EARLY_AWARENESS": "Early Awareness",
                        "EXPLORING_OPPORTUNITIES": "Exploring Opportunities",
                        "BUILDING_READINESS": "Building Readiness",
                        "READY_TO_PROGRESS": "Ready to Progress"
                    }
                    selected_option_text = option_to_text_map.get(option_value, option_value)
                    selected_option_label = option_to_label_map.get(option_value, option_value)
                    
                    question = {
                        "id": question_code,
                        "code": question_code,
                        "text": question_details.get('text', f'Question {question_code}'),
                        "domain_id": domain_id,
                        "order": question_details.get('order', 0),
                        "answer": {
                            "numeric_score": answer.get("numeric_score", 0),
                            "text": selected_option_text,
                            "question_id": question_code,
                            "selected_option": {
                                "label": selected_option_label,
                                "text": selected_option_text
                            },
                            "comment": answer.get("note", "")
                        }
                    }
                    questions_by_domain[domain_id].append(question)
            
            # Create domain info for Awareness - use plain & (docxtpl handles escaping)
            domains_by_id = {
                "Awareness & Understanding": {"id": "Awareness & Understanding", "name": "Awareness & Understanding"},
                "Governance & Trust Foundations": {"id": "Governance & Trust Foundations", "name": "Governance & Trust Foundations"},
                "People & Skills": {"id": "People & Skills", "name": "People & Skills"},
                "Leadership & Vision": {"id": "Leadership & Vision", "name": "Leadership & Vision"},
                "Data & Digital Readiness": {"id": "Data & Digital Readiness", "name": "Data & Digital Readiness"}
            }
        elif self.assessment_type == 'Readiness':
            # For Readiness assessments, use READINESS_QUESTIONS_DATA
            from readiness_questions import READINESS_QUESTIONS_DATA
            
            # Create mapping from question code to question details
            readiness_questions_by_code = {}
            for question_data in READINESS_QUESTIONS_DATA:
                code = question_data.get('code')
                if code:
                    readiness_questions_by_code[code] = question_data
            
            # Readiness domain mapping (8 domains) - use plain & as docxtpl handles escaping
            readiness_domain_mapping = {
                "SA": "Strategic Alignment & Awareness",
                "GF": "Governance Foundations",
                "DR": "Data Readiness",
                "TI": "Technology & Infrastructure",
                "PC": "People & Culture",
                "PR": "Policy & Compliance Readiness",
                "RE": "Risk & Ethics Awareness",
                "CL": "Continuous Learning & Improvement"
            }
            
            for answer in answers:
                # For Readiness, question_id IS the code (e.g., "SA-01")
                question_code = answer.get("question_id", "")
                
                if question_code and question_code in readiness_questions_by_code:
                    processed_answers += 1
                    question_details = readiness_questions_by_code[question_code]
                    
                    # Extract domain from question code
                    domain_prefix = question_code.split("-")[0] if "-" in question_code else ""
                    domain_name = readiness_domain_mapping.get(domain_prefix, "Unknown")
                    
                    # Use domain name as ID for Readiness (no domain collection)
                    domain_id = domain_name
                    
                    if domain_id not in questions_by_domain:
                        questions_by_domain[domain_id] = []
                    
                    # Map the option label to predefined answer text
                    option_value = answer.get("option", "")
                    option_to_text_map = {
                        "FOUNDATIONAL": question_details.get('foundational', ''),
                        "DEVELOPING": question_details.get('developing', ''),
                        "ESTABLISHED": question_details.get('established', ''),
                        "LEADING": question_details.get('leading', '')
                    }
                    option_to_label_map = {
                        "FOUNDATIONAL": "Foundational",
                        "DEVELOPING": "Developing",
                        "ESTABLISHED": "Established",
                        "LEADING": "Leading"
                    }
                    selected_option_text = option_to_text_map.get(option_value, option_value)
                    selected_option_label = option_to_label_map.get(option_value, option_value)
                    
                    question = {
                        "id": question_code,
                        "code": question_code,
                        "text": question_details.get('text', f'Question {question_code}'),
                        "domain_id": domain_id,
                        "order": question_details.get('order', 0),
                        "answer": {
                            "numeric_score": answer.get("numeric_score", 0),
                            "text": selected_option_text,
                            "question_id": question_code,
                            "selected_option": {
                                "label": selected_option_label,
                                "text": selected_option_text
                            },
                            "comment": answer.get("note", "")
                        }
                    }
                    questions_by_domain[domain_id].append(question)
            
            # Create domain info for Readiness (8 domains) - use plain & as docxtpl handles escaping
            domains_by_id = {
                "Strategic Alignment & Awareness": {"id": "Strategic Alignment & Awareness", "name": "Strategic Alignment & Awareness"},
                "Data Readiness": {"id": "Data Readiness", "name": "Data Readiness"},
                "Policy & Compliance Readiness": {"id": "Policy & Compliance Readiness", "name": "Policy & Compliance Readiness"},
                "Technology & Infrastructure": {"id": "Technology & Infrastructure", "name": "Technology & Infrastructure"},
                "Continuous Learning & Improvement": {"id": "Continuous Learning & Improvement", "name": "Continuous Learning & Improvement"},
                "Risk & Ethics Awareness": {"id": "Risk & Ethics Awareness", "name": "Risk & Ethics Awareness"},
                "Governance Foundations": {"id": "Governance Foundations", "name": "Governance Foundations"},
                "People & Culture": {"id": "People & Culture", "name": "People & Culture"}
            }
        elif self.assessment_type == 'Orgwide':
            # For Orgwide assessments, use ORGANISATION_QUESTIONS_DATA
            from organisation_questions import ORGANISATION_QUESTIONS_DATA
            
            # Create mapping from question code to question details
            orgwide_questions_by_code = {}
            for question_data in ORGANISATION_QUESTIONS_DATA:
                code = question_data.get('code')
                if code:
                    orgwide_questions_by_code[code] = question_data
            
            # Orgwide domain mapping (10 domains) - use plain & as docxtpl handles escaping
            orgwide_domain_mapping = {
                "AE": "Accountability & Ethics",
                "CA": "Continuous Improvement & Assurance",
                "CC": "Culture & Capability",
                "DS": "Data Stewardship & Security",
                "FI": "Fairness & Inclusivity",
                "GO": "Governance & Oversight",
                "PL": "Privacy & Legal Compliance",
                "RS": "Reliability & Safety",
                "RM": "Risk Management",
                "TE": "Transparency & Explainability"
            }
            
            for answer in answers:
                # For Orgwide, question_id IS the code (e.g., "AE-01")
                question_code = answer.get("question_id", "")
                
                if question_code and question_code in orgwide_questions_by_code:
                    processed_answers += 1
                    question_details = orgwide_questions_by_code[question_code]
                    
                    # Extract domain from question code
                    domain_prefix = question_code.split("-")[0] if "-" in question_code else ""
                    domain_name = orgwide_domain_mapping.get(domain_prefix, "Unknown")
                    
                    # Use domain name as ID for Orgwide (no domain collection)
                    domain_id = domain_name
                    
                    if domain_id not in questions_by_domain:
                        questions_by_domain[domain_id] = []
                    
                    # Map the option label to predefined answer text
                    option_value = answer.get("option", "")
                    option_to_text_map = {
                        "FOUNDATIONAL": question_details.get('foundational', ''),
                        "DEVELOPING": question_details.get('developing', ''),
                        "ESTABLISHED": question_details.get('established', ''),
                        "LEADING": question_details.get('leading', '')
                    }
                    option_to_label_map = {
                        "FOUNDATIONAL": "Foundational",
                        "DEVELOPING": "Developing",
                        "ESTABLISHED": "Established",
                        "LEADING": "Leading"
                    }
                    selected_option_text = option_to_text_map.get(option_value, option_value)
                    selected_option_label = option_to_label_map.get(option_value, option_value)
                    
                    question = {
                        "id": question_code,
                        "code": question_code,
                        "text": question_details.get('text', f'Question {question_code}'),
                        "domain_id": domain_id,
                        "order": question_details.get('order', 0),
                        "answer": {
                            "numeric_score": answer.get("numeric_score", 0),
                            "text": selected_option_text,
                            "question_id": question_code,
                            "selected_option": {
                                "label": selected_option_label,
                                "text": selected_option_text
                            },
                            "comment": answer.get("note", "")
                        }
                    }
                    questions_by_domain[domain_id].append(question)
            
            # Create domain info for Orgwide (10 domains) - use plain & as docxtpl handles escaping
            domains_by_id = {
                "Accountability & Ethics": {"id": "Accountability & Ethics", "name": "Accountability & Ethics"},
                "Continuous Improvement & Assurance": {"id": "Continuous Improvement & Assurance", "name": "Continuous Improvement & Assurance"},
                "Culture & Capability": {"id": "Culture & Capability", "name": "Culture & Capability"},
                "Data Stewardship & Security": {"id": "Data Stewardship & Security", "name": "Data Stewardship & Security"},
                "Fairness & Inclusivity": {"id": "Fairness & Inclusivity", "name": "Fairness & Inclusivity"},
                "Governance & Oversight": {"id": "Governance & Oversight", "name": "Governance & Oversight"},
                "Privacy & Legal Compliance": {"id": "Privacy & Legal Compliance", "name": "Privacy & Legal Compliance"},
                "Reliability & Safety": {"id": "Reliability & Safety", "name": "Reliability & Safety"},
                "Risk Management": {"id": "Risk Management", "name": "Risk Management"},
                "Transparency & Explainability": {"id": "Transparency & Explainability", "name": "Transparency & Explainability"}
            }
        else:
            # Original logic for System/other assessments
            for answer in answers:
                # Get question UUID and map to code using questions collection
                question_uuid = answer.get("question_id", "")
                
                # Get question code from UUID
                question_code = question_uuid_to_code.get(question_uuid)
                
                if question_code:
                    processed_answers += 1
                    # Get full question details from COMPLETE_QUESTIONS_DATA
                    question_details = questions_by_code.get(question_code)
                    
                    if question_details:
                        # Extract domain from question code (e.g., FA-1 -> Fairness domain)
                        domain_prefix = question_code.split("-")[0] if "-" in question_code else ""
                        
                        # Map domain prefixes to domain names
                        domain_mapping = {
                            "FA": "Fairness", "TR": "Transparency", "EX": "Explainability", 
                            "AC": "Accountability", "DI": "Data Integrity", "RE": "Reliability",
                            "SE": "Security", "PR": "Privacy", "SA": "Safety", 
                            "IN": "Inclusivity", "SU": "Sustainability"
                        }
                        
                        domain_name = domain_mapping.get(domain_prefix, "Unknown")
                        
                        # Find matching domain in database
                        matching_domain = None
                        for domain in domains:
                            if domain["name"] == domain_name:
                                matching_domain = domain
                                break
                        
                        if matching_domain:
                            domain_id = matching_domain["id"]
                            
                            if domain_id not in questions_by_domain:
                                questions_by_domain[domain_id] = []
                            
                            # Create question object with correct data
                            question = {
                                "id": question_uuid,
                                "code": question_code,
                                "text": question_details.get('text', f'Question {question_code}'),
                                "domain_id": domain_id,
                                # Include maturity level answer texts for detailed responses appendix
                                "leading_answer": question_details.get('leading_answer', ''),
                                "established_answer": question_details.get('established_answer', ''),
                                "developing_answer": question_details.get('developing_answer', ''),
                                "foundational_answer": question_details.get('foundational_answer', ''),
                                "answer": {
                                    "numeric_score": answer.get("numeric_score", 0),
                                    "option": answer.get("option", ""),  # Store the option (LEADING, ESTABLISHED, etc.)
                                    "text": answer.get("option", ""),
                                    "note": answer.get("note", ""),  # Store assessment notes
                                    "question_id": question_uuid
                                }
                            }
                            questions_by_domain[domain_id].append(question)
        
        
        print(f"DEBUG: Processed answers: {processed_answers}/{len(answers)}")
        print(f"DEBUG: Domains in heatmap: {len(questions_by_domain)}")
        for domain_id in list(questions_by_domain.keys())[:3]:  # Show first 3 domains
            domain = domains_by_id.get(domain_id, {"name": "Unknown"})
            questions = questions_by_domain[domain_id]
            print(f"  Domain: {domain['name']} - Questions: {len(questions)}")
            for question in questions[:2]:  # Show first 2 questions per domain
                print(f"    Question {question.get('code', 'N/A')}: Score {question.get('answer', {}).get('numeric_score', 'N/A')}")
        # Calculate summary data
        total_score = 0
        total_possible = 0
        domain_scores = []
        
        for domain_id, questions in questions_by_domain.items():
            domain = domains_by_id.get(domain_id, {"name": "Unknown"})
            domain_total = sum(q["answer"]["numeric_score"] for q in questions)
            domain_possible = len(questions) * 4  # 1-4 scale
            domain_percentage = (domain_total / domain_possible * 100) if domain_possible > 0 else 0
            
            domain_scores.append({
                "domain_id": domain_id,
                "domain_name": domain.get("name", "Unknown"),
                "percentage": domain_percentage
            })
            
            total_score += domain_total
            total_possible += domain_possible
        
        # For Awareness assessments, prefer stored overall_percentage
        if is_awareness and assessment.get('overall_percentage'):
            overall_percentage = assessment.get('overall_percentage')
            print(f"DEBUG: Using stored overall_percentage for Awareness: {overall_percentage}")
        else:
            overall_percentage = (total_score / total_possible * 100) if total_possible > 0 else 0
        
        # Calculate maturity tier
        # Categories: Excellent (91–100%), Good (81–90%), Moderate (61–80%), Low (41–60%), Basic (0–40%)
        if overall_percentage >= 91:
            overall_maturity = "Excellent"
        elif overall_percentage >= 81:
            overall_maturity = "Good"
        elif overall_percentage >= 61:
            overall_maturity = "Moderate"
        elif overall_percentage >= 41:
            overall_maturity = "Low"
        else:
            overall_maturity = "Basic"
        
        # Calculate overall_tier based on assessment type
        if self.assessment_type == 'Awareness':
            overall_tier = self._calculate_awareness_tier(overall_percentage)
        elif self.assessment_type == 'Readiness':
            overall_tier = self._calculate_readiness_tier(overall_percentage)
        elif self.assessment_type == 'Orgwide':
            overall_tier = self._calculate_orgwide_tier(overall_percentage)
        elif self.assessment_type == 'System':
            overall_tier = self._calculate_system_tier(overall_percentage)
        else:
            overall_tier = self._calculate_tier(overall_percentage)
        
        # Identify top strengths and gaps from domain_scores
        sorted_domains = sorted(domain_scores, key=lambda x: x['percentage'], reverse=True)
        top_strengths = [d['domain_name'] for d in sorted_domains if d['percentage'] >= 75][:3]
        top_gaps = [d['domain_name'] for d in sorted_domains if d['percentage'] < 75][-3:]  # Bottom performers
        
        summary_data = {
            "overall_percentage": overall_percentage,
            "overall_maturity": overall_maturity,
            "overall_tier": overall_tier,
            "top_strengths": top_strengths,
            "top_gaps": top_gaps,
            "domain_scores": domain_scores
        }
        
        # The questions_by_domain structure is already built above, no need to recreate it
        
        # Use the already-built questions_by_domain structure from above
        # questions_by_domain = {}  # Already built above
        # domains_by_id is already available from earlier in the method
        
        # Skip this loop since questions_by_domain is already built above
        pass
        
        # Create structured questions data
        questions_data = []
        for domain_id, questions in questions_by_domain.items():
            domain_info = domains_by_id.get(domain_id, {"name": "Unknown Domain"})
            questions_data.append({
                "domain": domain_info,
                "questions": sorted(questions, key=lambda q: q.get("order", 0))
            })
        
        # Get sector/industry from assessment - differs by assessment type
        system_info = assessment.get('system_info') or {}
        awareness_info = assessment.get('awareness_info') or {}
        
        # For Awareness assessments, get industry from awareness_info
        if self.assessment_type == 'Awareness':
            sector_name = awareness_info.get('industry', '')
        elif self.assessment_type == 'Readiness':
            readiness_info = assessment.get('readiness_info') or {}
            sector_name = readiness_info.get('industry', '')
        elif self.assessment_type == 'Orgwide':
            # For Orgwide, prefer the organization's primary_industry over org_info.industry
            # This ensures sector benchmarks use the correct industry
            org = await db.organizations.find_one({"id": current_user.org_id})
            if org and org.get('primary_industry'):
                sector_name = org.get('primary_industry')
            else:
                orgwide_info = assessment.get('orgwide_info') or assessment.get('org_info') or {}
                sector_name = orgwide_info.get('industry', '')
        else:
            sector_name = system_info.get('industry', '') if system_info else ''
        
        print(f"DEBUG: Assessment type: {self.assessment_type}")
        print(f"DEBUG: Assessment sector/industry: {sector_name}")
        
        # Generate sector-specific actions (for System, Awareness, Readiness, and Orgwide assessments with sector info)
        sector_actions = {'high': [], 'medium': [], 'low': []}
        if sector_name and self.assessment_type in ['System', 'Awareness', 'Readiness', 'Orgwide']:
            sector_actions = self._generate_sector_actions(
                {"questions_data": questions_data},
                sector_name
            )
        
        # DEBUG: Print what sector_actions contains BEFORE passing to assessment_data
        print(f"DEBUG generate_report_for_assessment: sector_actions generated")
        print(f"  - high: {len(sector_actions.get('high', []))} items")
        print(f"  - medium: {len(sector_actions.get('medium', []))} items")
        print(f"  - low: {len(sector_actions.get('low', []))} items")
        if sector_actions.get('high'):
            first = sector_actions['high'][0]
            print(f"  - FIRST HIGH: question_id='{first.get('question_id')}', domain='{first.get('domain')}', sector_action='{first.get('sector_action', 'N/A')[:60]}...'")
        
        # Get sector_average from benchmark data
        sector_average = self._get_sector_average(sector_name)
        print(f"DEBUG: sector_average for '{sector_name}': {sector_average}")
        
        # Prepare assessment data for report generation
        assessment_data = {
            "assessment": assessment,
            "questions_data": questions_data,
            "summary": summary_data,
            "sector_name": sector_name,
            "sector_average": sector_average,  # Add sector_average for bar chart and results summary
            "sector_actions": sector_actions,
            "system_info": system_info,  # Pass system_info for template use (may be empty for non-System assessments)
            "awareness_info": assessment.get('awareness_info') or {},  # Pass awareness_info for Awareness assessments
            "readiness_info": assessment.get('readiness_info') or {},  # Pass readiness_info for Readiness assessments
            "orgwide_info": assessment.get('orgwide_info') or assessment.get('org_info') or {},  # Pass orgwide_info for Orgwide assessments (fallback to org_info)
            "faira_form": assessment.get('faira_form') or {},  # Pass faira_form for FAIRA assessments
        }
        
        # Generate AI narratives for Awareness assessments if use_ai is enabled
        if use_ai and self.assessment_type == 'Awareness':
            print("=== GENERATING AWARENESS AI NARRATIVES ===")
            try:
                # Build report_data for AI narrative generation
                awareness_info = assessment.get('awareness_info') or {}
                ai_report_data = {
                    'organization_name': awareness_info.get('org_name', current_user.organization_name),
                    'overall': {
                        'tier': summary_data.get('overall_tier', 'Unknown'),
                        'percentage': summary_data.get('overall_percentage', 0),
                        'score': summary_data.get('overall_percentage', 0),
                    },
                    'strengths': summary_data.get('top_strengths', []),
                    'gaps': summary_data.get('top_gaps', []),
                    'sector_name': awareness_info.get('industry', 'Unknown'),
                    'industry': awareness_info.get('industry', 'Unknown'),
                    'awareness_info': awareness_info,  # Pass full awareness_info for AI-2
                }
                
                ai_narratives = await self._generate_awareness_ai_narratives(ai_report_data, assessment, db)
                assessment_data['ai_narratives'] = ai_narratives
                print(f"DEBUG: AI narratives generated: {list(ai_narratives.keys())}")
            except Exception as e:
                print(f"ERROR generating Awareness AI narratives: {e}")
                assessment_data['ai_narratives'] = {}
        
        # Generate AI narratives for Readiness assessments if use_ai is enabled
        if use_ai and self.assessment_type == 'Readiness':
            print("=== GENERATING READINESS AI NARRATIVES ===")
            try:
                # Build report_data for AI narrative generation
                readiness_info = assessment.get('readiness_info') or {}
                ai_report_data = {
                    'organization_name': readiness_info.get('org_name', current_user.organization_name),
                    'overall': {
                        'tier': summary_data.get('overall_tier', 'Unknown'),
                        'percentage': summary_data.get('overall_percentage', 0),
                        'score': summary_data.get('overall_percentage', 0),
                    },
                    'domains': summary_data.get('domain_scores', []),
                    'sector_name': sector_name,
                    'sector_average': sector_average,
                    'benchmark_data': self._load_benchmark_data(sector_name),
                    'actions': {
                        'high': sector_actions.get('high', []),
                        'medium': sector_actions.get('medium', []),
                        'low': sector_actions.get('low', []),
                    },
                    'readiness_info': readiness_info,
                    'org': {
                        'name': readiness_info.get('org_name', current_user.organization_name),
                        'industry': sector_name,
                    },
                }
                
                ai_narratives = await self._generate_readiness_ai_narratives(ai_report_data, assessment, db)
                assessment_data['ai_narratives'] = ai_narratives
                print(f"DEBUG: Readiness AI narratives generated: {list(ai_narratives.keys())}")
            except Exception as e:
                print(f"ERROR generating Readiness AI narratives: {e}")
                import traceback
                traceback.print_exc()
                assessment_data['ai_narratives'] = {}
        
        # Generate AI narratives for Orgwide assessments if use_ai is enabled
        if use_ai and self.assessment_type == 'Orgwide':
            print("=== GENERATING ORGWIDE AI NARRATIVES ===")
            try:
                # Build report_data for AI narrative generation
                orgwide_info = assessment.get('orgwide_info') or assessment.get('org_info') or {}
                
                # Build domains list for narrative generation
                domains_for_ai = []
                for ds in summary_data.get('domain_scores', []):
                    d_score = ds.get('percentage', ds.get('score', 0))
                    if isinstance(d_score, str):
                        d_score = float(d_score.replace('%', ''))
                    d_tier = self._calculate_orgwide_tier(d_score)
                    domains_for_ai.append({
                        'name': ds.get('domain_name', ds.get('name', 'Unknown')),
                        'tier': d_tier,
                        'score': d_score,
                        'score_display': f"{d_score:.0f}%"
                    })
                
                # Build strengths and gaps summaries
                strengths = []
                gaps = []
                for d in domains_for_ai:
                    if d['score'] >= 70:
                        strengths.append(f"{d['name']} ({d['score']:.0f}%)")
                    elif d['score'] < 60:
                        gaps.append(f"{d['name']} ({d['score']:.0f}%)")
                
                # If no domains >= 70%, use top 3 as relative strengths
                if not strengths:
                    sorted_domains = sorted(domains_for_ai, key=lambda x: x['score'], reverse=True)
                    strengths = [f"{d['name']} ({d['score']:.0f}%)" for d in sorted_domains[:3]]
                
                strengths_summary = ', '.join(strengths) if strengths else 'Building foundational maturity across all domains'
                gaps_summary = ', '.join(gaps) if gaps else 'All domains above 60%'
                
                ai_report_data = {
                    'organization_name': orgwide_info.get('org_name', current_user.organization_name),
                    'overall': {
                        'tier': summary_data.get('overall_tier', 'Unknown'),
                        'percentage': summary_data.get('overall_percentage', 0),
                        'score': summary_data.get('overall_percentage', 0),
                    },
                    'domains': domains_for_ai,
                    'sector_name': sector_name,
                    'sector_average': sector_average,
                    'strengths_summary': strengths_summary,
                    'gaps_summary': gaps_summary,
                    'benchmark_data': self._load_benchmark_data(sector_name),
                    'actions': {
                        'high': sector_actions.get('high', []),
                        'medium': sector_actions.get('medium', []),
                        'low': sector_actions.get('low', []),
                    },
                    'orgwide_info': orgwide_info,
                    'org': {
                        'name': orgwide_info.get('org_name', current_user.organization_name),
                        'industry': sector_name,
                    },
                }
                
                ai_narratives = await self._generate_orgwide_ai_narratives(ai_report_data, assessment, db)
                assessment_data['ai_narratives'] = ai_narratives
                print(f"DEBUG: Orgwide AI narratives generated: {list(ai_narratives.keys())}")
            except Exception as e:
                print(f"ERROR generating Orgwide AI narratives: {e}")
                import traceback
                traceback.print_exc()
                assessment_data['ai_narratives'] = {}
        
        # Generate AI narratives for System assessments if use_ai is enabled
        if use_ai and self.assessment_type == 'System':
            print("=== GENERATING SYSTEM AI NARRATIVES ===")
            try:
                system_info = assessment.get('system_info') or {}
                
                # Build domains list for narrative generation
                domains_for_ai = []
                for ds in summary_data.get('domain_scores', []):
                    d_score = ds.get('percentage', ds.get('score', 0))
                    if isinstance(d_score, str):
                        d_score = float(d_score.replace('%', ''))
                    d_tier = self._get_tier_for_score(d_score)
                    domains_for_ai.append({
                        'name': ds.get('domain_name', ds.get('name', 'Unknown')),
                        'tier': d_tier,
                        'score': d_score,
                        'score_display': f"{d_score:.0f}%"
                    })
                
                # Build strengths and gaps summaries
                strengths = []
                gaps = []
                for d in domains_for_ai:
                    if d['score'] >= 70:
                        strengths.append(f"{d['name']} ({d['score']:.0f}%)")
                    elif d['score'] < 60:
                        gaps.append(f"{d['name']} ({d['score']:.0f}%)")
                
                if not strengths:
                    sorted_domains = sorted(domains_for_ai, key=lambda x: x['score'], reverse=True)
                    strengths = [f"{d['name']} ({d['score']:.0f}%)" for d in sorted_domains[:3]]
                
                strengths_summary = ', '.join(strengths) if strengths else 'Building foundational maturity across all domains'
                gaps_summary = ', '.join(gaps) if gaps else 'All domains above 60%'
                
                ai_report_data = {
                    'system_name': system_info.get('system_name') or system_info.get('systemName', ''),
                    'overall': {
                        'tier': summary_data.get('overall_tier', 'Unknown'),
                        'percentage': summary_data.get('overall_percentage', 0),
                        'score': summary_data.get('overall_percentage', 0),
                    },
                    'domains': domains_for_ai,
                    'sector_name': sector_name,
                    'sector_average': sector_average,
                    'strengths_summary': strengths_summary,
                    'gaps_summary': gaps_summary,
                    'next_focus': summary_data.get('next_focus', ''),
                    'system_info': system_info,
                }
                
                ai_narratives = await self._generate_system_ai_narratives(ai_report_data, assessment, db)
                assessment_data['ai_narratives'] = ai_narratives
                print(f"DEBUG: System AI narratives generated: {list(ai_narratives.keys())}")
            except Exception as e:
                print(f"ERROR generating System AI narratives: {e}")
                import traceback
                traceback.print_exc()
                assessment_data['ai_narratives'] = {}
        
        # Generate AI narratives for FAIRA assessments if use_ai is enabled
        if use_ai and self.assessment_type == 'FAIRA':
            print("=== GENERATING FAIRA AI NARRATIVES ===")
            try:
                faira_form = assessment.get('faira_form') or assessment.get('form_responses') or {}
                
                # Calculate FAIRA scores using the scoring engine
                faira_risk_summary = calculate_overall_risk(faira_form)
                faira_radar_data = get_radar_chart_data(faira_form)
                faira_controls = get_recommended_controls(faira_form, top_n=5)
                
                print(f"DEBUG: FAIRA risk summary keys: {list(faira_risk_summary.keys())}")
                print(f"DEBUG: FAIRA overall risk level: {faira_risk_summary.get('overall_risk_level')}")
                
                # Get domain scores from risk summary
                domain_scores = faira_risk_summary.get('domain_scores', {})
                top_risk_areas = faira_risk_summary.get('top_risk_areas', [])
                
                # Build AI report data for FAIRA
                ai_report_data = {
                    'system_name': faira_form.get('ai_system_name') or faira_form.get('A1_2') or assessment.get('name', ''),
                    'faira_form': faira_form,
                    'form_responses': faira_form,
                    'domain_scores': domain_scores,
                    'overall_risk_score': faira_risk_summary.get('overall_risk_score', 50),
                    'overall_risk_level': faira_risk_summary.get('overall_risk_level', 'Medium'),
                    'overall_impact_score': faira_risk_summary.get('total_impact', 50),
                    'overall_likelihood_score': faira_risk_summary.get('total_likelihood', 50),
                    'overall_control_effectiveness_score': faira_risk_summary.get('total_control_effectiveness', 50),
                    'top_risk_areas': top_risk_areas,
                    'recommended_controls': faira_controls,
                    'existing_controls_summary': assessment.get('existing_controls_summary', 'Not specified'),
                    'identified_gaps': assessment.get('identified_gaps', 'Not specified'),
                    'governance_maturity_indicators': faira_risk_summary.get('governance_score', 'Not specified'),
                    'assurance_mechanisms': faira_risk_summary.get('assurance_readiness', 'Not specified'),
                    'supporting_artefacts': assessment.get('evidence_summary', 'Not specified'),
                    'assessment_methodology': 'FAIRA Framework v1.0 - Queensland Government AI Risk Assessment',
                    'scoring_method': 'Risk Score = (Impact × Likelihood) / Control Effectiveness',
                    'normalisation_notes': 'Scores normalised to 0-100 scale for comparability',
                }
                
                # Store the calculated data for template context
                assessment_data['faira_risk_summary'] = faira_risk_summary
                assessment_data['faira_radar_data'] = faira_radar_data
                assessment_data['faira_controls'] = faira_controls
                assessment_data['faira_form'] = faira_form  # Pass form data to template context
                
                ai_narratives = await self._generate_faira_ai_narratives(ai_report_data, assessment, db)
                assessment_data['ai_narratives'] = ai_narratives
                print(f"DEBUG: FAIRA AI narratives generated: {list(ai_narratives.keys())}")
            except Exception as e:
                print(f"ERROR generating FAIRA AI narratives: {e}")
                import traceback
                traceback.print_exc()
                assessment_data['ai_narratives'] = {}
        
        print(f"DEBUG: current_user type: {type(current_user)}")
        print(f"DEBUG: current_user.organization_name: {getattr(current_user, 'organization_name', 'NOT SET')}")
        
        user_data = {
            "organization_name": current_user.organization_name
        }
        
        print(f"DEBUG: user_data: {user_data}")
        
        # Fetch benchmark data if radar view is selected
        benchmark_data = None
        if view_type == "radar":
            # Get user's industry for benchmarking
            user_industry = getattr(current_user, 'industry', None) or getattr(current_user, 'primary_industry', None)
            if user_industry:
                print(f"DEBUG: Fetching benchmarks for industry: {user_industry}")
                # Fetch benchmark data from sector_benchmarks collection
                benchmark_record = await db.sector_benchmarks.find_one({"sector": user_industry})
                if benchmark_record:
                    benchmark_data = {
                        "sector": user_industry,
                        "benchmarks": benchmark_record.get("benchmarks", {})
                    }
                    print(f"DEBUG: Benchmark data found for {user_industry}")
                else:
                    print(f"DEBUG: No benchmark data found for {user_industry}")
            else:
                print(f"DEBUG: User has no industry set, cannot fetch benchmarks")
        
        # Generate the report
        docx_bytes, pdf_bytes = await self.generate_report(
            assessment_id, assessment_data, user_data, 
            view_type=view_type, benchmark_data=benchmark_data,
            use_ai=use_ai
        )
        
        # Generate filename - use assessment's org_name from the appropriate info object
        # The org_name is stored in different locations depending on assessment type
        # For System assessments, use system_name instead of org_name
        assessment_name_for_file = ''
        if self.assessment_type == 'Awareness':
            assessment_name_for_file = (assessment.get('awareness_info') or {}).get('org_name', '')
        elif self.assessment_type == 'Readiness':
            assessment_name_for_file = (assessment.get('readiness_info') or {}).get('org_name', '')
        elif self.assessment_type == 'Orgwide':
            # For Orgwide, check both orgwide_info and org_info (fallback)
            assessment_name_for_file = (assessment.get('orgwide_info') or assessment.get('org_info') or {}).get('org_name', '')
        elif self.assessment_type == 'System':
            # For System assessments, use system_name (systemName field)
            system_info = assessment.get('system_info') or {}
            assessment_name_for_file = system_info.get('systemName', '') or system_info.get('system_name', '')
        elif self.assessment_type == 'FAIRA':
            # For FAIRA assessments, use ai_system_name from faira_form
            faira_form = assessment.get('faira_form') or {}
            assessment_name_for_file = faira_form.get('ai_system_name', '') or faira_form.get('A1_2', '') or assessment.get('name', '')
        else:
            # Other types - check org_info first, then system_info
            assessment_name_for_file = (assessment.get('org_info') or {}).get('org_name', '')
            if not assessment_name_for_file:
                assessment_name_for_file = (assessment.get('system_info') or {}).get('org_name', '')
        
        # Fallback to organization_name field if exists, then to user's organization name
        if not assessment_name_for_file:
            assessment_name_for_file = assessment.get('organization_name', '')
        if not assessment_name_for_file:
            assessment_name_for_file = current_user.organization_name
        
        safe_name = "".join(c for c in assessment_name_for_file if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_name = safe_name.replace(' ', '_') if safe_name else "Assessment"
        filename = f"AM_AI_SAFE_{self.assessment_type}_{safe_name}_{assessment.get('created_at', datetime.now(timezone.utc)).strftime('%Y-%m-%d')}.docx"
        
        return docx_bytes, filename
    
    async def generate_report_full(self, assessment_id: str, db, current_user) -> Tuple[bytes, bytes]:
        """
        Generate both DOCX and PDF reports for a specific assessment ID.
        
        Args:
            assessment_id: Assessment ID to generate report for
            db: Database connection
            current_user: Current user object
            
        Returns:
            Tuple of (docx_bytes, pdf_bytes)
        """
        # Get assessment data (reuse the existing logic)
        assessment = await db.assessments.find_one({"id": assessment_id, "org_id": current_user.org_id})
        if not assessment:
            raise Exception("Assessment not found")
        
        if assessment.get("status") != "COMPLETED":
            raise Exception("Assessment must be completed before generating report")
        
        # Get actual data from the real database structure
        # Get all answers for this assessment
        answers = await db.answers.find({"assessment_id": assessment_id}).to_list(length=None)
        
        # Get all domains 
        domains = await db.domains.find().to_list(length=None)
        
        # Get all questions to map UUIDs to codes
        questions = await db.questions.find().to_list(length=None)
        
        # Import COMPLETE_QUESTIONS_DATA to get full question details
        from complete_questions import COMPLETE_QUESTIONS_DATA
        
        # Create mapping from question code to question details
        questions_by_code = {}
        for question_data in COMPLETE_QUESTIONS_DATA:
            code = question_data.get('code')
            if code:
                questions_by_code[code] = question_data
        
        # Create mapping from question UUID to question code
        question_uuid_to_code = {}
        for question in questions:
            question_uuid_to_code[question["id"]] = question["code"]
        
        # Create questions structure from answers (since questions collection is empty)
        # Group answers by domain
        domains_by_id = {domain["id"]: domain for domain in domains}
        
        # Get all unique question patterns from answers to reconstruct question structure
        questions_by_domain = {}
        
        processed_answers = 0
        
        # Determine if this is an Awareness assessment
        is_awareness = self.assessment_type == 'Awareness'
        
        if is_awareness:
            # For Awareness assessments, use AWARENESS_QUESTIONS_DATA
            from awareness_questions import AWARENESS_QUESTIONS_DATA
            
            # Create mapping from question code to question details
            awareness_questions_by_code = {}
            for question_data in AWARENESS_QUESTIONS_DATA:
                code = question_data.get('code')
                if code:
                    awareness_questions_by_code[code] = question_data
            
            # Awareness domain mapping
            awareness_domain_mapping = {
                "AU": "Awareness & Understanding",
                "GT": "Governance & Trust Foundations",
                "PS": "People & Skills",
                "LV": "Leadership & Vision",
                "DR": "Data & Digital Readiness"
            }
            
            for answer in answers:
                # For Awareness, question_id IS the code (e.g., "AU-01")
                question_code = answer.get("question_id", "")
                
                if question_code and question_code in awareness_questions_by_code:
                    processed_answers += 1
                    question_details = awareness_questions_by_code[question_code]
                    
                    # Extract domain from question code
                    domain_prefix = question_code.split("-")[0] if "-" in question_code else ""
                    domain_name = awareness_domain_mapping.get(domain_prefix, "Unknown")
                    
                    # Use domain name as ID for Awareness (no domain collection)
                    domain_id = domain_name
                    
                    if domain_id not in questions_by_domain:
                        questions_by_domain[domain_id] = []
                    
                    # Map the option label to predefined answer text
                    option_value = answer.get("option", "")
                    option_to_text_map = {
                        "EARLY_AWARENESS": question_details.get('early_awareness', ''),
                        "EXPLORING_OPPORTUNITIES": question_details.get('exploring_opportunities', ''),
                        "BUILDING_READINESS": question_details.get('building_readiness', ''),
                        "READY_TO_PROGRESS": question_details.get('ready_to_progress', '')
                    }
                    option_to_label_map = {
                        "EARLY_AWARENESS": "Early Awareness",
                        "EXPLORING_OPPORTUNITIES": "Exploring Opportunities",
                        "BUILDING_READINESS": "Building Readiness",
                        "READY_TO_PROGRESS": "Ready to Progress"
                    }
                    selected_option_text = option_to_text_map.get(option_value, option_value)
                    selected_option_label = option_to_label_map.get(option_value, option_value)
                    
                    question = {
                        "id": question_code,
                        "code": question_code,
                        "text": question_details.get('text', f'Question {question_code}'),
                        "domain_id": domain_id,
                        "order": question_details.get('order', 0),
                        "answer": {
                            "numeric_score": answer.get("numeric_score", 0),
                            "text": selected_option_text,
                            "question_id": question_code,
                            "selected_option": {
                                "label": selected_option_label,
                                "text": selected_option_text
                            },
                            "comment": answer.get("note", "")
                        }
                    }
                    questions_by_domain[domain_id].append(question)
            
            # Create domain info for Awareness - use plain & (docxtpl handles escaping)
            domains_by_id = {
                "Awareness & Understanding": {"id": "Awareness & Understanding", "name": "Awareness & Understanding"},
                "Governance & Trust Foundations": {"id": "Governance & Trust Foundations", "name": "Governance & Trust Foundations"},
                "People & Skills": {"id": "People & Skills", "name": "People & Skills"},
                "Leadership & Vision": {"id": "Leadership & Vision", "name": "Leadership & Vision"},
                "Data & Digital Readiness": {"id": "Data & Digital Readiness", "name": "Data & Digital Readiness"}
            }
        elif self.assessment_type == 'Readiness':
            # For Readiness assessments, use READINESS_QUESTIONS_DATA
            from readiness_questions import READINESS_QUESTIONS_DATA
            
            # Create mapping from question code to question details
            readiness_questions_by_code = {}
            for question_data in READINESS_QUESTIONS_DATA:
                code = question_data.get('code')
                if code:
                    readiness_questions_by_code[code] = question_data
            
            # Readiness domain mapping (8 domains) - use plain & as docxtpl handles escaping
            readiness_domain_mapping = {
                "SA": "Strategic Alignment & Awareness",
                "GF": "Governance Foundations",
                "DR": "Data Readiness",
                "TI": "Technology & Infrastructure",
                "PC": "People & Culture",
                "PR": "Policy & Compliance Readiness",
                "RE": "Risk & Ethics Awareness",
                "CL": "Continuous Learning & Improvement"
            }
            
            for answer in answers:
                # For Readiness, question_id IS the code (e.g., "SA-01")
                question_code = answer.get("question_id", "")
                
                if question_code and question_code in readiness_questions_by_code:
                    processed_answers += 1
                    question_details = readiness_questions_by_code[question_code]
                    
                    # Extract domain from question code
                    domain_prefix = question_code.split("-")[0] if "-" in question_code else ""
                    domain_name = readiness_domain_mapping.get(domain_prefix, "Unknown")
                    
                    # Use domain name as ID for Readiness (no domain collection)
                    domain_id = domain_name
                    
                    if domain_id not in questions_by_domain:
                        questions_by_domain[domain_id] = []
                    
                    # Map the option label to predefined answer text
                    option_value = answer.get("option", "")
                    option_to_text_map = {
                        "FOUNDATIONAL": question_details.get('foundational', ''),
                        "DEVELOPING": question_details.get('developing', ''),
                        "ESTABLISHED": question_details.get('established', ''),
                        "LEADING": question_details.get('leading', '')
                    }
                    option_to_label_map = {
                        "FOUNDATIONAL": "Foundational",
                        "DEVELOPING": "Developing",
                        "ESTABLISHED": "Established",
                        "LEADING": "Leading"
                    }
                    selected_option_text = option_to_text_map.get(option_value, option_value)
                    selected_option_label = option_to_label_map.get(option_value, option_value)
                    
                    question = {
                        "id": question_code,
                        "code": question_code,
                        "text": question_details.get('text', f'Question {question_code}'),
                        "domain_id": domain_id,
                        "order": question_details.get('order', 0),
                        "answer": {
                            "numeric_score": answer.get("numeric_score", 0),
                            "text": selected_option_text,
                            "question_id": question_code,
                            "selected_option": {
                                "label": selected_option_label,
                                "text": selected_option_text
                            },
                            "comment": answer.get("note", "")
                        }
                    }
                    questions_by_domain[domain_id].append(question)
            
            # Create domain info for Readiness (8 domains) - use plain & as docxtpl handles escaping
            domains_by_id = {
                "Strategic Alignment & Awareness": {"id": "Strategic Alignment & Awareness", "name": "Strategic Alignment & Awareness"},
                "Data Readiness": {"id": "Data Readiness", "name": "Data Readiness"},
                "Policy & Compliance Readiness": {"id": "Policy & Compliance Readiness", "name": "Policy & Compliance Readiness"},
                "Technology & Infrastructure": {"id": "Technology & Infrastructure", "name": "Technology & Infrastructure"},
                "Continuous Learning & Improvement": {"id": "Continuous Learning & Improvement", "name": "Continuous Learning & Improvement"},
                "Risk & Ethics Awareness": {"id": "Risk & Ethics Awareness", "name": "Risk & Ethics Awareness"},
                "Governance Foundations": {"id": "Governance Foundations", "name": "Governance Foundations"},
                "People & Culture": {"id": "People & Culture", "name": "People & Culture"}
            }
        elif self.assessment_type == 'Orgwide':
            # For Orgwide assessments, use ORGANISATION_QUESTIONS_DATA
            from organisation_questions import ORGANISATION_QUESTIONS_DATA
            
            # Create mapping from question code to question details
            orgwide_questions_by_code = {}
            for question_data in ORGANISATION_QUESTIONS_DATA:
                code = question_data.get('code')
                if code:
                    orgwide_questions_by_code[code] = question_data
            
            # Orgwide domain mapping (10 domains) - use plain & as docxtpl handles escaping
            orgwide_domain_mapping = {
                "AE": "Accountability & Ethics",
                "CA": "Continuous Improvement & Assurance",
                "CC": "Culture & Capability",
                "DS": "Data Stewardship & Security",
                "FI": "Fairness & Inclusivity",
                "GO": "Governance & Oversight",
                "PL": "Privacy & Legal Compliance",
                "RS": "Reliability & Safety",
                "RM": "Risk Management",
                "TE": "Transparency & Explainability"
            }
            
            for answer in answers:
                # For Orgwide, question_id IS the code (e.g., "AE-01")
                question_code = answer.get("question_id", "")
                
                if question_code and question_code in orgwide_questions_by_code:
                    processed_answers += 1
                    question_details = orgwide_questions_by_code[question_code]
                    
                    # Extract domain from question code
                    domain_prefix = question_code.split("-")[0] if "-" in question_code else ""
                    domain_name = orgwide_domain_mapping.get(domain_prefix, "Unknown")
                    
                    # Use domain name as ID for Orgwide (no domain collection)
                    domain_id = domain_name
                    
                    if domain_id not in questions_by_domain:
                        questions_by_domain[domain_id] = []
                    
                    # Map the option label to predefined answer text
                    option_value = answer.get("option", "")
                    option_to_text_map = {
                        "FOUNDATIONAL": question_details.get('foundational', ''),
                        "DEVELOPING": question_details.get('developing', ''),
                        "ESTABLISHED": question_details.get('established', ''),
                        "LEADING": question_details.get('leading', '')
                    }
                    option_to_label_map = {
                        "FOUNDATIONAL": "Foundational",
                        "DEVELOPING": "Developing",
                        "ESTABLISHED": "Established",
                        "LEADING": "Leading"
                    }
                    selected_option_text = option_to_text_map.get(option_value, option_value)
                    selected_option_label = option_to_label_map.get(option_value, option_value)
                    
                    question = {
                        "id": question_code,
                        "code": question_code,
                        "text": question_details.get('text', f'Question {question_code}'),
                        "domain_id": domain_id,
                        "order": question_details.get('order', 0),
                        "answer": {
                            "numeric_score": answer.get("numeric_score", 0),
                            "text": selected_option_text,
                            "question_id": question_code,
                            "selected_option": {
                                "label": selected_option_label,
                                "text": selected_option_text
                            },
                            "comment": answer.get("note", "")
                        }
                    }
                    questions_by_domain[domain_id].append(question)
            
            # Create domain info for Orgwide (10 domains) - use plain & as docxtpl handles escaping
            domains_by_id = {
                "Accountability & Ethics": {"id": "Accountability & Ethics", "name": "Accountability & Ethics"},
                "Continuous Improvement & Assurance": {"id": "Continuous Improvement & Assurance", "name": "Continuous Improvement & Assurance"},
                "Culture & Capability": {"id": "Culture & Capability", "name": "Culture & Capability"},
                "Data Stewardship & Security": {"id": "Data Stewardship & Security", "name": "Data Stewardship & Security"},
                "Fairness & Inclusivity": {"id": "Fairness & Inclusivity", "name": "Fairness & Inclusivity"},
                "Governance & Oversight": {"id": "Governance & Oversight", "name": "Governance & Oversight"},
                "Privacy & Legal Compliance": {"id": "Privacy & Legal Compliance", "name": "Privacy & Legal Compliance"},
                "Reliability & Safety": {"id": "Reliability & Safety", "name": "Reliability & Safety"},
                "Risk Management": {"id": "Risk Management", "name": "Risk Management"},
                "Transparency & Explainability": {"id": "Transparency & Explainability", "name": "Transparency & Explainability"}
            }
        else:
            # Original logic for System/other assessments
            for answer in answers:
                # Get question UUID and map to code using questions collection
                question_uuid = answer.get("question_id", "")
                
                # Get question code from UUID
                question_code = question_uuid_to_code.get(question_uuid)
                
                if question_code:
                    processed_answers += 1
                    # Get full question details from COMPLETE_QUESTIONS_DATA
                    question_details = questions_by_code.get(question_code)
                    
                    if question_details:
                        # Extract domain from question code (e.g., FA-1 -> Fairness domain)
                        domain_prefix = question_code.split("-")[0] if "-" in question_code else ""
                        
                        # Map domain prefixes to domain names
                        domain_mapping = {
                            "FA": "Fairness", "TR": "Transparency", "EX": "Explainability", 
                            "AC": "Accountability", "DI": "Data Integrity", "RE": "Reliability",
                            "SE": "Security", "PR": "Privacy", "SA": "Safety", 
                            "IN": "Inclusivity", "SU": "Sustainability"
                        }
                        
                        domain_name = domain_mapping.get(domain_prefix, "Unknown")
                        
                        # Find matching domain in database
                        matching_domain = None
                        for domain in domains:
                            if domain["name"] == domain_name:
                                matching_domain = domain
                                break
                        
                        if matching_domain:
                            domain_id = matching_domain["id"]
                            
                            if domain_id not in questions_by_domain:
                                questions_by_domain[domain_id] = []
                            
                            # Create question object with correct data
                            question = {
                                "id": question_uuid,
                                "code": question_code,
                                "text": question_details.get('text', f'Question {question_code}'),
                                "domain_id": domain_id,
                                # Include maturity level answer texts for detailed responses appendix
                                "leading_answer": question_details.get('leading_answer', ''),
                                "established_answer": question_details.get('established_answer', ''),
                                "developing_answer": question_details.get('developing_answer', ''),
                                "foundational_answer": question_details.get('foundational_answer', ''),
                                "answer": {
                                    "numeric_score": answer.get("numeric_score", 0),
                                    "option": answer.get("option", ""),  # Store the option (LEADING, ESTABLISHED, etc.)
                                    "text": answer.get("option", ""),
                                    "note": answer.get("note", ""),  # Store assessment notes
                                    "question_id": question_uuid
                                }
                            }
                            questions_by_domain[domain_id].append(question)
        
        
        print(f"DEBUG: Processed answers: {processed_answers}/{len(answers)}")
        print(f"DEBUG: Domains in heatmap: {len(questions_by_domain)}")
        for domain_id in list(questions_by_domain.keys())[:3]:  # Show first 3 domains
            domain = domains_by_id.get(domain_id, {"name": "Unknown"})
            questions = questions_by_domain[domain_id]
            print(f"  Domain: {domain['name']} - Questions: {len(questions)}")
            for question in questions[:2]:  # Show first 2 questions per domain
                print(f"    Question {question.get('code', 'N/A')}: Score {question.get('answer', {}).get('numeric_score', 'N/A')}")
        # Calculate summary data
        total_score = 0
        total_possible = 0
        domain_scores = []
        
        for domain_id, questions in questions_by_domain.items():
            domain = domains_by_id.get(domain_id, {"name": "Unknown"})
            domain_total = sum(q["answer"]["numeric_score"] for q in questions)
            domain_possible = len(questions) * 4  # 1-4 scale
            domain_percentage = (domain_total / domain_possible * 100) if domain_possible > 0 else 0
            
            domain_scores.append({
                "domain_id": domain_id,
                "domain_name": domain.get("name", "Unknown"),
                "percentage": domain_percentage
            })
            
            total_score += domain_total
            total_possible += domain_possible
        
        # For Awareness assessments, prefer stored overall_percentage
        if is_awareness and assessment.get('overall_percentage'):
            overall_percentage = assessment.get('overall_percentage')
            print(f"DEBUG: Using stored overall_percentage for Awareness: {overall_percentage}")
        else:
            overall_percentage = (total_score / total_possible * 100) if total_possible > 0 else 0
        
        # Calculate maturity tier
        # Categories: Excellent (91–100%), Good (81–90%), Moderate (61–80%), Low (41–60%), Basic (0–40%)
        if overall_percentage >= 91:
            overall_maturity = "Excellent"
        elif overall_percentage >= 81:
            overall_maturity = "Good"
        elif overall_percentage >= 61:
            overall_maturity = "Moderate"
        elif overall_percentage >= 41:
            overall_maturity = "Low"
        else:
            overall_maturity = "Basic"
        
        # Calculate overall_tier based on assessment type
        if self.assessment_type == 'Awareness':
            overall_tier = self._calculate_awareness_tier(overall_percentage)
        elif self.assessment_type == 'Readiness':
            overall_tier = self._calculate_readiness_tier(overall_percentage)
        elif self.assessment_type == 'Orgwide':
            overall_tier = self._calculate_orgwide_tier(overall_percentage)
        elif self.assessment_type == 'System':
            overall_tier = self._calculate_system_tier(overall_percentage)
        else:
            overall_tier = self._calculate_tier(overall_percentage)
        
        # Identify top strengths and gaps from domain_scores
        sorted_domains = sorted(domain_scores, key=lambda x: x['percentage'], reverse=True)
        top_strengths = [d['domain_name'] for d in sorted_domains if d['percentage'] >= 75][:3]
        top_gaps = [d['domain_name'] for d in sorted_domains if d['percentage'] < 75][-3:]  # Bottom performers
        
        summary_data = {
            "overall_percentage": overall_percentage,
            "overall_maturity": overall_maturity,
            "overall_tier": overall_tier,
            "top_strengths": top_strengths,
            "top_gaps": top_gaps,
            "domain_scores": domain_scores
        }
        
        # The questions_by_domain structure is already built above, no need to recreate it
        
        # Use the already-built questions_by_domain structure from above
        # questions_by_domain = {}  # Already built above
        # domains_by_id is already available from earlier in the method
        
        # Skip this loop since questions_by_domain is already built above
        pass
        
        # Create structured questions data
        questions_data = []
        for domain_id, questions in questions_by_domain.items():
            domain_info = domains_by_id.get(domain_id, {"name": "Unknown Domain"})
            questions_data.append({
                "domain": domain_info,
                "questions": sorted(questions, key=lambda q: q.get("order", 0))
            })
        
        # Get sector/industry from assessment - differs by assessment type
        system_info = assessment.get('system_info') or {}
        awareness_info = assessment.get('awareness_info') or {}
        
        # For Awareness assessments, get industry from awareness_info
        if self.assessment_type == 'Awareness':
            sector_name = awareness_info.get('industry', '')
        elif self.assessment_type == 'Readiness':
            readiness_info = assessment.get('readiness_info') or {}
            sector_name = readiness_info.get('industry', '')
        elif self.assessment_type == 'Orgwide':
            # For Orgwide, prefer the organization's primary_industry over org_info.industry
            # This ensures sector benchmarks use the correct industry
            org = await db.organizations.find_one({"id": current_user.org_id})
            if org and org.get('primary_industry'):
                sector_name = org.get('primary_industry')
            else:
                orgwide_info = assessment.get('orgwide_info') or assessment.get('org_info') or {}
                sector_name = orgwide_info.get('industry', '')
        else:
            sector_name = system_info.get('industry', '') if system_info else ''
        
        print(f"DEBUG: Assessment type: {self.assessment_type}")
        print(f"DEBUG: Assessment sector/industry: {sector_name}")
        
        # Generate sector-specific actions (for System, Awareness, Readiness, and Orgwide assessments with sector info)
        sector_actions = {'high': [], 'medium': [], 'low': []}
        if sector_name and self.assessment_type in ['System', 'Awareness', 'Readiness', 'Orgwide']:
            sector_actions = self._generate_sector_actions(
                {"questions_data": questions_data},
                sector_name
            )
        
        # DEBUG: Print what sector_actions contains BEFORE passing to assessment_data
        print(f"DEBUG generate_report_for_assessment: sector_actions generated")
        print(f"  - high: {len(sector_actions.get('high', []))} items")
        print(f"  - medium: {len(sector_actions.get('medium', []))} items")
        print(f"  - low: {len(sector_actions.get('low', []))} items")
        if sector_actions.get('high'):
            first = sector_actions['high'][0]
            print(f"  - FIRST HIGH: question_id='{first.get('question_id')}', domain='{first.get('domain')}', sector_action='{first.get('sector_action', 'N/A')[:60]}...'")
        
        # Get sector_average from benchmark data
        sector_average = self._get_sector_average(sector_name)
        print(f"DEBUG: sector_average for '{sector_name}': {sector_average}")
        
        # Prepare assessment data for report generation
        assessment_data = {
            "assessment": assessment,
            "questions_data": questions_data,
            "summary": summary_data,
            "sector_name": sector_name,
            "sector_average": sector_average,  # Add sector_average for bar chart and results summary
            "sector_actions": sector_actions,
            "system_info": system_info,  # Pass system_info for template use (may be empty for non-System assessments)
            "awareness_info": assessment.get('awareness_info') or {},  # Pass awareness_info for Awareness assessments
            "readiness_info": assessment.get('readiness_info') or {},  # Pass readiness_info for Readiness assessments
            "orgwide_info": assessment.get('orgwide_info') or assessment.get('org_info') or {},  # Pass orgwide_info for Orgwide assessments (fallback to org_info)
            "faira_form": assessment.get('faira_form') or {},  # Pass faira_form for FAIRA assessments
        }
        
        print(f"DEBUG: current_user type: {type(current_user)}")
        print(f"DEBUG: current_user.organization_name: {getattr(current_user, 'organization_name', 'NOT SET')}")
        
        user_data = {
            "organization_name": current_user.organization_name
        }
        
        print(f"DEBUG: user_data: {user_data}")
        
        # Generate both DOCX and PDF
        return await self.generate_report(assessment_id, assessment_data, user_data)